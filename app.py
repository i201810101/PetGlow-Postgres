from flask import Flask, render_template, redirect, url_for, flash, request, session, jsonify, request, session
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
from xml.etree.ElementTree import Element, SubElement, tostring
import re
import resend
import json
import os
import smtplib
from flask import current_app
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv
import psycopg2
from psycopg2 import Error
import psycopg2.extras
import hashlib
from functools import wraps


# Importar configuración y base de datos
from config.config import config
from config.database import db

# Cargar variables de entorno
load_dotenv()

# Inicializar la aplicación
app = Flask(__name__)
app.secret_key = config.SECRET_KEY
app.config['SECRET_KEY'] = config.SECRET_KEY
app.config['SQLALCHEMY_DATABASE_URI'] = config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = config.SQLALCHEMY_TRACK_MODIFICATIONS

# Configuración de correo directamente de config
app.config['MAIL_SERVER'] = config.MAIL_SERVER
app.config['MAIL_PORT'] = config.MAIL_PORT
app.config['MAIL_USERNAME'] = config.MAIL_USERNAME
app.config['MAIL_PASSWORD'] = config.MAIL_PASSWORD
app.config['MAIL_DEFAULT_SENDER'] = config.MAIL_DEFAULT_SENDER
app.config['ADMIN_EMAIL'] = config.ADMIN_EMAIL

# TLS/SSL según configuración
app.config['MAIL_USE_TLS'] = config.MAIL_USE_TLS
app.config['MAIL_USE_SSL'] = config.MAIL_USE_SSL

# 🔥 NUEVO: Configurar Resend
app.config['RESEND_API_KEY'] = config.RESEND_API_KEY

# 🔥 Añade esto también para inicializar Resend
try:
    if config.RESEND_API_KEY:
        resend.api_key = config.RESEND_API_KEY
        print(f"✅ Resend configurado con API Key: {config.RESEND_API_KEY[:10]}...")
    else:
        print("⚠️ RESEND_API_KEY no configurada en config.py")
except Exception as e:
    print(f"⚠️ Error configurando Resend: {e}")
    
def parse_xml(xml_string):
    """Parse XML string to Element"""
    from xml.etree.ElementTree import fromstring
    return fromstring(xml_string)

def nsdecls(*prefixes):
    """Generate namespace declarations for Word XML"""
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    }
    
    declarations = []
    for prefix in prefixes:
        if prefix in nsmap:
            declarations.append(f'xmlns:{prefix}="{nsmap[prefix]}"')
    
    return ' '.join(declarations)

def login_required(f):
    """Decorador para requerir inicio de sesión"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'id_usuario' not in session:
            flash('Por favor inicia sesión para acceder a esta página.', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    """Decorador para requerir rol de administrador"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'id_usuario' not in session:
            flash('Por favor inicia sesión.', 'warning')
            return redirect(url_for('login'))
        if session.get('rol') != 'admin':
            flash('No tienes permisos para acceder a esta página.', 'danger')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function



# Configuración de la base de datos MySQL
def get_db_connection():
    try:
        DATABASE_URL = os.getenv("DATABASE_URL")

        conn = psycopg2.connect(
            DATABASE_URL,
            cursor_factory=psycopg2.extras.DictCursor
        )

        return conn

    except Exception as e:
        print("ERROR CONEXION POSTGRES:", e)
        return None


# Middleware para establecer valores por defecto en sesión
@app.before_request
def set_default_session():
    """Establecer valores por defecto en sesión si no existen"""
    if 'rol' not in session:
        session['rol'] = 'admin'  # O 'cajero'
    if 'nombre' not in session:
        session['nombre'] = 'Administrador'
    if 'id_empleado' not in session:
        session['id_empleado'] = 1  # ID del empleado admin
    
    # Pasar la fecha/hora actual a todos los templates
    from datetime import datetime
    app.jinja_env.globals['now'] = datetime.now
    app.jinja_env.globals.update(max=max, min=min, abs=abs, round=round)
# ================= RUTAS =================

@app.route('/')
def index():
    """Página principal - Redirige directamente al dashboard"""
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
@login_required
def dashboard():
    """Panel de control principal"""
    conn = get_db_connection()
    stats = {
        'total_clientes': 15,
        'total_mascotas': 42,
        'reservas_hoy': 5,
        'ventas_hoy': 850.0
    }
    
    ultimas_reservas = []  # ← Añadir esta línea
    
    if conn:
        try:
            cursor = conn.cursor()
            
            # Verificar si las tablas existen
            cursor.execute("SELECT to_regclass( public.clientes)")
            if cursor.fetchone():
                # Total clientes
                cursor.execute("SELECT COUNT(*) as total FROM clientes")
                result = cursor.fetchone()
                stats['total_clientes'] = result['total'] if result else 15
            
            cursor.execute("SELECT to_regclass (public.mascotas)")
            if cursor.fetchone():
                # Total mascotas
                cursor.execute("SELECT COUNT(*) as total FROM mascotas")
                result = cursor.fetchone()
                stats['total_mascotas'] = result['total'] if result else 42
            
            cursor.execute("SELECT to_regclass( public.reservas)")
            if cursor.fetchone():
                # Reservas de hoy
                cursor.execute("SELECT COUNT(*) as total FROM reservas WHERE DATE(fecha_reserva) = CURRENT_DATE")
                result = cursor.fetchone()
                stats['reservas_hoy'] = result['total'] if result else 5
                
                # OBTENER ÚLTIMAS RESERVAS ← Añadir esta consulta
                cursor.execute("""
                    SELECT r.*, 
                           m.nombre as mascota_nombre, m.especie,
                           c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                           e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                           string_agg(DISTINCT s.nombre SEPARATOR ', ') as servicios_nombres
                    FROM reservas r
                    JOIN mascotas m ON r.id_mascota = m.id_mascota
                    JOIN clientes c ON m.id_cliente = c.id_cliente
                    JOIN empleados e ON r.id_empleado = e.id_empleado
                    LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
                    LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
                    GROUP BY r.id_reserva
                    ORDER BY r.fecha_reserva DESC
                    LIMIT 10
                """)
                ultimas_reservas = cursor.fetchall()  # ← Obtener las reservas reales
            
            cursor.execute("SELECT to_regclass( public.facturas)")
            if cursor.fetchone():
                # Ventas de hoy
                cursor.execute("""
                    SELECT COALESCE(SUM(total), 0) as total 
                    FROM facturas 
                    WHERE DATE(fecha_emision) = CURRENT_DATE AND estado = 'pagada'
                """)
                result = cursor.fetchone()
                stats['ventas_hoy'] = float(result['total']) if result and result['total'] else 850.0
            
            cursor.close()
            
        except Error as e:
            print(f"⚠️  Error obteniendo estadísticas: {e}")
            # Mantener valores demo
        finally:
            conn.close()
    
    # Pasar las últimas reservas a la plantilla
    return render_template('dashboard.html', ultimas_reservas=ultimas_reservas, **stats)

def hash_password(password):
    """Generar hash seguro para contraseñas"""   
    return generate_password_hash(password)

def verify_password(stored_hash, provided_password):
    return check_password_hash(stored_hash, provided_password)

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Página de inicio de sesión - CON DEBUG"""
    if 'id_usuario' in session:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        remember = request.form.get('remember') == 'on'
        
        print(f"DEBUG: Usuario intentando login: {username}")
        
        if not username or not password:
            flash('Usuario y contraseña son requeridos.', 'danger')
            return render_template('login/login.html')
        
        conn = get_db_connection()
        if not conn:
            flash('Error de conexión a la base de datos.', 'danger')
            return render_template('login/login.html')
        
        cursor = None
        try:
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT u.*, e.id_empleado, e.nombre, e.apellido, e.email as empleado_email
                FROM usuarios u
                LEFT JOIN empleados e ON u.id_empleado = e.id_empleado
                WHERE u.username = %s AND u.activo = TRUE
            """, (username,))
            
            usuario = cursor.fetchone()
            
            if not usuario:
                print(f"DEBUG: Usuario '{username}' no encontrado")
                flash('Usuario o contraseña incorrectos.', 'danger')
                return render_template('login/login.html')
            
            print(f"DEBUG: Usuario encontrado: {usuario['username']}")
            
            # Verificar contraseña
            stored_hash = usuario['password_hash']
            
            # Si el hash empieza con 'sha256$'
            if stored_hash.startswith('sha256$'):
                print("DEBUG: Hash detectado como formato 'sha256$'")
                hash_parts = stored_hash.split('$')
                if len(hash_parts) == 2:
                    stored_hash_only = hash_parts[1]
                    print(f"DEBUG: Hash puro: {stored_hash_only}")
                    
                    # Calcular hash de la contraseña ingresada
                    import hashlib
                    input_hash = hashlib.sha256(password.encode()).hexdigest()
                    print(f"DEBUG: Hash calculado de input: {input_hash}")
                    print(f"DEBUG: ¿Son iguales? {input_hash == stored_hash_only}")
                    
                    if input_hash != stored_hash_only:
                        print("DEBUG: Hash NO coincide")
                        flash('Usuario o contraseña incorrectos.', 'danger')
                        return render_template('login/login.html')
                else:
                    print("DEBUG: Formato de hash inválido")
                    flash('Formato de hash inválido.', 'danger')
                    return render_template('login/login.html')
            else:
                # Si no es formato sha256$, comparar directamente
                print("DEBUG: Hash NO tiene formato sha256$")
                import hashlib
                input_hash = hashlib.sha256(password.encode()).hexdigest()
                if input_hash != stored_hash:
                    print("DEBUG: Hash NO coincide")
                    flash('Usuario o contraseña incorrectos.', 'danger')
                    return render_template('login/login.html')
            
            # Si llega aquí, la contraseña es correcta
            print("DEBUG: Contraseña VERIFICADA correctamente")
            
            # 1. ACTUALIZAR último login en tabla usuarios
            print("DEBUG: Actualizando ultimo_login...")
            cursor.execute("""
                UPDATE usuarios 
                SET ultimo_login = CURRENT_TIMESTAMP 
                WHERE id_usuario = %s
            """, (usuario['id_usuario'],))
            
            # 2. REGISTRAR en login_history
            print("DEBUG: Registrando en login_history...")
            ip_address = request.remote_addr
            user_agent = request.headers.get('User-Agent', '')
            
            cursor.execute("""
                INSERT INTO login_history (id_usuario, ip_address, user_agent)
                VALUES (%s, %s, %s)
            """, (usuario['id_usuario'], ip_address, user_agent))
            
            conn.commit()
            print("DEBUG: Base de datos actualizada correctamente")
            
            # Configurar sesión
            session['id_usuario'] = usuario['id_usuario']
            session['username'] = usuario['username']
            session['rol'] = usuario['rol']
            session['nombre'] = usuario.get('nombre', 'Administrador')
            session['apellido'] = usuario.get('apellido', '')
            session['email'] = usuario.get('empleado_email', '')
            session['id_empleado'] = usuario.get('id_empleado')
            session['last_activity'] = datetime.now().isoformat()
            
            if remember:
                session.permanent = True
                app.permanent_session_lifetime = timedelta(days=7)
            else:
                session.permanent = True
                app.permanent_session_lifetime = timedelta(hours=8)
            
            flash(f'¡Bienvenido(a), {session["nombre"]}!', 'success')
            
            if usuario['rol'] == 'admin':
                return redirect(url_for('dashboard'))
            elif usuario['rol'] == 'gerente':
                return redirect(url_for('reservas'))
            elif usuario['rol'] == 'cajero':
                return redirect(url_for('ventas'))
            else:
                return redirect(url_for('dashboard'))
            
        except Exception as e:
            print(f"DEBUG: ERROR: {str(e)}")
            flash(f'Error en el inicio de sesión: {str(e)}', 'danger')
            return render_template('login/login.html')
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
    
    return render_template('login/login.html')
@app.route('/logout')
def logout():
    """Cerrar sesión"""
    # Registrar logout
    if 'id_usuario' in session:
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE usuarios 
                    SET ultimo_logout = NOW()
                    WHERE id_usuario = %s
                """, (session['id_usuario'],))
                conn.commit()
            except:
                pass
            finally:
                cursor.close()
                conn.close()
    
    # Limpiar sesión (esto eliminará también el acceso a reportes)
    session.clear()
    flash('Sesión cerrada exitosamente.', 'info')
    return redirect(url_for('login'))

# ==================== RUTA DE VERIFICACIÓN DE SESIÓN ====================

@app.before_request
def before_request():
    """Verificar sesión antes de cada petición"""
    # Excluir rutas públicas
    public_routes = ['login', 'static']
    if request.endpoint in public_routes:
        return
    
    # Verificar si hay sesión activa
    if 'id_usuario' not in session:
        return redirect(url_for('login'))
    
    # Verificar timeout de sesión (8 horas)
    if 'last_activity' in session:
        last_activity = datetime.fromisoformat(session['last_activity'])
        time_difference = datetime.now() - last_activity
        
        if time_difference.total_seconds() > 28800:  # 8 horas en segundos
            session.clear()
            flash('Tu sesión ha expirado por inactividad.', 'warning')
            return redirect(url_for('login'))
    
    # Actualizar última actividad
    session['last_activity'] = datetime.now().isoformat()

   
@app.route('/ventas')
def ventas():
    """Listar ventas/facturas"""
    conn = get_db_connection()
    ventas_list = []
    
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT f.*, 
                       c.nombre as cliente_nombre, 
                       c.apellido as cliente_apellido,
                       r.codigo_reserva
                FROM facturas f
                LEFT JOIN clientes c ON f.id_cliente = c.id_cliente
                LEFT JOIN reservas r ON f.id_reserva = r.id_reserva
                ORDER BY f.fecha_emision DESC 
                LIMIT 50
            """)
            ventas_raw = cursor.fetchall()  # Cambiar nombre
            
            # Convertir cada RealDictRow a dict mutable
            for venta in ventas_raw:
                venta_dict = dict(venta)  # ← CONVERTIR A DICT
                
                estado_clases = {
                    'pendiente': 'bg-warning text-dark',
                    'pagada': 'bg-success',
                    'anulada': 'bg-danger',
                    'credito': 'bg-info text-dark'
                }
                venta_dict['estado_clase'] = estado_clases.get(venta_dict['estado'], 'bg-secondary')
                venta_dict['estado_texto'] = venta_dict['estado'].capitalize()
                
                # Añadir icono según tipo
                if venta_dict['tipo_comprobante'] == 'factura':
                    venta_dict['tipo_icono'] = 'fa-file-invoice'
                    venta_dict['tipo_color'] = 'text-primary'
                else:
                    venta_dict['tipo_icono'] = 'fa-receipt'
                    venta_dict['tipo_color'] = 'text-success'
                
                ventas_list.append(venta_dict)  # Agregar dict mutable
            
            cursor.close()
        except Error as e:
            flash(f'Error obteniendo ventas: {e}', 'danger')
        finally:
            conn.close()
    
    return render_template('ventas/listar.html', ventas=ventas_list)
    
@app.route('/clientes')
@login_required
def clientes():
    """Listar clientes con paginación"""
    # Obtener parámetro de página (default: 1)
    page = request.args.get('page', 1, type=int)
    per_page = 10  # Número de clientes por página
    
    conn = get_db_connection()
    clientes_list = []
    total_clientes = 0
    total_pages = 0
    
    if conn:
        try:
            cursor = conn.cursor()
            
            # 1. Contar total de clientes
            cursor.execute("SELECT COUNT(*) as total FROM clientes")
            total_clientes = cursor.fetchone()['total']
            
            # 2. Calcular total de páginas
            total_pages = (total_clientes + per_page - 1) // per_page
            
            # 3. Calcular offset para la paginación
            offset = (page - 1) * per_page
            
            # 4. Obtener clientes paginados
            cursor.execute("""
                SELECT * FROM clientes 
                ORDER BY fecha_registro DESC 
                LIMIT %s OFFSET %s
            """, (per_page, offset))
            clientes_list = cursor.fetchall()
            
            cursor.close()
        except Error as e:
            flash(f'Error obteniendo clientes: {e}', 'danger')
        finally:
            conn.close()
    else:
        # Datos demo (con paginación simulada)
        total_clientes = 25  # Simulamos 25 clientes
        per_page = 10
        total_pages = 3
        
        clientes_demo = [
            {'id_cliente': i, 'nombre': f'Cliente {i}', 'apellido': f'Apellido {i}', 
             'telefono': f'555-{i:04d}', 'email': f'cliente{i}@email.com', 
             'fecha_registro': datetime.now(), 'dni': f'7{i:07d}' if i % 3 != 0 else None}
            for i in range((page-1)*per_page + 1, min(page*per_page, total_clientes) + 1)
        ]
        clientes_list = clientes_demo
    
    return render_template('clientes/listar.html', 
                         clientes=clientes_list,
                         page=page,
                         per_page=per_page,
                         total_clientes=total_clientes,
                         total_pages=total_pages)

@app.route('/clientes/crear', methods=['GET', 'POST'])
def crear_cliente():
    """Crear nuevo cliente"""
    if request.method == 'POST':
        # Obtener datos del formulario (¡FALTAN CAMPOS!)
        dni = request.form.get('dni', '').strip()
        nombre = request.form.get('nombre', '').strip()
        apellido = request.form.get('apellido', '').strip()
        telefono = request.form.get('telefono', '').strip()
        email = request.form.get('email', '').strip()
        
        # ¡AGREGA ESTOS CAMPOS QUE FALTAN!
        direccion = request.form.get('direccion', '').strip()
        notas = request.form.get('notas', '').strip()
        
        if not nombre or not telefono:
            flash('Nombre y teléfono son obligatorios.', 'danger')
            return render_template('clientes/crear.html')
        
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                # ¡ACTUALIZA EL INSERT CON TODOS LOS CAMPOS!
                cursor.execute("""
                    INSERT INTO clientes (dni, nombre, apellido, telefono, email, direccion, notas)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (dni or None, nombre, apellido, telefono, 
                      email or None, direccion or None, notas or None))
                conn.commit()
                
                flash(f'Cliente {nombre} {apellido} creado exitosamente.', 'success')
                return redirect(url_for('clientes'))
                
            except Error as e:
                flash(f'Error creando cliente: {e}', 'danger')
                # Si hay error, volver al formulario con los datos ingresados
                return render_template('clientes/crear.html', 
                                      dni=dni, nombre=nombre, apellido=apellido,
                                      telefono=telefono, email=email,
                                      direccion=direccion, notas=notas)
            finally:
                cursor.close()
                conn.close()
        else:
            flash('No hay conexión a la base de datos. Cliente guardado en modo demo.', 'warning')
            return redirect(url_for('clientes'))
    
    return render_template('clientes/crear.html')

# Añade estas rutas después de la ruta crear_cliente en app.py

@app.route('/clientes/editar/<int:id>', methods=['GET', 'POST'])
def editar_cliente(id):
    """Editar cliente existente"""
    conn = get_db_connection()
    cliente = None
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('clientes'))
    
    try:
        cursor = conn.cursor()
        
        if request.method == 'POST':
            # Obtener TODOS los datos del formulario
            dni = request.form.get('dni', '').strip()
            nombre = request.form.get('nombre', '').strip()
            apellido = request.form.get('apellido', '').strip()
            telefono = request.form.get('telefono', '').strip()
            email = request.form.get('email', '').strip()
            direccion = request.form.get('direccion', '').strip()
            notas = request.form.get('notas', '').strip()
            
            # Validaciones
            if not nombre or not telefono:
                flash('Nombre y teléfono son obligatorios.', 'danger')
                # Volver a la edición con los datos ingresados
                return render_template('clientes/editar.html', 
                                      cliente={'id_cliente': id, 'dni': dni, 'nombre': nombre, 
                                               'apellido': apellido, 'telefono': telefono,
                                               'email': email, 'direccion': direccion, 'notas': notas})
            
            try:
                # Actualizar cliente con TODOS los campos
                cursor.execute("""
                    UPDATE clientes 
                    SET dni = %s, nombre = %s, apellido = %s, 
                        telefono = %s, email = %s, direccion = %s,
                        notas = %s
                    WHERE id_cliente = %s
                """, (dni or None, nombre, apellido, telefono, 
                      email or None, direccion or None, notas or None, id))
                
                conn.commit()
                
                # Verificar si se actualizó
                if cursor.rowcount > 0:
                    flash(f'Cliente {nombre} {apellido} actualizado exitosamente.', 'success')
                else:
                    flash('No se realizaron cambios o el cliente no existe.', 'warning')
                
                return redirect(url_for('ver_cliente', id=id))
                
            except Error as e:
                conn.rollback()
                flash(f'Error actualizando cliente: {str(e)}', 'danger')
                # Volver a la edición con los datos ingresados
                return render_template('clientes/editar.html', 
                                      cliente={'id_cliente': id, 'dni': dni, 'nombre': nombre, 
                                               'apellido': apellido, 'telefono': telefono,
                                               'email': email, 'direccion': direccion, 'notas': notas})
        
        # GET request: Obtener datos del cliente para mostrar en el formulario
        cursor.execute("SELECT * FROM clientes WHERE id_cliente = %s", (id,))
        cliente = cursor.fetchone()
        
        if not cliente:
            flash('Cliente no encontrado.', 'danger')
            return redirect(url_for('clientes'))
        
    except Error as e:
        flash(f'Error obteniendo datos del cliente: {str(e)}', 'danger')
        return redirect(url_for('clientes'))
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
    
    return render_template('clientes/editar.html', cliente=cliente)

@app.route('/clientes/eliminar/<int:id>', methods=['POST'])
def eliminar_cliente(id):
    """Eliminar cliente (soft delete o hard delete)"""
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el cliente tiene mascotas asociadas
        cursor.execute("SELECT COUNT(*) as total FROM mascotas WHERE id_cliente = %s", (id,))
        result = cursor.fetchone()
        
        if result and result['total'] > 0:
            return jsonify({
                'success': False, 
                'message': f'No se puede eliminar el cliente porque tiene {result["total"]} mascota(s) asociada(s). Primero elimina o transfiere las mascotas.'
            })
        
        # Si no tiene mascotas, proceder con eliminación
        cursor.execute("DELETE FROM clientes WHERE id_cliente = %s", (id,))
        conn.commit()
        
        if cursor.rowcount > 0:
            return jsonify({'success': True, 'message': 'Cliente eliminado exitosamente.'})
        else:
            return jsonify({'success': False, 'message': 'Cliente no encontrado.'}), 404
            
    except Error as e:
        return jsonify({'success': False, 'message': f'Error eliminando cliente: {str(e)}'}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# También puedes agregar una ruta para ver detalles del cliente
@app.route('/clientes/ver/<int:id>')
def ver_cliente(id):
    """Ver detalles del cliente"""
    conn = get_db_connection()
    cliente = None
    mascotas = []
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('clientes'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener datos del cliente
        cursor.execute("SELECT * FROM clientes WHERE id_cliente = %s", (id,))
        cliente = cursor.fetchone()
        
        if not cliente:
            flash('Cliente no encontrado.', 'danger')
            return redirect(url_for('clientes'))
        
        # Obtener mascotas del cliente
        cursor.execute("SELECT * FROM mascotas WHERE id_cliente = %s", (id,))
        mascotas = cursor.fetchall()
        
    except Error as e:
        flash(f'Error obteniendo datos del cliente: {e}', 'danger')
        return redirect(url_for('clientes'))
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
    
    return render_template('clientes/ver.html', cliente=cliente, mascotas=mascotas)


@app.route('/mascotas')
def mascotas():
    """Listar mascotas con paginación"""
    # Obtener parámetro de página (default: 1)
    page = request.args.get('page', 1, type=int)
    per_page = 10  # Número de mascotas por página
    
    conn = get_db_connection()
    mascotas_list = []
    
    if conn:
        try:
            cursor = conn.cursor()
            
            # Contar total de mascotas
            cursor.execute("SELECT COUNT(*) as total FROM mascotas")
            total_result = cursor.fetchone()
            total_mascotas = total_result['total'] if total_result else 0
            
            # Calcular total de páginas
            total_pages = (total_mascotas + per_page - 1) // per_page
            
            # Calcular offset para la paginación
            offset = (page - 1) * per_page
            
            # Obtener mascotas paginadas
            cursor.execute("""
                SELECT m.*, c.nombre as cliente_nombre, c.apellido as cliente_apellido
                FROM mascotas m
                LEFT JOIN clientes c ON m.id_cliente = c.id_cliente
                ORDER BY m.fecha_registro DESC 
                LIMIT %s OFFSET %s
            """, (per_page, offset))
            mascotas_list = cursor.fetchall()
            
            cursor.close()
        except Error as e:
            flash(f'Error obteniendo mascotas: {e}', 'danger')
        finally:
            conn.close()
    else:
        # Datos demo
        total_mascotas = 3
        total_pages = 1
        mascotas_list = [
            {'id_mascota': 1, 'nombre': 'Max', 'especie': 'perro', 'raza': 'Labrador', 'cliente_nombre': 'Juan', 'cliente_apellido': 'Pérez', 'fecha_registro': datetime.now()},
            {'id_mascota': 2, 'nombre': 'Luna', 'especie': 'gato', 'raza': 'Siamés', 'cliente_nombre': 'María', 'cliente_apellido': 'García', 'fecha_registro': datetime.now()},
            {'id_mascota': 3, 'nombre': 'Rocky', 'especie': 'perro', 'raza': 'Bulldog', 'cliente_nombre': 'Carlos', 'cliente_apellido': 'Ruiz', 'fecha_registro': datetime.now()}
        ]
    
    return render_template('mascotas/listar.html', 
                         mascotas=mascotas_list,
                         page=page,
                         per_page=per_page,
                         total_mascotas=total_mascotas,
                         total_pages=total_pages)
    

@app.route('/mascotas/crear', methods=['GET', 'POST'])
def crear_mascota():
    """Crear nueva mascota"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('mascotas'))
    
    if request.method == 'POST':
        try:
            # Obtener datos del formulario
            id_cliente = request.form.get('id_cliente', '').strip()
            nombre = request.form.get('nombre', '').strip()
            especie = request.form.get('especie', 'perro').strip()
            raza = request.form.get('raza', '').strip()
            tamano = request.form.get('tamano', '').strip()
            fecha_nacimiento = request.form.get('fecha_nacimiento', '').strip()
            peso = request.form.get('peso', '').strip()
            color = request.form.get('color', '').strip()
            corte = request.form.get('corte', '').strip()
            caracteristicas = request.form.get('caracteristicas', '').strip()
            alergias = request.form.get('alergias', '').strip()
            
            # Validaciones básicas
            if not id_cliente or not nombre:
                flash('Cliente y nombre de mascota son obligatorios.', 'danger')
                return redirect(url_for('crear_mascota'))
            
            # Convertir fecha
            fecha_nacimiento_dt = None
            if fecha_nacimiento:
                try:
                    fecha_nacimiento_dt = datetime.strptime(fecha_nacimiento, '%Y-%m-%d')
                except ValueError:
                    flash('Formato de fecha inválido. Use YYYY-MM-DD.', 'danger')
                    return redirect(url_for('crear_mascota'))
            
            # Convertir peso
            peso_float = None
            if peso:
                try:
                    peso_float = float(peso)
                except ValueError:
                    flash('Peso debe ser un número válido.', 'danger')
                    return redirect(url_for('crear_mascota'))
            
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO mascotas (
                    id_cliente, nombre, especie, raza, tamano, 
                    fecha_nacimiento, peso, color,corte, caracteristicas, alergias
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                id_cliente, nombre, especie, raza or None, tamano or None,
                fecha_nacimiento_dt, peso_float, color or None, corte or None,
                caracteristicas or None, alergias or None
            ))
            conn.commit()
            
            flash(f'Mascota {nombre} creada exitosamente.', 'success')
            return redirect(url_for('mascotas'))
            
        except Error as e:
            flash(f'Error creando mascota: {e}', 'danger')
        finally:
            cursor.close()
            conn.close()
    
    # GET: Obtener clientes para el select
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT id_cliente, nombre, apellido FROM clientes ORDER BY apellido, nombre")
        clientes = cursor.fetchall()
        cursor.close()
    except Error as e:
        flash(f'Error obteniendo clientes: {e}', 'danger')
        clientes = []
    finally:
        conn.close()
    
    return render_template('mascotas/crear.html', clientes=clientes)

@app.route('/mascotas/editar/<int:id>', methods=['GET', 'POST'])
def editar_mascota(id):
    """Editar mascota existente"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('mascotas'))
    
    try:
        cursor = conn.cursor()
        
        if request.method == 'POST':
            # Obtener datos del formulario
            id_cliente = request.form.get('id_cliente', '').strip()
            nombre = request.form.get('nombre', '').strip()
            especie = request.form.get('especie', 'perro').strip()
            raza = request.form.get('raza', '').strip()
            tamano = request.form.get('tamano', '').strip()
            fecha_nacimiento = request.form.get('fecha_nacimiento', '').strip()
            peso = request.form.get('peso', '').strip()
            color = request.form.get('color', '').strip()
            corte = request.form.get('corte', '').strip()  # Asegurar captura del corte
            caracteristicas = request.form.get('caracteristicas', '').strip()
            alergias = request.form.get('alergias', '').strip()
            
            # Validaciones básicas
            if not id_cliente or not nombre:
                flash('Cliente y nombre de mascota son obligatorios.', 'danger')
                return redirect(url_for('editar_mascota', id=id))
            
            # Convertir fecha
            fecha_nacimiento_dt = None
            if fecha_nacimiento:
                try:
                    fecha_nacimiento_dt = datetime.strptime(fecha_nacimiento, '%Y-%m-%d')
                except ValueError:
                    flash('Formato de fecha inválido. Use YYYY-MM-DD.', 'danger')
                    return redirect(url_for('editar_mascota', id=id))
            
            # Convertir peso
            peso_float = None
            if peso:
                try:
                    peso_float = float(peso)
                except ValueError:
                    flash('Peso debe ser un número válido.', 'danger')
                    return redirect(url_for('editar_mascota', id=id))
            
            # Obtener el corte actual ANTES de actualizar
            cursor.execute("SELECT corte, nombre FROM mascotas WHERE id_mascota = %s", (id,))
            mascota_actual = cursor.fetchone()
            
            if not mascota_actual:
                flash('Mascota no encontrada.', 'danger')
                return redirect(url_for('mascotas'))
            
            corte_anterior = mascota_actual['corte'] if mascota_actual else None
            nombre_mascota = mascota_actual['nombre']
            
            # Actualizar mascota
            cursor.execute("""
                UPDATE mascotas 
                SET id_cliente = %s, nombre = %s, especie = %s, raza = %s, tamano = %s,
                    fecha_nacimiento = %s, peso = %s, color = %s, corte = %s, 
                    caracteristicas = %s, alergias = %s
                WHERE id_mascota = %s
            """, (
                id_cliente, nombre, especie, raza or None, tamano or None,
                fecha_nacimiento_dt, peso_float, color or None, corte or None,
                caracteristicas or None, alergias or None, id
            ))
            
            # Registrar en historial si el corte cambió
            if corte and corte != corte_anterior:
                try:
                    # Obtener id del empleado de la sesión o usar NULL
                    id_empleado = session.get('id_empleado') if 'id_empleado' in session else None
                    
                    descripcion = f"Cambio de corte: {corte_anterior or 'Sin corte'} → {corte}"
                    notas = f"Actualizado al editar mascota"
                    
                    cursor.execute("""
                        INSERT INTO historial_cortes (id_mascota, tipo_corte, descripcion, id_empleado, notas)
                        VALUES (%s, %s, %s, %s, %s)
                    """, (id, corte, descripcion, id_empleado, notas))
                except Exception as e:
                    print(f"Error registrando en historial de cortes: {e}")
                    # No interrumpir por error en historial
            
            conn.commit()
            
            flash(f'Mascota {nombre} actualizada exitosamente.', 'success')
            return redirect(url_for('ver_mascota', id=id))  # Redirigir a ver, no a listar
        
        # GET: Obtener datos de la mascota
        cursor.execute("SELECT * FROM mascotas WHERE id_mascota = %s", (id,))
        mascota_raw = cursor.fetchone()
        
        if not mascota_raw:
            flash('Mascota no encontrada.', 'danger')
            return redirect(url_for('mascotas'))
        
        # Convertir a diccionario
        mascota = dict(mascota_raw)
        
        # Formatear fecha para input type="date"
        if mascota.get('fecha_nacimiento'):
            mascota['fecha_nacimiento_str'] = mascota['fecha_nacimiento'].strftime('%Y-%m-%d')
        else:
            mascota['fecha_nacimiento_str'] = ''
        
        # Asegurar valores para la plantilla
        mascota['raza'] = mascota.get('raza') or ''
        mascota['color'] = mascota.get('color') or ''
        mascota['corte'] = mascota.get('corte') or ''
        mascota['peso'] = mascota.get('peso') or ''
        mascota['caracteristicas'] = mascota.get('caracteristicas') or ''
        mascota['alergias'] = mascota.get('alergias') or ''
        mascota['tamano'] = mascota.get('tamano') or ''
        
        # Obtener clientes para el select
        cursor.execute("SELECT id_cliente, nombre, apellido FROM clientes ORDER BY apellido, nombre")
        clientes_raw = cursor.fetchall()
        
        # Convertir clientes a diccionarios
        clientes = []
        for cliente in clientes_raw:
            clientes.append(dict(cliente))
        
    except Exception as e:
        flash(f'Error editando mascota: {e}', 'danger')
        print(f"Error detallado en editar_mascota: {e}")
        return redirect(url_for('mascotas'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('mascotas/editar.html', mascota=mascota, clientes=clientes)
@app.route('/mascotas/eliminar/<int:id>', methods=['POST'])
def eliminar_mascota(id):
    """Eliminar mascota (soft delete)"""
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si la mascota tiene reservas asociadas
        cursor.execute("SELECT COUNT(*) as total FROM reservas WHERE id_mascota = %s", (id,))
        result = cursor.fetchone()
        
        if result and result[0] > 0:
            return jsonify({
                'success': False, 
                'message': f'No se puede eliminar la mascota porque tiene {result[0]} reserva(s) asociada(s).'
            })
        
        # Soft delete (cambiar activo a false)
        cursor.execute("UPDATE mascotas SET activo = FALSE WHERE id_mascota = %s", (id,))
        conn.commit()
        
        if cursor.rowcount > 0:
            return jsonify({'success': True, 'message': 'Mascota eliminada exitosamente.'})
        else:
            return jsonify({'success': False, 'message': 'Mascota no encontrada.'}), 404
            
    except Error as e:
        return jsonify({'success': False, 'message': f'Error eliminando mascota: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/mascotas/ver/<int:id>')
def ver_mascota(id):
    """Ver detalles de la mascota"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('mascotas'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener datos de la mascota con información del cliente
        cursor.execute("""
            SELECT m.*, 
                   c.id_cliente, c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                   c.telefono as cliente_telefono, c.email as cliente_email
            FROM mascotas m
            LEFT JOIN clientes c ON m.id_cliente = c.id_cliente
            WHERE m.id_mascota = %s
        """, (id,))
        mascota_raw = cursor.fetchone()
        
        if not mascota_raw:
            flash('Mascota no encontrada.', 'danger')
            return redirect(url_for('mascotas'))
        
        # Convertir a diccionario
        mascota = dict(mascota_raw)
        
        # CALCULAR EDAD SI HAY FECHA DE NACIMIENTO
        if mascota.get('fecha_nacimiento'):
            hoy = datetime.now()
            nacimiento = mascota['fecha_nacimiento']
            
            # Calcular años
            años = hoy.year - nacimiento.year
            # Ajustar si aún no ha pasado el cumpleaños este año
            if (hoy.month, hoy.day) < (nacimiento.month, nacimiento.day):
                años -= 1
            
            # Calcular meses restantes
            meses = hoy.month - nacimiento.month
            if hoy.day < nacimiento.day:
                meses -= 1
            if meses < 0:
                meses += 12
            
            mascota['edad_anios'] = años
            mascota['edad_meses'] = meses
        else:
            mascota['edad_anios'] = None
            mascota['edad_meses'] = None
        
        # Asegurar que todos los campos tengan valores
        mascota['raza'] = mascota.get('raza') or ''
        mascota['color'] = mascota.get('color') or ''
        mascota['corte'] = mascota.get('corte') or ''
        mascota['peso'] = mascota.get('peso')
        mascota['caracteristicas'] = mascota.get('caracteristicas') or ''
        mascota['alergias'] = mascota.get('alergias') or ''
        mascota['cliente_telefono'] = mascota.get('cliente_telefono') or ''
        mascota['cliente_email'] = mascota.get('cliente_email') or ''
        
        # Obtener historial de reservas
        try:
            cursor.execute("""
                SELECT r.*, 
                       e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                       string_agg(s.nombre, ', ') as servicios_nombres
                FROM reservas r
                LEFT JOIN empleados e ON r.id_empleado = e.id_empleado
                LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
                LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
                WHERE r.id_mascota = %s
                GROUP BY r.id_reserva, e.id_empleado
                ORDER BY r.fecha_reserva DESC
                LIMIT 10
            """, (id,))
            reservas_raw = cursor.fetchall()
            
            # Convertir a lista de diccionarios
            reservas = []
            for reserva in reservas_raw:
                r = dict(reserva)
                # Asegurar servicios_nombres
                r['servicios_nombres'] = r.get('servicios_nombres') or ''
                reservas.append(r)
        except Exception as e:
            print(f"Advertencia: Error obteniendo reservas: {e}")
            reservas = []
        
        # OBTENER HISTORIAL DE CORTES
        try:
            cursor.execute("""
                SELECT 
                    hc.*,
                    e.nombre || ' ' || e.apellido as empleado_nombre,
                    hc.fecha_registro
                FROM historial_cortes hc
                LEFT JOIN empleados e ON hc.id_empleado = e.id_empleado
                WHERE hc.id_mascota = %s
                ORDER BY hc.fecha_registro DESC
                LIMIT 10
            """, (id,))
            historial_cortes_raw = cursor.fetchall()
            
            # Convertir y formatear
            historial_cortes = []
            for corte in historial_cortes_raw:
                c = dict(corte)
                if c.get('fecha_registro'):
                    c['fecha_formateada'] = c['fecha_registro'].strftime('%d/%m/%Y %H:%M')
                else:
                    c['fecha_formateada'] = 'Fecha no disponible'
                
                # Asegurar campos
                c['descripcion'] = c.get('descripcion') or ''
                c['empleado_nombre'] = c.get('empleado_nombre') or 'No asignado'
                
                historial_cortes.append(c)
        except Exception as e:
            print(f"Advertencia: Error obteniendo historial de cortes: {e}")
            historial_cortes = []
        
    except Exception as e:
        flash(f'Error obteniendo datos de la mascota: {e}', 'danger')
        print(f"Error detallado: {e}")  # Para debugging
        return redirect(url_for('mascotas'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('mascotas/ver.html', 
                         mascota=mascota, 
                         reservas=reservas,
                         historial_cortes=historial_cortes)

def obtener_historial_cortes(id_mascota):
    """Obtener historial de cortes de una mascota - CORREGIDO PARA POSTGRESQL"""
    conn = get_db_connection()
    if not conn:
        return []
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT 
                hc.*,
                CONCAT(e.nombre, ' ', e.apellido) as empleado_nombre,
                hc.fecha_registro
            FROM historial_cortes hc
            LEFT JOIN empleados e ON hc.id_empleado = e.id_empleado
            WHERE hc.id_mascota = %s
            ORDER BY hc.fecha_registro DESC
            LIMIT 20
        """, (id_mascota,))
        
        historial = cursor.fetchall()
        
        # Formatear fechas en Python (no usar DATE_FORMAT de MySQL)
        for corte in historial:
            if corte['fecha_registro']:
                # Formato: día/mes/año hora:minutos
                corte['fecha_formateada'] = corte['fecha_registro'].strftime('%d/%m/%Y %H:%M')
                # Si necesitas fecha_simple y hora por separado
                corte['fecha_simple'] = corte['fecha_registro'].strftime('%Y-%m-%d')
                corte['hora'] = corte['fecha_registro'].strftime('%H:%M:%S')
            else:
                corte['fecha_formateada'] = 'Fecha no disponible'
                corte['fecha_simple'] = ''
                corte['hora'] = ''
        
        return historial
    except Exception as e:  # Cambia Error por Exception
        print(f"Error obteniendo historial de cortes: {e}")
        return []
    finally:
        cursor.close()
        conn.close()

@app.route('/mascotas/<int:id>/registrar-corte', methods=['POST'])
@app.route('/mascotas/<int:id>/registrar-corte', methods=['POST'])
def registrar_corte(id):
    """Registrar un nuevo corte en el historial - CORREGIDO PARA POSTGRESQL"""
    if request.method == 'POST':
        conn = get_db_connection()
        if not conn:
            flash('No hay conexión a la base de datos.', 'danger')
            return redirect(url_for('ver_mascota', id=id))
        
        cursor = None
        try:
            # Obtener datos del formulario
            tipo_corte = request.form.get('tipo_corte', '').strip()
            descripcion = request.form.get('descripcion', '').strip()
            notas = request.form.get('notas', '').strip()
            
            # Validar
            if not tipo_corte:
                flash('El tipo de corte es obligatorio.', 'danger')
                return redirect(url_for('ver_mascota', id=id))
            
            cursor = conn.cursor()
            
            # 1. Actualizar el corte actual en la mascota
            cursor.execute("""
                UPDATE mascotas SET corte = %s WHERE id_mascota = %s
            """, (tipo_corte, id))
            
            # 2. Registrar en el historial
            # Obtener id_empleado de la sesión si existe
            id_empleado = session.get('id_empleado')
            
            cursor.execute("""
                INSERT INTO historial_cortes 
                (id_mascota, tipo_corte, descripcion, id_empleado, notas)
                VALUES (%s, %s, %s, %s, %s)
            """, (id, tipo_corte, descripcion, id_empleado, notas))
            
            conn.commit()
            flash(f'Corte "{tipo_corte}" registrado exitosamente.', 'success')
            
        except Exception as e:  # Cambia Error por Exception
            if conn:
                conn.rollback()
            flash(f'Error registrando corte: {e}', 'danger')
        finally:
            if cursor:
                cursor.close()
            if conn:
                conn.close()
        
        return redirect(url_for('ver_mascota', id=id))

@app.route('/servicios')
def servicios():
    """Listar servicios"""
    conn = get_db_connection()
    servicios_list = []
    
    if conn:
        try:
            cursor = conn.cursor()
            # CORRECCIÓN: Usar TRUE en lugar de 1 para booleanos
            cursor.execute("""
                SELECT * FROM servicios 
                WHERE activo = TRUE  -- Cambiado de 1 a TRUE
                ORDER BY categoria, precio
            """)
            servicios_list = cursor.fetchall()
            
            # Calcular márgenes si no están en la BD
            for servicio in servicios_list:
                if servicio['margen'] is None and servicio['costo'] and servicio['precio']:
                    if servicio['costo'] > 0:
                        servicio['margen'] = ((servicio['precio'] - servicio['costo']) / servicio['costo']) * 100
                    else:
                        servicio['margen'] = 0
            
            cursor.close()
        except Error as e:
            flash(f'Error obteniendo servicios: {e}', 'danger')
        finally:
            conn.close()
    else:
        # Datos demo con márgenes calculados
        servicios_list = [
            {'id_servicio': 1, 'nombre': 'Baño Básico', 'precio': 25.00, 'costo': 10.00, 'margen': 150.00, 'descripcion': 'Baño con shampoo especial', 'categoria': 'baño', 'duracion_min': 45, 'activo': True},
            {'id_servicio': 2, 'nombre': 'Corte de Pelo', 'precio': 35.00, 'costo': 12.00, 'margen': 191.67, 'descripcion': 'Corte profesional', 'categoria': 'corte', 'duracion_min': 60, 'activo': True},
            {'id_servicio': 3, 'nombre': 'Baño + Corte', 'precio': 50.00, 'costo': 20.00, 'margen': 150.00, 'descripcion': 'Servicio completo', 'categoria': 'spa', 'duracion_min': 90, 'activo': True},
            {'id_servicio': 4, 'nombre': 'Limpieza Dental', 'precio': 20.00, 'costo': 8.00, 'margen': 150.00, 'descripcion': 'Limpieza dental especializada', 'categoria': 'salud', 'duracion_min': 30, 'activo': True},
        ]
    
    return render_template('servicios/listar.html', servicios=servicios_list)

# ================= RUTAS DE SERVICIOS =================

@app.route('/servicios/crear', methods=['GET', 'POST'])
def crear_servicio():
    """Crear nuevo servicio"""
    if request.method == 'POST':
        # Obtener datos del formulario
        codigo = request.form.get('codigo', '').strip()
        nombre = request.form.get('nombre', '').strip()
        categoria = request.form.get('categoria', 'baño').strip()
        descripcion = request.form.get('descripcion', '').strip()
        duracion_min = request.form.get('duracion_min', '60').strip()
        costo = request.form.get('costo', '0').strip()
        precio = request.form.get('precio', '0').strip()
        
        # Validaciones básicas
        if not codigo or not nombre:
            flash('Código y nombre son obligatorios.', 'danger')
            return redirect(url_for('crear_servicio'))
        
        try:
            duracion_int = int(duracion_min) if duracion_min else 60
            costo_float = float(costo) if costo else 0.0
            precio_float = float(precio) if precio else 0.0
        except ValueError:
            flash('Duración, costo y precio deben ser números válidos.', 'danger')
            return redirect(url_for('crear_servicio'))
        
        if precio_float <= 0:
            flash('El precio debe ser mayor a 0.', 'danger')
            return redirect(url_for('crear_servicio'))
        
        conn = get_db_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO servicios (codigo, nombre, categoria, descripcion, 
                                          duracion_min, costo, precio)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (codigo, nombre, categoria, descripcion or None, 
                      duracion_int, costo_float, precio_float))
                conn.commit()
                
                flash(f'Servicio {nombre} creado exitosamente.', 'success')
                return redirect(url_for('servicios'))
                
            except Error as e:
                flash(f'Error creando servicio: {e}', 'danger')
            finally:
                cursor.close()
                conn.close()
        else:
            flash('No hay conexión a la base de datos.', 'danger')
    
    return render_template('servicios/crear.html')

@app.route('/servicios/editar/<int:id>', methods=['GET', 'POST'])
def editar_servicio(id):
    """Editar servicio existente"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('servicios'))
    
    try:
        cursor = conn.cursor()
        
        if request.method == 'POST':
            # Obtener datos del formulario
            codigo = request.form.get('codigo', '').strip()
            nombre = request.form.get('nombre', '').strip()
            categoria = request.form.get('categoria', 'baño').strip()
            descripcion = request.form.get('descripcion', '').strip()
            duracion_min = request.form.get('duracion_min', '60').strip()
            costo = request.form.get('costo', '0').strip()
            precio = request.form.get('precio', '0').strip()
            
            # Validaciones
            if not codigo or not nombre:
                flash('Código y nombre son obligatorios.', 'danger')
                return redirect(url_for('editar_servicio', id=id))
            
            try:
                duracion_int = int(duracion_min) if duracion_min else 60
                costo_float = float(costo) if costo else 0.0
                precio_float = float(precio) if precio else 0.0
            except ValueError:
                flash('Duración, costo y precio deben ser números válidos.', 'danger')
                return redirect(url_for('editar_servicio', id=id))
            
            if precio_float <= 0:
                flash('El precio debe ser mayor a 0.', 'danger')
                return redirect(url_for('editar_servicio', id=id))
            
            # Actualizar servicio
            cursor.execute("""
                UPDATE servicios 
                SET codigo = %s, nombre = %s, categoria = %s, descripcion = %s,
                    duracion_min = %s, costo = %s, precio = %s,
                    margen = ((%s - %s) / NULLIF(%s, 0)) * 100
                WHERE id_servicio = %s
            """, (codigo, nombre, categoria, descripcion or None, 
                  duracion_int, costo_float, precio_float,
                  precio_float, costo_float, costo_float, id))
            conn.commit()
            
            flash(f'Servicio {nombre} actualizado exitosamente.', 'success')
            return redirect(url_for('servicios'))
        
        # GET: Obtener datos del servicio
        cursor.execute("""
            SELECT id_servicio, codigo, nombre, categoria, descripcion, 
                   costo, precio, activo, duracion_min, margen, iva
            FROM servicios 
            WHERE id_servicio = %s
        """, (id,))
        
        servicio_raw = cursor.fetchone()
        
        if not servicio_raw:
            flash('Servicio no encontrado.', 'danger')
            return redirect(url_for('servicios'))
        
        # Convertir a diccionario explícitamente
        servicio = dict(servicio_raw)
        
        # Calcular margen si no está en la BD
        if servicio.get('margen') is None and servicio.get('costo') and servicio.get('precio'):
            if servicio['costo'] > 0:
                servicio['margen'] = ((servicio['precio'] - servicio['costo']) / servicio['costo']) * 100
            else:
                servicio['margen'] = 0
        
        # Asegurar tipos de datos correctos
        servicio['costo'] = float(servicio['costo']) if servicio['costo'] is not None else 0.0
        servicio['precio'] = float(servicio['precio']) if servicio['precio'] is not None else 0.0
        servicio['margen'] = float(servicio['margen']) if servicio['margen'] is not None else 0.0
        servicio['activo'] = bool(servicio['activo']) if servicio['activo'] is not None else True
        
        # Obtener estadísticas para la vista (opcional)
        try:
            cursor.execute("""
                SELECT COUNT(*) as total_reservas,
                       SUM(rs.subtotal) as total_ingresos
                FROM reserva_servicios rs
                WHERE rs.id_servicio = %s
            """, (id,))
            estadisticas = cursor.fetchone()
            
            if estadisticas:
                servicio['total_reservas'] = estadisticas['total_reservas'] or 0
                servicio['total_ingresos'] = float(estadisticas['total_ingresos']) if estadisticas['total_ingresos'] else 0.0
        except:
            # Si falla, establecer valores por defecto
            servicio['total_reservas'] = 0
            servicio['total_ingresos'] = 0.0
        
        cursor.close()
        conn.close()
        
        return render_template('servicios/editar.html', servicio=servicio)
        
    except Error as e:
        flash(f'Error editando servicio: {e}', 'danger')
        print(f"Error en editar_servicio: {e}")  # Para debugging
        return redirect(url_for('servicios'))

@app.route('/servicios/eliminar/<int:id>', methods=['POST'])
def eliminar_servicio(id):
    """Eliminar servicio (soft delete)"""
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el servicio está en uso en reservas
        cursor.execute("""
            SELECT COUNT(*) as total 
            FROM reserva_servicios 
            WHERE id_servicio = %s
        """, (id,))
        result = cursor.fetchone()
        
        if result and result[0] > 0:
            return jsonify({
                'success': False, 
                'message': f'No se puede eliminar el servicio porque está en {result[0]} reserva(s).'
            })
        
        # Soft delete - CORRECCIÓN: Usar FALSE en lugar de false
        cursor.execute("UPDATE servicios SET activo = FALSE WHERE id_servicio = %s", (id,))  # Cambiado false por FALSE
        conn.commit()
        
        if cursor.rowcount > 0:
            return jsonify({'success': True, 'message': 'Servicio eliminado exitosamente.'})
        else:
            return jsonify({'success': False, 'message': 'Servicio no encontrado.'}), 404
            
    except Error as e:
        return jsonify({'success': False, 'message': f'Error eliminando servicio: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/servicios/ver/<int:id>')
def ver_servicio(id):
    """Ver detalles del servicio"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('servicios'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener datos del servicio
        cursor.execute("""
            SELECT id_servicio, codigo, nombre, categoria, descripcion, 
                   costo, precio, activo, duracion_min, margen, iva
            FROM servicios 
            WHERE id_servicio = %s
        """, (id,))
        
        servicio_raw = cursor.fetchone()
        
        if not servicio_raw:
            flash('Servicio no encontrado.', 'danger')
            return redirect(url_for('servicios'))
        
        # Convertir a diccionario explícitamente
        servicio = dict(servicio_raw)
        
        # Calcular margen si no está en la BD
        if servicio.get('margen') is None and servicio.get('costo') and servicio.get('precio'):
            if servicio['costo'] > 0:
                servicio['margen'] = ((servicio['precio'] - servicio['costo']) / servicio['costo']) * 100
            else:
                servicio['margen'] = 0
        
        # Obtener estadísticas de uso
        cursor.execute("""
            SELECT COUNT(*) as total_reservas,
                   SUM(rs.cantidad) as total_veces,
                   SUM(rs.subtotal) as total_ingresos
            FROM reserva_servicios rs
            WHERE rs.id_servicio = %s
        """, (id,))
        
        estadisticas = cursor.fetchone()
        
        # Asegurarnos de que existen las claves, incluso si son 0
        servicio['total_reservas'] = estadisticas['total_reservas'] if estadisticas and estadisticas['total_reservas'] else 0
        servicio['total_veces'] = estadisticas['total_veces'] if estadisticas and estadisticas['total_veces'] else 0
        servicio['total_ingresos'] = float(estadisticas['total_ingresos']) if estadisticas and estadisticas['total_ingresos'] else 0.0
        
        # Asegurar tipos de datos correctos
        servicio['costo'] = float(servicio['costo']) if servicio['costo'] is not None else 0.0
        servicio['precio'] = float(servicio['precio']) if servicio['precio'] is not None else 0.0
        servicio['margen'] = float(servicio['margen']) if servicio['margen'] is not None else 0.0
        servicio['activo'] = bool(servicio['activo']) if servicio['activo'] is not None else True
        
        cursor.close()
        conn.close()
        
        return render_template('servicios/ver.html', servicio=servicio)
        
    except Error as e:
        flash(f'Error obteniendo detalles del servicio: {e}', 'danger')
        print(f"Error en ver_servicio: {e}")  # Para debugging
        return redirect(url_for('servicios'))
# En tu app.py, agrega esta ruta al final
@app.route('/api/calendario/reservas')
def api_calendario_reservas():
    """API para obtener reservas para el calendario"""
    try:
        # Datos de ejemplo - LUEGO REEMPLAZA CON DATOS REALES
        from datetime import datetime, timedelta
        
        eventos = []
        hoy = datetime.now()
        
        # Ejemplo 1: Hoy a las 10:00
        eventos.append({
            'title': 'RES-001 - Max (Baño completo)',
            'start': hoy.replace(hour=10, minute=0, second=0).isoformat(),
            'end': hoy.replace(hour=11, minute=0, second=0).isoformat(),
            'backgroundColor': '#0dcaf0',
            'estado': 'confirmada'
        })
        
        # Ejemplo 2: Hoy a las 14:00
        eventos.append({
            'title': 'RES-002 - Luna (Corte y peinado)',
            'start': hoy.replace(hour=14, minute=0, second=0).isoformat(),
            'end': hoy.replace(hour=15, minute=30, second=0).isoformat(),
            'backgroundColor': '#ffc107',
            'estado': 'pendiente'
        })
        
        # Ejemplo 3: Mañana
        manana = hoy + timedelta(days=1)
        eventos.append({
            'title': 'RES-003 - Rocky (Limpieza dental)',
            'start': manana.replace(hour=9, minute=0, second=0).isoformat(),
            'end': manana.replace(hour=10, minute=0, second=0).isoformat(),
            'backgroundColor': '#198754',
            'estado': 'completada'
        })
        
        return jsonify({
            'success': True,
            'eventos': eventos
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/reservas')
def reservas():
    """Listar reservas"""
    # Obtener parámetros de paginación
    pagina = request.args.get('pagina', 1, type=int)
    items_por_pagina = 10
    
    conn = get_db_connection()
    reservas_list = []
    total_reservas = 0
    total_paginas = 0
    
    if conn:
        try:
            cursor = conn.cursor()
            
            # **PRIMERO: Contar el TOTAL de reservas**
            cursor.execute("SELECT COUNT(*) as total FROM reservas")
            total_reservas = cursor.fetchone()['total']
            
            # **Calcular total de páginas**
            if total_reservas > 0:
                total_paginas = (total_reservas + items_por_pagina - 1) // items_por_pagina
            
            # **Calcular OFFSET para paginación**
            offset = (pagina - 1) * items_por_pagina
            
            # **SEGUNDO: Obtener reservas con LIMIT y OFFSET**
            cursor.execute("""
                SELECT r.*, 
                       m.nombre as mascota_nombre, 
                       m.especie,
                       c.nombre as cliente_nombre, 
                       c.apellido as cliente_apellido,
                       e.nombre as empleado_nombre, 
                       e.apellido as empleado_apellido,
                       string_agg(DISTINCT s.nombre, ', ') as servicios_nombres
                FROM reservas r
                JOIN mascotas m ON r.id_mascota = m.id_mascota
                JOIN clientes c ON m.id_cliente = c.id_cliente
                JOIN empleados e ON r.id_empleado = e.id_empleado
                LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
                LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
                GROUP BY r.id_reserva, m.id_mascota, c.id_cliente, e.id_empleado
                ORDER BY r.fecha_reserva DESC 
                LIMIT %s OFFSET %s
            """, (items_por_pagina, offset))
            
            reservas_list_raw = cursor.fetchall()  # Cambio de nombre aquí
            
            # ====== ¡IMPORTANTE!: CONVERTIR A DICT MUTABLE ======
            for reserva in reservas_list_raw:
                # Convertir RealDictRow a diccionario mutable
                reserva_dict = dict(reserva)  # <-- ¡ESTO ES CLAVE!
                
                # Estado con clase CSS
                estado_clases = {
                    'pendiente': 'bg-warning',
                    'confirmada': 'bg-info',
                    'en_proceso': 'bg-primary',
                    'completada': 'bg-success',
                    'cancelada': 'bg-danger',
                    'no_show': 'bg-secondary'
                }
                reserva_dict['estado_clase'] = estado_clases.get(reserva_dict['estado'], 'bg-secondary')
                reserva_dict['estado_texto'] = reserva_dict['estado'].replace('_', ' ').title()
                
                # Verificar si está vencida
                if reserva_dict['estado'] == 'pendiente' and reserva_dict['fecha_reserva'] < datetime.now():
                    reserva_dict['vencida'] = True
                    reserva_dict['estado_clase'] = 'bg-dark'
                    reserva_dict['estado_texto'] = 'Vencida'
                else:
                    reserva_dict['vencida'] = False
                    
                # Asegurar que servicios_nombres sea string (puede ser None si no hay servicios)
                if reserva_dict['servicios_nombres'] is None:
                    reserva_dict['servicios_nombres'] = ''
                
                # Agregar a la lista final
                reservas_list.append(reserva_dict)  # <-- Agregar el dict mutable
            
            cursor.close()
        except Error as e:
            flash(f'Error obteniendo reservas: {e}', 'danger')
            print(f"Error SQL: {e}")
        finally:
            conn.close()
    else:
        # Datos demo (ya son diccionarios)
        total_reservas = 1
        total_paginas = 1
        reservas_list = [
            {
                'id_reserva': 1, 
                'codigo_reserva': 'RES-231201-0001',
                'fecha_reserva': datetime.now(), 
                'mascota_nombre': 'Max', 
                'cliente_nombre': 'Juan', 
                'cliente_apellido': 'Pérez', 
                'empleado_nombre': 'Ana', 
                'empleado_apellido': 'López',
                'servicios_nombres': 'Baño Básico, Corte',
                'estado': 'completada',
                'estado_clase': 'bg-success',
                'estado_texto': 'Completada',
                'vencida': False
            }
        ]
    
    return render_template('reservas/listar.html', 
                         reservas=reservas_list,
                         pagina_actual=pagina,
                         total_paginas=total_paginas,
                         total_reservas=total_reservas)
# ================= RUTAS DE RESERVAS =================

@app.route('/reservas/crear', methods=['GET', 'POST'])
def crear_reserva():
    """Crear nueva reserva"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reservas'))
    
    try:
        cursor = conn.cursor()
        
        if request.method == 'POST':
            # Obtener datos del formulario
            id_mascota = request.form.get('id_mascota', '').strip()
            id_empleado = request.form.get('id_empleado', '').strip()
            fecha_reserva = request.form.get('fecha_reserva', '').strip()
            hora_reserva = request.form.get('hora_reserva', '').strip()
            servicios = request.form.getlist('servicios[]')
            notas = request.form.get('notas', '').strip()
            
            # Validaciones básicas
            if not id_mascota or not id_empleado or not fecha_reserva or not hora_reserva:
                flash('Todos los campos obligatorios deben ser completados.', 'danger')
                return redirect(url_for('crear_reserva'))
            
            if not servicios:
                flash('Debe seleccionar al menos un servicio.', 'danger')
                return redirect(url_for('crear_reserva'))
            
            # Combinar fecha y hora
            try:
                fecha_hora_str = f"{fecha_reserva} {hora_reserva}"
                fecha_hora_dt = datetime.strptime(fecha_hora_str, '%Y-%m-%d %H:%M')
                ahora = datetime.now()
                
                # CORRECCIÓN: Comparar solo año, mes, día, hora y minutos
                fecha_hora_sin_segundos = fecha_hora_dt.replace(second=0, microsecond=0)
                ahora_sin_segundos = ahora.replace(second=0, microsecond=0)
                
                if fecha_hora_sin_segundos < (ahora_sin_segundos - timedelta(minutes=1)):
                    flash('No se pueden crear reservas en fechas u horas pasadas.', 'danger')
                    return redirect(url_for('crear_reserva'))
                
                # Horario de atención LUNES A DOMINGO de 9:00 AM a 6:00 PM
                hora = fecha_hora_dt.hour
                minuto = fecha_hora_dt.minute
                
                # Validar horario: 9:00 AM - 6:00 PM todos los días
                if hora < 9 or (hora == 18 and minuto > 0) or hora >= 19:
                    flash('Horario de atención: Lunes a Domingo de 9:00 AM a 6:00 PM.', 'danger')
                    return redirect(url_for('crear_reserva'))
                
                # Calcular duración total de servicios para verificar disponibilidad
                if servicios:
                    # Obtener duración de todos los servicios seleccionados
                    servicios_tuple = tuple(map(int, servicios))
                    placeholders = ','.join(['%s'] * len(servicios))
                    cursor.execute(f"""
                        SELECT SUM(duracion_min) as duracion_total
                        FROM servicios 
                        WHERE id_servicio IN ({placeholders})
                    """, servicios_tuple)
                    resultado = cursor.fetchone()
                    duracion_total = int(resultado['duracion_total']) if resultado and resultado['duracion_total'] else 60
                    
                    # OBTENER INFORMACIÓN DEL EMPLEADO PARA VERIFICAR SI PUEDE MÚLTIPLES RESERVAS
                    cursor.execute("""
                        SELECT nombre, apellido 
                        FROM empleados 
                        WHERE id_empleado = %s
                    """, (id_empleado,))
                    empleado_info = cursor.fetchone()
                    
                    # Verificar si es administrador/sistema (puede múltiples reservas)
                    es_administrador = False
                    if empleado_info:
                        nombre_completo = f"{empleado_info['nombre']} {empleado_info['apellido']}".lower()
                        if any(keyword in nombre_completo for keyword in ['admin', 'sistema', 'administrador']):
                            es_administrador = True
                    
                    # Verificar disponibilidad del empleado SOLO si NO es administrador
                    if not es_administrador:
                        cursor.execute("""
                            SELECT r.id_reserva, r.fecha_reserva, s.duracion_min
                            FROM reservas r
                            JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
                            JOIN servicios s ON rs.id_servicio = s.id_servicio
                            WHERE r.id_empleado = %s 
                            AND r.estado NOT IN ('cancelada', 'no_show')
                            AND DATE(r.fecha_reserva) = DATE(%s)
                        """, (id_empleado, fecha_reserva))
                        
                        reservas_existentes = cursor.fetchall()
                        
                        # Calcular hora de inicio y fin de la nueva reserva
                        nueva_inicio = fecha_hora_dt
                        nueva_fin = fecha_hora_dt + timedelta(minutes=int(duracion_total))
                        
                        # Verificar superposiciones
                        for reserva in reservas_existentes:
                            reserva_inicio = reserva['fecha_reserva']
                            reserva_duracion = int(reserva['duracion_min']) if reserva['duracion_min'] else 60
                            reserva_fin = reserva_inicio + timedelta(minutes=int(reserva_duracion))
                            
                            # Verificar si hay superposición
                            if (nueva_inicio < reserva_fin and nueva_fin > reserva_inicio):
                                flash(f'El empleado ya tiene una reserva de {reserva_inicio.strftime("%H:%M")} a {reserva_fin.strftime("%H:%M")}.', 'danger')
                                return redirect(url_for('crear_reserva'))
                
            except ValueError:
                flash('Formato de fecha u hora inválido.', 'danger')
                return redirect(url_for('crear_reserva'))
            
            # Generar código de reserva único
            cursor.execute("SELECT COALESCE(MAX(id_reserva), 0) + 1 as next_id FROM reservas")
            next_id = cursor.fetchone()['next_id']
            codigo_reserva = f"RES-{datetime.now().strftime('%y%m%d')}-{next_id:04d}"
            
            # ===== CORRECCIÓN PARA POSTGRESQL =====
            # Crear reserva y obtener el ID insertado usando RETURNING
            cursor.execute("""
                INSERT INTO reservas (codigo_reserva, id_mascota, id_empleado, fecha_reserva, notas)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id_reserva
            """, (codigo_reserva, id_mascota, id_empleado, fecha_hora_dt, notas or None))
            
            # Obtener el ID de la reserva creada
            result = cursor.fetchone()
            id_reserva = result['id_reserva'] if result else None
            
            if not id_reserva:
                flash('Error al crear la reserva: No se pudo obtener el ID.', 'danger')
                return redirect(url_for('crear_reserva'))
            
            print(f"DEBUG: Reserva creada con ID: {id_reserva}")
            
            # Agregar servicios a la reserva
            for servicio_id in servicios:
                cursor.execute("SELECT precio FROM servicios WHERE id_servicio = %s", (servicio_id,))
                servicio = cursor.fetchone()
                if servicio:
                    cursor.execute("""
                        INSERT INTO reserva_servicios (id_reserva, id_servicio, precio_unitario)
                        VALUES (%s, %s, %s)
                    """, (id_reserva, servicio_id, servicio['precio']))
                    print(f"DEBUG: Servicio {servicio_id} agregado a reserva {id_reserva}")
            
            conn.commit()
            
            flash(f'Reserva {codigo_reserva} creada exitosamente.', 'success')
            return redirect(url_for('ver_reserva', id=id_reserva))
        
        # GET: Obtener datos para formulario
        # Mascotas
        cursor.execute("""
            SELECT m.id_mascota, m.nombre as mascota_nombre, 
                   c.nombre as cliente_nombre, c.apellido as cliente_apellido
            FROM mascotas m
            JOIN clientes c ON m.id_cliente = c.id_cliente
            WHERE m.activo = TRUE
            ORDER BY m.nombre
        """)
        mascotas = cursor.fetchall()
        
        # Empleados - Ordenar con administrador primero
        cursor.execute("""
            SELECT id_empleado, nombre, apellido, especialidad
            FROM empleados
            WHERE activo = TRUE
            ORDER BY 
                CASE 
                    WHEN LOWER(nombre) LIKE '%admin%' OR LOWER(nombre) LIKE '%sistema%' THEN 1
                    ELSE 2
                END,
                nombre
        """)
        empleados = cursor.fetchall()
        
        # Servicios
        cursor.execute("""
            SELECT id_servicio, nombre, precio, categoria, duracion_min
            FROM servicios
            WHERE activo = TRUE
            ORDER BY categoria, nombre
        """)
        servicios = cursor.fetchall()
        
        # Agrupar servicios por categoría
        servicios_por_categoria = {}
        for servicio in servicios:
            categoria = servicio['categoria']
            if categoria not in servicios_por_categoria:
                servicios_por_categoria[categoria] = []
            servicios_por_categoria[categoria].append(servicio)
        
        # Fecha mínima (siempre hoy)
        fecha_minima = datetime.now().strftime('%Y-%m-%d')
        
        # Hora mínima dinámica
        ahora = datetime.now()
        hora_minima = f"{ahora.hour:02d}:{ahora.minute:02d}"
        
        # Ajustar si está fuera de horario
        hora_actual = ahora.hour
        minuto_actual = ahora.minute
        
        if hora_actual > 18 or (hora_actual == 18 and minuto_actual > 0):
            fecha_minima = (ahora + timedelta(days=1)).strftime('%Y-%m-%d')
            hora_minima = '09:00'
        elif hora_actual < 9:
            hora_minima = '09:00'
        else:
            if hora_actual >= 18:
                hora_minima = '17:59'
        
    except Error as e:
        flash(f'Error creando reserva: {e}', 'danger')
        return redirect(url_for('reservas'))
    finally:
        cursor.close()
        conn.close()
    
    # Buscar el ID del empleado "Admin Sistema" para marcarlo como seleccionado
    empleado_admin_id = None
    for empleado in empleados:
        if 'admin' in empleado['nombre'].lower() or 'sistema' in empleado['nombre'].lower():
            empleado_admin_id = empleado['id_empleado']
            break
    
    return render_template('reservas/crear.html', 
                         mascotas=mascotas, 
                         empleados=empleados,
                         servicios_por_categoria=servicios_por_categoria,
                         fecha_minima=fecha_minima,
                         hora_minima=hora_minima,
                         empleado_admin_id=empleado_admin_id)
@app.route('/reservas/editar/<int:id>', methods=['GET', 'POST'])
def editar_reserva(id):
    """Editar reserva existente"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reservas'))
    
    try:
        cursor = conn.cursor()
        
        # Primero obtener la reserva actual
        cursor.execute("""
            SELECT r.*, m.nombre as mascota_nombre, m.id_cliente,
                   e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                   m.especie
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN empleados e ON r.id_empleado = e.id_empleado
            WHERE r.id_reserva = %s
        """, (id,))
        reserva = cursor.fetchone()
        
        if not reserva:
            flash('Reserva no encontrada.', 'danger')
            return redirect(url_for('reservas'))
        
        # ========== ¡CONVERTIR A DICT MUTABLE! ==========
        reserva = dict(reserva)
        # ==============================================
        
        if request.method == 'POST':
            # Obtener datos del formulario
            id_mascota = request.form.get('id_mascota', '').strip()
            id_empleado = request.form.get('id_empleado', '').strip()
            fecha_reserva = request.form.get('fecha_reserva', '').strip()
            hora_reserva = request.form.get('hora_reserva', '').strip()
            servicios = request.form.getlist('servicios[]')
            notas = request.form.get('notas', '').strip()
            estado = request.form.get('estado', 'pendiente').strip()
            
            # Validaciones básicas
            if not id_mascota or not id_empleado or not fecha_reserva or not hora_reserva:
                flash('Todos los campos obligatorios deben ser completados.', 'danger')
                return redirect(url_for('editar_reserva', id=id))
            
            if not servicios:
                flash('Debe seleccionar al menos un servicio.', 'danger')
                return redirect(url_for('editar_reserva', id=id))
            
            # Combinar fecha y hora
            try:
                fecha_hora_str = f"{fecha_reserva} {hora_reserva}"
                fecha_hora_dt = datetime.strptime(fecha_hora_str, '%Y-%m-%d %H:%M')
                ahora = datetime.now()
                
                # Para edición: solo validar si la nueva fecha/hora es diferente a la actual
                # y si está en el futuro
                fecha_hora_actual = reserva['fecha_reserva']
                
                # Si se cambió la fecha/hora, aplicar validaciones
                if fecha_hora_dt != fecha_hora_actual:
                    # Verificar que no sea fecha/hora pasada
                    if fecha_hora_dt < ahora:
                        flash('No se pueden cambiar a fechas/horas pasadas.', 'danger')
                        return redirect(url_for('editar_reserva', id=id))
                    
                    # Para reservas que ya estaban confirmadas/comenzadas, permitir cambios con menos restricciones
                    if reserva['estado'] in ['pendiente', 'confirmada']:
                        # Solo validar horario de atención
                        dia_semana = fecha_hora_dt.weekday()  # 0 = lunes, 6 = domingo
                        hora = fecha_hora_dt.hour
                        minuto = fecha_hora_dt.minute
                        
                        # Validar domingo
                        if dia_semana == 6:
                            flash('Domingo cerrado. No se pueden programar reservas.', 'danger')
                            return redirect(url_for('editar_reserva', id=id))
                        
                        # Validar sábado (9:00 - 14:00)
                        if dia_semana == 5:
                            if hora < 9 or (hora == 14 and minuto > 0) or hora >= 15:
                                flash('Sábados: horario de atención 9:00 - 14:00.', 'danger')
                                return redirect(url_for('editar_reserva', id=id))
                        
                        # Validar lunes a viernes (9:00 - 18:00)
                        else:
                            if hora < 9 or (hora == 18 and minuto > 0) or hora >= 19:
                                flash('Lunes a Viernes: horario de atención 9:00 - 18:00.', 'danger')
                                return redirect(url_for('editar_reserva', id=id))
                    
                    # Para reservas en estado diferente, validar más estrictamente
                    else:
                        margen_minutos = 30
                        if fecha_hora_dt < (ahora + timedelta(minutes=margen_minutos)):
                            flash(f'Las reservas deben hacerse con al menos {margen_minutos} minutos de anticipación.', 'danger')
                            return redirect(url_for('editar_reserva', id=id))
                        
                        # Validar horario de atención
                        dia_semana = fecha_hora_dt.weekday()
                        hora = fecha_hora_dt.hour
                        minuto = fecha_hora_dt.minute
                        
                        # Validar domingo
                        if dia_semana == 6:
                            flash('Domingo cerrado.', 'danger')
                            return redirect(url_for('editar_reserva', id=id))
                        
                        # Validar sábado (9:00 - 14:00)
                        if dia_semana == 5:
                            if hora < 9 or (hora == 14 and minuto > 0) or hora >= 15:
                                flash('Sábados: horario de atención 9:00 - 14:00.', 'danger')
                                return redirect(url_for('editar_reserva', id=id))
                        
                        # Validar lunes a viernes (9:00 - 18:00)
                        else:
                            if hora < 9 or (hora == 18 and minuto > 0) or hora >= 19:
                                flash('Lunes a Viernes: horario de atención 9:00 - 18:00.', 'danger')
                                return redirect(url_for('editar_reserva', id=id))
                
                # Verificar disponibilidad del empleado si se cambió empleado o fecha/hora
                if id_empleado != str(reserva['id_empleado']) or fecha_hora_dt != fecha_hora_actual:
                    # OBTENER INFORMACIÓN DEL EMPLEADO
                    cursor.execute("SELECT nombre, apellido FROM empleados WHERE id_empleado = %s", (id_empleado,))
                    empleado_info = cursor.fetchone()
    
                    es_administrador = False
                    if empleado_info:
                        nombre_completo = f"{empleado_info['nombre']} {empleado_info['apellido']}".lower()
                        if any(keyword in nombre_completo for keyword in ['admin', 'sistema', 'administrador']):
                            es_administrador = True    
                    # Calcular duración total de servicios para verificar disponibilidad
                    if servicios:
                        servicios_tuple = tuple(map(int, servicios))
                        placeholders = ','.join(['%s'] * len(servicios))
                        cursor.execute(f"""
                            SELECT SUM(duracion_min) as duracion_total
                            FROM servicios 
                            WHERE id_servicio IN ({placeholders})
                        """, servicios_tuple)
                        resultado = cursor.fetchone()
                        duracion_total = int(resultado['duracion_total']) if resultado and resultado['duracion_total'] else 60
                        
                        # Verificar disponibilidad del empleado (excluyendo esta reserva)
                        cursor.execute("""
                            SELECT r.id_reserva, r.fecha_reserva, s.duracion_min
                            FROM reservas r
                            JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
                            JOIN servicios s ON rs.id_servicio = s.id_servicio
                            WHERE r.id_empleado = %s 
                            AND r.id_reserva != %s
                            AND r.estado NOT IN ('cancelada', 'no_show')
                            AND DATE(r.fecha_reserva) = DATE(%s)
                        """, (id_empleado, id, fecha_reserva))
                        
                        reservas_existentes = cursor.fetchall()
                        
                        # Calcular hora de inicio y fin de la nueva reserva
                        nueva_inicio = fecha_hora_dt
                        nueva_fin = fecha_hora_dt + timedelta(minutes=int(duracion_total))
                        
                        # Verificar superposiciones
                        for reserva_existente in reservas_existentes:
                            reserva_inicio = reserva_existente['fecha_reserva']
                            reserva_duracion = int(reserva_existente['duracion_min']) if reserva_existente['duracion_min'] else 60
                            reserva_fin = reserva_inicio + timedelta(minutes=int(reserva_duracion))
                            
                            # Verificar si hay superposición
                            if (nueva_inicio < reserva_fin and nueva_fin > reserva_inicio):
                                flash(f'El empleado ya tiene una reserva de {reserva_inicio.strftime("%H:%M")} a {reserva_fin.strftime("%H:%M")}.', 'danger')
                                return redirect(url_for('editar_reserva', id=id))
                
            except ValueError:
                flash('Formato de fecha u hora inválido.', 'danger')
                return redirect(url_for('editar_reserva', id=id))
            
            # Actualizar reserva
            cursor.execute("""
                UPDATE reservas 
                SET id_mascota = %s, id_empleado = %s, fecha_reserva = %s, 
                    notas = %s, estado = %s
                WHERE id_reserva = %s
            """, (id_mascota, id_empleado, fecha_hora_dt, notas or None, estado, id))
            
            # Eliminar servicios anteriores y agregar nuevos
            cursor.execute("DELETE FROM reserva_servicios WHERE id_reserva = %s", (id,))
            
            # Agregar servicios a la reserva
            for servicio_id in servicios:
                # Obtener precio del servicio
                cursor.execute("SELECT precio FROM servicios WHERE id_servicio = %s", (servicio_id,))
                servicio = cursor.fetchone()
                if servicio:
                    cursor.execute("""
                        INSERT INTO reserva_servicios (id_reserva, id_servicio, precio_unitario)
                        VALUES (%s, %s, %s)
                    """, (id, servicio_id, servicio['precio']))
            
            conn.commit()
            
            flash(f'Reserva actualizada exitosamente.', 'success')
            return redirect(url_for('ver_reserva', id=id))
        
        # GET: Obtener datos de la reserva
        # Obtener servicios de la reserva
        cursor.execute("""
            SELECT rs.id_servicio, s.nombre, rs.precio_unitario
            FROM reserva_servicios rs
            JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE rs.id_reserva = %s
        """, (id,))
        servicios_reserva = cursor.fetchall()
        servicios_ids = [str(s['id_servicio']) for s in servicios_reserva]
        
        # Formatear fecha para inputs
        fecha_hora = reserva['fecha_reserva']
        reserva['fecha_str'] = fecha_hora.strftime('%Y-%m-%d')
        reserva['hora_str'] = fecha_hora.strftime('%H:%M')
        
        # Calcular si la reserva está vencida
        ahora = datetime.now()
        reserva['vencida'] = fecha_hora < ahora and reserva['estado'] in ['pendiente', 'confirmada']
        
        # Datos para formulario
        # Mascotas
        cursor.execute("""
            SELECT m.id_mascota, m.nombre as mascota_nombre, 
                   c.nombre as cliente_nombre, c.apellido as cliente_apellido
            FROM mascotas m
            JOIN clientes c ON m.id_cliente = c.id_cliente
            WHERE m.activo = TRUE
            ORDER BY m.nombre
        """)
        mascotas = cursor.fetchall()
        
        # Empleados
        cursor.execute("""
            SELECT id_empleado, nombre, apellido, especialidad
            FROM empleados
            WHERE activo = TRUE
            ORDER BY nombre
        """)
        empleados = cursor.fetchall()
        
        # Servicios
        cursor.execute("""
            SELECT id_servicio, nombre, precio, categoria, duracion_min
            FROM servicios
            WHERE activo = TRUE
            ORDER BY categoria, nombre
        """)
        servicios = cursor.fetchall()
        
        # Agrupar servicios por categoría
        servicios_por_categoria = {}
        for servicio in servicios:
            categoria = servicio['categoria']
            if categoria not in servicios_por_categoria:
                servicios_por_categoria[categoria] = []
            servicios_por_categoria[categoria].append(servicio)
        
    except Error as e:
        flash(f'Error editando reserva: {e}', 'danger')
        return redirect(url_for('reservas'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reservas/editar.html', 
                         reserva=reserva,
                         servicios_ids=servicios_ids,
                         mascotas=mascotas, 
                         empleados=empleados,
                         servicios_por_categoria=servicios_por_categoria)

@app.route('/reservas/eliminar/<int:id>', methods=['POST'])
def eliminar_reserva(id):
    """Eliminar reserva"""
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si la reserva tiene factura asociada
        cursor.execute("SELECT COUNT(*) as total FROM facturas WHERE id_reserva = %s", (id,))
        result = cursor.fetchone()
        
        if result and result[0] > 0:
            return jsonify({
                'success': False, 
                'message': f'No se puede eliminar la reserva porque tiene una factura asociada.'
            })
        
        # Eliminar servicios de la reserva primero
        cursor.execute("DELETE FROM reserva_servicios WHERE id_reserva = %s", (id,))
        
        # Eliminar reserva
        cursor.execute("DELETE FROM reservas WHERE id_reserva = %s", (id,))
        conn.commit()
        
        if cursor.rowcount > 0:
            return jsonify({'success': True, 'message': 'Reserva eliminada exitosamente.'})
        else:
            return jsonify({'success': False, 'message': 'Reserva no encontrada.'}), 404
            
    except Error as e:
        return jsonify({'success': False, 'message': f'Error eliminando reserva: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

def obtener_correo_admin():
    """Obtiene el correo del administrador del sistema"""
    try:
        # Usar el correo de configuración de Flask
        correo = current_app.config.get('ADMIN_EMAIL')
        if not correo:
            correo = 'ayumu798@gmail.com'  # Correo por defecto
        print(f"✅ Usando correo admin: {correo}")
        return correo
    except Exception as e:
        print(f"⚠️ Error al obtener correo admin: {e}")
        return 'ayumu798@gmail.com'

def enviar_correo_reserva_completada(reserva_dict):
    """Envía correo usando Resend (funciona en Render)"""
    try:
        print(f"\n🔍 DEBUG: Datos recibidos en reserva_dict:")
        for key, value in reserva_dict.items():
            print(f"   {key}: {value}")
        
        # 1. Configurar API key
        api_key = app.config.get('RESEND_API_KEY')
        if not api_key:
            print("❌ RESEND_API_KEY no configurada")
            return False
        
        resend.api_key = api_key
        
        # 2. Preparar datos
        codigo_reserva = reserva_dict.get('codigo_reserva', 'Sin código')
        mascota = reserva_dict.get('mascota_nombre', 'N/A')
        cliente = f"{reserva_dict.get('cliente_nombre', '')} {reserva_dict.get('cliente_apellido', '')}".strip()
        
        # Obtener total con manejo robusto
        total_raw = reserva_dict.get('total', 0)
        print(f"🔍 total_raw: {total_raw} (tipo: {type(total_raw)})")
        
        try:
            if total_raw is None:
                total = 0.0
            elif isinstance(total_raw, (int, float)):
                total = float(total_raw)
            else:
                total = float(str(total_raw))
        except (ValueError, TypeError) as e:
            print(f"⚠️ Error convirtiendo total: {e}")
            total = 0.0
        
        fecha = datetime.now().strftime('%d/%m/%Y %H:%M')
        servicios_texto = reserva_dict.get('servicios_texto', 'No especificados')
        
        print(f"📧 Preparando correo para: {codigo_reserva}")
        print(f"   Total: S/ {total:.2f}")
        print(f"   Servicios: {servicios_texto}")
        
        # 3. Crear contenido HTML
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; color: #333; line-height: 1.6; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ background: #4CAF50; color: white; padding: 20px; text-align: center; border-radius: 10px 10px 0 0; }}
                .content {{ background: #f9f9f9; padding: 30px; border-radius: 0 0 10px 10px; }}
                .badge {{ background: #4CAF50; color: white; padding: 5px 15px; border-radius: 20px; display: inline-block; }}
                .total {{ font-size: 24px; color: #4CAF50; font-weight: bold; }}
                .footer {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 12px; }}
                .details {{ background: white; padding: 15px; border-radius: 5px; margin: 15px 0; }}
                .servicios {{ background: #e8f5e9; padding: 10px; border-radius: 5px; margin: 10px 0; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>✅ Reserva Completada</h1>
                    <span class="badge">{codigo_reserva}</span>
                </div>
                
                <div class="content">
                    <p>¡Excelente! Una reserva ha sido completada exitosamente.</p>
                    
                    <div class="details">
                        <h3>📋 Detalles de la Reserva</h3>
                        <p><strong>🦮 Mascota:</strong> {mascota}</p>
                        <p><strong>👤 Cliente:</strong> {cliente}</p>
                        <p><strong>📅 Fecha:</strong> {fecha}</p>
                        
                        <div class="servicios">
                            <strong>🛠️ Servicios:</strong><br>
                            {servicios_texto}
                        </div>
                        
                        <p><strong>💰 Total:</strong> <span class="total">S/ {total:.2f}</span></p>
                    </div>
                    
                    <p>Puedes revisar los detalles completos en el sistema PetGlow.</p>
                    
                    <div class="footer">
                        <p>Este es un correo automático del sistema PetGlow.</p>
                        <p>Fecha de envío: {fecha}</p>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        # 4. Contenido texto plano
        text_content = f"""
        RESERVA COMPLETADA - PETGLOW
        =============================
        
        ✅ Reserva completada exitosamente
        
        Código: {codigo_reserva}
        Mascota: {mascota}
        Cliente: {cliente}
        Servicios: {servicios_texto}
        Total: S/ {total:.2f}
        Fecha: {fecha}
        
        Puedes revisar los detalles en el sistema PetGlow.
        
        Este es un correo automático del sistema PetGlow.
        Fecha de envío: {fecha}
        """
        
        # 5. Enviar correo
        response = resend.Emails.send({
            "from": "PetGlow <onboarding@resend.dev>",
            "to": ["ayumu798@gmail.com"],
            "subject": f"✅ Reserva Completada - {codigo_reserva}",
            "html": html_content,
            "text": text_content
        })
        
        print(f"✅ Correo enviado exitosamente!")
        print(f"   ID del correo: {response.get('id')}")
        print(f"   Total mostrado: S/ {total:.2f}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error con Resend: {str(e)}")
        import traceback
        traceback.print_exc()
        return False
# ===================== RUTAS API PARA EMPLEADOS ====================

@app.route('/api/empleado/info')
@login_required
def api_empleado_info():
    """Obtener información del empleado actual"""
    id_empleado = session.get('id_empleado')
    
    if not id_empleado:
        return jsonify({'success': False, 'message': 'No se encontró empleado'})
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Error de conexión'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Obtener información del empleado
        cursor.execute("""
            SELECT e.id_empleado, e.nombre, e.apellido, e.especialidad, e.email, e.telefono,
                   e.fecha_contratacion, u.username, u.rol
            FROM empleados e
            LEFT JOIN usuarios u ON e.id_empleado = u.id_empleado
            WHERE e.id_empleado = %s
        """, (id_empleado,))
        
        empleado = cursor.fetchone()
        
        if empleado:
            return jsonify({
                'success': True,
                'id': empleado['id_empleado'],
                'nombre': f"{empleado['nombre']} {empleado['apellido']}",
                'nombre_corto': empleado['nombre'],
                'apellido': empleado['apellido'],
                'especialidad': empleado['especialidad'],
                'email': empleado['email'],
                'telefono': empleado['telefono'],
                'username': empleado.get('username'),
                'rol': empleado.get('rol'),
                'fecha_contratacion': empleado['fecha_contratacion'].strftime('%d/%m/%Y') if empleado['fecha_contratacion'] else None
            })
        else:
            return jsonify({'success': False, 'message': 'Empleado no encontrado'})
            
    except Error as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cursor.close()
        conn.close()
# ==================== VISTA DE EMPLEADOS ====================

@app.route('/empleado/reservas')
@login_required
def empleado_reservas():
    """Vista especial para empleados - Ver sus reservas asignadas"""
    id_empleado = session.get('id_empleado')
    
    if not id_empleado:
        flash('No se encontró información del empleado.', 'danger')
        return redirect(url_for('dashboard'))
    
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener reservas asignadas a este empleado
        # Hoy + próximos 7 días
        cursor.execute("""
            SELECT 
                r.*,
                m.id_mascota,
                m.nombre as mascota_nombre,
                m.especie,
                m.raza,
                m.color,
                m.tamano,
                c.id_cliente,
                c.nombre as cliente_nombre,
                c.apellido as cliente_apellido,
                c.telefono as cliente_telefono,
                CONCAT(e.nombre, ' ', e.apellido) as empleado_nombre,
                string_agg(DISTINCT s.nombre SEPARATOR ', ') as servicios_nombres,
                SUM(s.duracion_min) as duracion_total
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            JOIN empleados e ON r.id_empleado = e.id_empleado
            LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
            LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE r.id_empleado = %s
            AND DATE(r.fecha_reserva) BETWEEN CURRENT_DATE AND DATE_ADD(CURRENT_DATE, INTERVAL 7 DAY)
            AND r.estado IN ('pendiente', 'confirmada', 'en_proceso')
            GROUP BY r.id_reserva
            ORDER BY 
                CASE r.estado 
                    WHEN 'en_proceso' THEN 1
                    WHEN 'confirmada' THEN 2
                    WHEN 'pendiente' THEN 3
                    ELSE 4
                END,
                r.fecha_reserva ASC
        """, (id_empleado,))
        
        reservas_asignadas = cursor.fetchall()
        
        # Obtener estadísticas
        cursor.execute("""
            SELECT 
                COUNT(*) as total,
                SUM(CASE WHEN estado = 'pendiente' THEN 1 ELSE 0 END) as pendientes,
                SUM(CASE WHEN estado = 'confirmada' THEN 1 ELSE 0 END) as confirmadas,
                SUM(CASE WHEN estado = 'en_proceso' THEN 1 ELSE 0 END) as en_proceso
            FROM reservas 
            WHERE id_empleado = %s 
            AND DATE(fecha_reserva) >= CURRENT_DATE
            AND estado IN ('pendiente', 'confirmada', 'en_proceso')
        """, (id_empleado,))
        
        estadisticas = cursor.fetchone()
        
        # Obtener información del empleado
        cursor.execute("""
            SELECT nombre, apellido, especialidad 
            FROM empleados 
            WHERE id_empleado = %s
        """, (id_empleado,))
        
        empleado_info = cursor.fetchone()
        
    except Error as e:
        flash(f'Error obteniendo reservas: {e}', 'danger')
        return redirect(url_for('dashboard'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('empleados/reservas.html', 
                         reservas=reservas_asignadas,
                         estadisticas=estadisticas,
                         empleado=empleado_info)

@app.route('/api/empleado/reservas/<int:id>/estado', methods=['POST'])
@login_required
def api_cambiar_estado_reserva_empleado(id):
    """API para que empleados cambien estado de sus reservas"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'}), 401
    
    data = request.get_json()
    nuevo_estado = data.get('estado')
    
    if not nuevo_estado:
        return jsonify({'success': False, 'message': 'Estado no especificado.'}), 400
    
    estados_validos = ['confirmada', 'en_proceso', 'completada', 'cancelada']
    if nuevo_estado not in estados_validos:
        return jsonify({'success': False, 'message': 'Estado no válido.'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        id_empleado = session.get('id_empleado')
        
        # Verificar que la reserva pertenece a este empleado
        cursor.execute("""
            SELECT id_reserva FROM reservas 
            WHERE id_reserva = %s AND id_empleado = %s
        """, (id, id_empleado))
        
        reserva = cursor.fetchone()
        
        if not reserva:
            return jsonify({'success': False, 'message': 'Reserva no encontrada o no asignada.'}), 404
        
        # Cambiar estado
        cursor.execute("""
            UPDATE reservas 
            SET estado = %s, fecha_modificacion = NOW()
            WHERE id_reserva = %s
        """, (nuevo_estado, id))
        
        conn.commit()
        
        return jsonify({
            'success': True, 
            'message': f'Estado cambiado a {nuevo_estado.replace("_", " ").title()}',
            'estado': nuevo_estado
        })
            
    except Error as e:
        return jsonify({'success': False, 'message': f'Error cambiando estado: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleado/reservas/hoy')
@login_required
def api_reservas_hoy_empleado():
    """API para obtener reservas de hoy del empleado (para pantallas de trabajo)"""
    id_empleado = session.get('id_empleado')
    
    if not id_empleado:
        return jsonify({'success': False, 'message': 'No se encontró empleado'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Reservas de hoy para este empleado
        cursor.execute("""
            SELECT 
                r.id_reserva,
                r.codigo_reserva,
                r.fecha_reserva,
                r.estado,
                r.notas,
                m.nombre as mascota_nombre,
                m.especie,
                m.raza,
                m.color,
                c.nombre as cliente_nombre,
                c.apellido as cliente_apellido,
                c.telefono as cliente_telefono,
                string_agg(DISTINCT s.nombre SEPARATOR ', ') as servicios_nombres,
                SUM(s.duracion_min) as duracion_total
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
            LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE r.id_empleado = %s
            AND DATE(r.fecha_reserva) = CURRENT_DATE
            AND r.estado IN ('pendiente', 'confirmada', 'en_proceso')
            GROUP BY r.id_reserva
            ORDER BY r.fecha_reserva ASC
        """, (id_empleado,))
        
        reservas_hoy = cursor.fetchall()
        
        # Formatear datos
        for reserva in reservas_hoy:
            # Estado con clase CSS
            estado_clases = {
                'pendiente': 'warning',
                'confirmada': 'info',
                'en_proceso': 'primary',
                'completada': 'success',
                'cancelada': 'danger'
            }
            reserva['estado_color'] = estado_clases.get(reserva['estado'], 'secondary')
            
            # Formatear fecha/hora
            if reserva['fecha_reserva']:
                fecha_obj = reserva['fecha_reserva']
                if isinstance(fecha_obj, datetime):
                    reserva['hora_str'] = fecha_obj.strftime('%H:%M')
                    reserva['fecha_str'] = fecha_obj.strftime('%d/%m/%Y')
                    
                    # Calcular minutos restantes
                    ahora = datetime.now()
                    diferencia = (fecha_obj - ahora).total_seconds() / 60
                    reserva['minutos_restantes'] = int(diferencia) if diferencia > 0 else 0
                    reserva['es_proxima'] = 0 < diferencia <= 30
        
        return jsonify({
            'success': True,
            'reservas': reservas_hoy,
            'total': len(reservas_hoy),
            'hoy': datetime.now().strftime('%d/%m/%Y')
        })
        
    except Error as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/empleado/monitor')
@login_required
def empleado_monitor():
    """Pantalla de monitor/kiosk para empleados (solo lectura, auto-refresh)"""
    return render_template('empleados/monitor.html')

@app.route('/api/monitor/reservas/<int:id>/tomar', methods=['POST'])
@login_required
def api_tomar_reserva(id):
    """API para que un empleado tome una reserva del monitor"""
    
    data = request.get_json()
    id_empleado_seleccionado = data.get('id_empleado')
    
    # Si no se envía empleado, usar el de la sesión
    if not id_empleado_seleccionado:
        id_empleado_seleccionado = session.get('id_empleado')
    
    if not id_empleado_seleccionado:
        return jsonify({'success': False, 'message': 'No se especificó empleado.'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Error de conexión.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 1. Verificar que la reserva existe
        cursor.execute("""
            SELECT estado, id_empleado, codigo_reserva
            FROM reservas 
            WHERE id_reserva = %s
        """, (id,))
        
        reserva = cursor.fetchone()
        
        if not reserva:
            return jsonify({'success': False, 'message': 'Reserva no encontrada.'}), 404
        
        # 2. Verificar que está en estado correcto
        if reserva['estado'] not in ['pendiente', 'confirmada']:
            return jsonify({'success': False, 'message': 'Solo se pueden tomar reservas pendientes o confirmadas.'}), 400
        
        # 3. Verificar que no está ya asignada a otro empleado (que no sea Admin Sistema)
        if reserva['id_empleado'] != 1 and reserva['id_empleado'] != id_empleado_seleccionado:
            return jsonify({'success': False, 'message': 'Esta reserva ya está asignada a otro empleado.'}), 400
        
        # 4. Obtener nombre del empleado seleccionado
        cursor.execute("SELECT nombre, apellido FROM empleados WHERE id_empleado = %s", (id_empleado_seleccionado,))
        empleado_info = cursor.fetchone()
        
        if not empleado_info:
            return jsonify({'success': False, 'message': 'Empleado no encontrado.'}), 404
        
        nombre_empleado = f"{empleado_info['nombre']} {empleado_info['apellido']}"
        
        # 5. Cambiar el empleado asignado y el estado
        cursor.execute("""
            UPDATE reservas 
            SET id_empleado = %s,
                estado = 'en_proceso',
                fecha_modificacion = NOW()
            WHERE id_reserva = %s
        """, (id_empleado_seleccionado, id))
        
        conn.commit()
        
        return jsonify({
            'success': True, 
            'message': f'Reserva asignada a {nombre_empleado} y en proceso.',
            'estado': 'en_proceso',
            'nuevo_empleado': nombre_empleado,
            'id_empleado_nuevo': id_empleado_seleccionado
        })
            
    except Error as e:
        conn.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()
@app.route('/api/monitor/reservas')
@login_required
def api_monitor_reservas():
    """API para obtener reservas para el monitor de empleados - TODAS LAS RESERVAS"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Obtener ID del empleado actual DE LA SESIÓN
        id_empleado = session.get('id_empleado')
        nombre_empleado_actual = "Empleado"
        
        # Si no hay id_empleado en sesión, buscar por el usuario actual
        if not id_empleado and 'id_usuario' in session:
            cursor.execute("""
                SELECT e.id_empleado, e.nombre, e.apellido 
                FROM usuarios u
                JOIN empleados e ON u.id_empleado = e.id_empleado
                WHERE u.id_usuario = %s
            """, (session['id_usuario'],))
            usuario_info = cursor.fetchone()
            
            if usuario_info:
                id_empleado = usuario_info['id_empleado']
                nombre_empleado_actual = f"{usuario_info['nombre']} {usuario_info['apellido']}"
                # Actualizar la sesión con el id_empleado encontrado
                session['id_empleado'] = id_empleado
        
        # Si aún no hay id_empleado, usar el nombre de la sesión
        if not id_empleado:
            nombre_empleado_actual = session.get('nombre', 'Empleado')
        
        # Obtener TODAS las reservas de hoy
        cursor.execute("""
            SELECT 
                r.id_reserva,
                r.codigo_reserva,
                DATE_FORMAT(r.fecha_reserva, '%Y-%m-%dT%H:%i:%s') as fecha_reserva,
                r.estado,
                r.id_empleado,
                m.nombre as mascota_nombre,
                m.especie,
                m.raza,
                m.color,
                c.nombre as cliente_nombre,
                c.apellido as cliente_apellido,
                c.telefono as cliente_telefono,
                string_agg(DISTINCT s.nombre SEPARATOR ', ') as servicios_nombres,
                CONCAT(e.nombre, ' ', e.apellido) as empleado_asignado,
                e.nombre as empleado_nombre,
                e.apellido as empleado_apellido
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            LEFT JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
            LEFT JOIN servicios s ON rs.id_servicio = s.id_servicio
            LEFT JOIN empleados e ON r.id_empleado = e.id_empleado
            WHERE DATE(r.fecha_reserva) = CURRENT_DATE
            AND r.estado IN ('pendiente', 'confirmada', 'en_proceso', 'completada')
            GROUP BY r.id_reserva
            ORDER BY 
                CASE r.estado 
                    WHEN 'pendiente' THEN 1
                    WHEN 'confirmada' THEN 2
                    WHEN 'en_proceso' THEN 3
                    WHEN 'completada' THEN 4
                    ELSE 5
                END,
                r.fecha_reserva ASC
        """)
        
        todas_reservas = cursor.fetchall()
        
        # Separar por estado
        pendientes = [r for r in todas_reservas if r['estado'] in ['pendiente', 'confirmada']]
        en_proceso = [r for r in todas_reservas if r['estado'] == 'en_proceso']
        completadas = [r for r in todas_reservas if r['estado'] == 'completada']
        
        return jsonify({
            'success': True,
            'pendientes': pendientes,
            'en_proceso': en_proceso,
            'completadas': completadas,
            'id_empleado_actual': id_empleado,
            'nombre_empleado_actual': nombre_empleado_actual
        })
        
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/monitor/empleados')
def api_monitor_empleados():
    """API para obtener lista de empleados para el monitor"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'Error de conexión'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Obtener todos los empleados activos (EXCLUYENDO al Admin Sistema - ID 1)
        cursor.execute("""
            SELECT id_empleado, dni, nombre, apellido, email, 
                   especialidad, telefono, fecha_contratacion
            FROM empleados
            WHERE activo = TRUE AND id_empleado != 1  -- Excluir Admin Sistema
            ORDER BY nombre, apellido
        """)
        
        empleados = cursor.fetchall()
        
        return jsonify({
            'success': True,
            'empleados': empleados,
            'total': len(empleados)
        })
        
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/reservas/devolver/<int:id>', methods=['POST'])
@login_required
def devolver_reserva(id):
    """Devolver una reserva al Admin Sistema para que otros puedan tomarla"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'message': 'Error de conexión.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # 1. Verificar que la reserva existe y está en_proceso
        cursor.execute("""
            SELECT estado, id_empleado, codigo_reserva
            FROM reservas 
            WHERE id_reserva = %s
        """, (id,))
        
        reserva = cursor.fetchone()
        
        if not reserva:
            return jsonify({'success': False, 'message': 'Reserva no encontrada.'}), 404
        
        # 2. Verificar que está en proceso
        if reserva['estado'] != 'en_proceso':
            return jsonify({'success': False, 'message': 'Solo se pueden devolver reservas en proceso.'}), 400
        
        # 3. Verificar que el empleado actual es el asignado
        id_empleado_sesion = session.get('id_empleado')
        if not id_empleado_sesion:
            return jsonify({'success': False, 'message': 'No hay empleado en sesión.'}), 400
        
        if reserva['id_empleado'] != id_empleado_sesion:
            return jsonify({'success': False, 'message': 'Solo puedes devolver reservas asignadas a ti.'}), 403
        
        # 4. Devolver al Admin Sistema (ID 1) y cambiar estado a "confirmada"
        cursor.execute("""
            UPDATE reservas 
            SET id_empleado = 1,
                estado = 'confirmada',
                fecha_modificacion = NOW()
            WHERE id_reserva = %s
        """, (id,))
        
        conn.commit()
        
        return jsonify({
            'success': True, 
            'message': f'Reserva #{reserva["codigo_reserva"]} devuelta a disponibles.',
            'nuevo_estado': 'confirmada',
            'nuevo_empleado': 'Admin Sistema'
        })
            
    except Error as e:
        conn.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/mascota/<int:id>')
def obtener_datos_mascota(id):
    """Obtener datos completos de una mascota para AJAX - CORREGIDO PARA POSTGRESQL"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'No hay conexión a la base de datos'})
    
    try:
        cursor = conn.cursor()
        
        # Obtener datos completos de la mascota - CORREGIDO PARA POSTGRESQL
        cursor.execute("""
            SELECT 
                m.*, 
                c.nombre as cliente_nombre, 
                c.apellido as cliente_apellido,
                c.telefono as cliente_telefono,
                c.email as cliente_email
            FROM mascotas m
            JOIN clientes c ON m.id_cliente = c.id_cliente
            WHERE m.id_mascota = %s
        """, (id,))
        
        mascota_raw = cursor.fetchone()
        
        if not mascota_raw:
            return jsonify({'success': False, 'error': 'Mascota no encontrada'})
        
        # Convertir a diccionario
        mascota = dict(mascota_raw)
        
        # CALCULAR EDAD EN PYTHON (en lugar de TIMESTAMPDIFF de MySQL)
        if mascota.get('fecha_nacimiento'):
            from datetime import datetime
            hoy = datetime.now()
            nacimiento = mascota['fecha_nacimiento']
            
            años = hoy.year - nacimiento.year
            # Ajustar si aún no ha pasado el cumpleaños este año
            if (hoy.month, hoy.day) < (nacimiento.month, nacimiento.day):
                años -= 1
            
            # Calcular meses restantes
            meses = hoy.month - nacimiento.month
            if hoy.day < nacimiento.day:
                meses -= 1
            if meses < 0:
                meses += 12
            
            mascota['edad_anios'] = años
            mascota['edad_meses'] = meses
        else:
            mascota['edad_anios'] = None
            mascota['edad_meses'] = None
        
        # OBTENER HISTORIAL DE CORTES - ¡CORREGIDO PARA POSTGRESQL!
        cursor.execute("""
            SELECT 
                hc.*,
                e.nombre || ' ' || e.apellido as empleado_nombre,
                hc.fecha_registro
            FROM historial_cortes hc
            LEFT JOIN empleados e ON hc.id_empleado = e.id_empleado
            WHERE hc.id_mascota = %s
            ORDER BY hc.fecha_registro DESC
            LIMIT 10
        """, (id,))
        
        historial_raw = cursor.fetchall()
        
        # Convertir historial a lista de diccionarios y formatear
        historial_cortes = []
        for corte in historial_raw:
            corte_dict = dict(corte)
            
            # Formatear fecha
            if corte_dict.get('fecha_registro'):
                corte_dict['fecha_formateada'] = corte_dict['fecha_registro'].strftime('%d/%m/%Y %H:%M')
            else:
                corte_dict['fecha_formateada'] = 'Fecha no disponible'
            
            historial_cortes.append(corte_dict)
        
        cursor.close()
        conn.close()
        
        return jsonify({
            'success': True,
            'mascota': mascota,
            'historial_cortes': historial_cortes
        })
        
    except Exception as e:
        print(f"Error al obtener datos de mascota: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})
        
@app.route('/mascotas/actualizar-datos/<int:id>', methods=['POST'])
def actualizar_datos_mascota_reserva(id):
    """Actualizar datos de mascota desde la reserva - CORREGIDO PARA POSTGRESQL"""
    print(f"Recibiendo actualización para mascota ID: {id}")
    print(f"Datos recibidos: {request.form}")
    
    if request.method == 'POST':
        conn = get_db_connection()
        if not conn:
            print("ERROR: No hay conexión a la base de datos")
            return jsonify({'success': False, 'error': 'No hay conexión a la base de datos'})
        
        try:
            # Obtener datos del formulario con valores por defecto
            raza = request.form.get('raza', '').strip()
            color = request.form.get('color', '').strip()
            corte = request.form.get('corte', '').strip()
            tamano = request.form.get('tamano', '').strip().lower()
            peso = request.form.get('peso', '').strip()
            caracteristicas = request.form.get('caracteristicas', '').strip()
            alergias = request.form.get('alergias', '').strip()
            
            print(f"Procesando datos: raza={raza}, color={color}, corte={corte}, tamano={tamano}, peso={peso}")
            
            # Validar campos requeridos
            if not raza:
                return jsonify({'success': False, 'error': 'El campo Raza es obligatorio'})
            if not color:
                return jsonify({'success': False, 'error': 'El campo Color es obligatorio'})
            
            cursor = conn.cursor()
            
            # Verificar si la mascota existe
            cursor.execute("SELECT id_mascota, corte FROM mascotas WHERE id_mascota = %s", (id,))
            mascota_existente_raw = cursor.fetchone()
            
            if not mascota_existente_raw:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'error': 'Mascota no encontrada'})
            
            # Convertir a diccionario
            mascota_existente = dict(mascota_existente_raw)
            corte_anterior = mascota_existente.get('corte', None)
            
            # Preparar valores para la actualización
            valores = []
            campos = []
            
            # Agregar campos dinámicamente
            if raza:
                campos.append("raza = %s")
                valores.append(raza)
            
            if color:
                campos.append("color = %s")
                valores.append(color)
            
            if corte:
                campos.append("corte = %s")
                valores.append(corte)
            
            if tamano in ['pequeño', 'mediano', 'grande', 'gigante', 'pequeno']:
                # Normalizar "pequeno" sin ñ
                if tamano == 'pequeno':
                    tamano = 'pequeño'
                campos.append("tamano = %s")
                valores.append(tamano)
            
            if peso:
                try:
                    peso_float = float(peso)
                    campos.append("peso = %s")
                    valores.append(peso_float)
                except ValueError:
                    print(f"Advertencia: Peso no válido: {peso}")
            
            if caracteristicas:
                campos.append("caracteristicas = %s")
                valores.append(caracteristicas)
            
            if alergias:
                campos.append("alergias = %s")
                valores.append(alergias)
            
            # Agregar ID al final para la condición WHERE
            valores.append(id)
            
            # Construir y ejecutar la consulta
            if campos:
                sql = f"UPDATE mascotas SET {', '.join(campos)} WHERE id_mascota = %s"
                print(f"SQL ejecutado: {sql}")
                print(f"Valores: {valores}")
                
                cursor.execute(sql, valores)
                
                # Registrar en historial si el corte cambió - CORREGIDO PARA POSTGRESQL
                if corte and corte != corte_anterior:
                    id_empleado = session.get('id_empleado') if 'id_empleado' in session else None
                    
                    descripcion = f"Cambio de corte: {corte_anterior or 'Sin corte'} → {corte}"
                    notas = f"Actualizado desde el sistema de reservas"
                    
                    try:
                        cursor.execute("""
                            INSERT INTO historial_cortes 
                            (id_mascota, tipo_corte, descripcion, id_empleado, notas)
                            VALUES (%s, %s, %s, %s, %s)
                        """, (id, corte, descripcion, id_empleado, notas))
                        print("Registro en historial_cortes creado")
                    except Exception as hist_error:
                        print(f"Advertencia: No se pudo registrar en historial: {hist_error}")
                        # Continuar aunque falle el historial
                
                conn.commit()
                filas_afectadas = cursor.rowcount
                print(f"Filas afectadas: {filas_afectadas}")
                
                cursor.close()
                conn.close()
                
                if filas_afectadas > 0:
                    return jsonify({
                        'success': True, 
                        'message': 'Datos actualizados correctamente',
                        'corte_anterior': corte_anterior,
                        'corte_nuevo': corte
                    })
                else:
                    return jsonify({'success': False, 'error': 'No se realizaron cambios'})
            else:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'error': 'No se proporcionaron datos para actualizar'})
            
        except Exception as e:  # Cambia Error por Exception para PostgreSQL
            print(f"Error de PostgreSQL: {str(e)}")
            import traceback
            traceback.print_exc()
            if 'cursor' in locals() and cursor:
                cursor.close()
            if conn:
                conn.close()
            return jsonify({'success': False, 'error': f'Error de base de datos: {str(e)}'})
@app.route('/facturas/<int:id>')
def ver_factura(id):
    """Ver detalles de una factura - SOLO SERVICIOS"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('ventas'))
    
    try:
        cursor = conn.cursor()
        
        # 1. Obtener datos de la factura
        cursor.execute("""
            SELECT f.*, 
                   c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                   c.dni as cliente_dni, c.direccion as cliente_direccion,
                   c.telefono as cliente_telefono, c.email as cliente_email,
                   r.codigo_reserva, r.id_reserva, r.id_mascota,
                   e.nombre as empleado_nombre, e.apellido as empleado_apellido
            FROM facturas f
            LEFT JOIN clientes c ON f.id_cliente = c.id_cliente
            LEFT JOIN reservas r ON f.id_reserva = r.id_reserva
            LEFT JOIN empleados e ON f.id_empleado_cajero = e.id_empleado
            WHERE f.id_factura = %s
        """, (id,))
        
        factura_raw = cursor.fetchone()
        
        if not factura_raw:
            flash('Factura no encontrada.', 'danger')
            return redirect(url_for('ventas'))
        
        # Convertir a dict mutable
        factura = dict(factura_raw)
        
        print(f"🔍 Factura encontrada: {factura.get('numero', 'N/A')}")
        
        # 2. Obtener servicios de la factura (SOLO SERVICIOS)
        cursor.execute("""
            SELECT 
                fs.*, 
                s.categoria,
                COALESCE(fs.descripcion, s.nombre) as nombre_servicio
            FROM factura_servicios fs
            LEFT JOIN servicios s ON fs.id_servicio = s.id_servicio
            WHERE fs.id_factura = %s
            ORDER BY s.nombre
        """, (id,))
        
        servicios_raw = cursor.fetchall()
        
        # 3. Calcular totales
        total_servicios = 0
        servicios_procesados = []
        
        for servicio in servicios_raw:
            servicio_dict = dict(servicio)
            
            # Asegurar valores numéricos
            precio = float(servicio_dict.get('precio_unitario', 0))
            cantidad = int(servicio_dict.get('cantidad', 1))
            subtotal = float(servicio_dict.get('subtotal', 0))
            
            # Calcular subtotal si es 0
            if subtotal == 0 and precio > 0:
                subtotal = precio * cantidad
                servicio_dict['subtotal'] = subtotal
            
            total_servicios += subtotal
            servicios_procesados.append(servicio_dict)
        
        # 4. Preparar datos para el template
        factura['servicios'] = servicios_procesados
        factura['productos'] = []  # Lista vacía - NO MANEJAS PRODUCTOS
        
        # Si la factura no tiene total, usar el calculado
        if not factura.get('total') or float(factura.get('total', 0)) == 0:
            factura['total'] = total_servicios
        
        # Si no tiene subtotal/igv, calcular según tipo de comprobante
        if not factura.get('subtotal') or float(factura.get('subtotal', 0)) == 0:
            if factura.get('tipo_comprobante') == 'factura':
                # Para factura: calcular base e IGV
                subtotal_val = float(factura.get('total', 0)) / 1.18
                igv_val = subtotal_val * 0.18
                factura['subtotal'] = subtotal_val
                factura['igv'] = igv_val
            else:
                # Para boleta: subtotal = total, igv = 0
                factura['subtotal'] = float(factura.get('total', 0))
                factura['igv'] = 0.0
        
        factura['total_servicios'] = total_servicios
        factura['total_productos'] = 0
        factura['descuento_total'] = 0
        
        # 5. Formatear fecha para mostrar
        if factura.get('fecha_emision'):
            factura['fecha_emision_str'] = factura['fecha_emision'].strftime('%d/%m/%Y %H:%M')
        else:
            factura['fecha_emision_str'] = 'No disponible'
        
        # 6. Obtener datos de mascota si hay reserva
        if factura.get('id_mascota'):
            cursor.execute("""
                SELECT nombre, especie, raza 
                FROM mascotas 
                WHERE id_mascota = %s
            """, (factura['id_mascota'],))
            
            mascota = cursor.fetchone()
            if mascota:
                factura['mascota_nombre'] = mascota['nombre']
                factura['mascota_especie'] = mascota['especie']
                factura['mascota_raza'] = mascota['raza']
        
        # 7. Obtener pagos (opcional)
        try:
            cursor.execute("""
                SELECT * FROM pagos 
                WHERE id_factura = %s 
                ORDER BY fecha_pago DESC
            """, (id,))
            pagos = cursor.fetchall()
            factura['pagos'] = pagos
        except Exception as e:
            print(f"⚠️ Tabla pagos no existe: {e}")
            factura['pagos'] = []
        
        print(f"✅ Factura preparada:")
        print(f"   Número: {factura.get('numero')}")
        print(f"   Total: S/ {factura.get('total', 0):.2f}")
        print(f"   Servicios: {len(servicios_procesados)}")
        
    except Error as e:
        flash(f'Error obteniendo factura: {e}', 'danger')
        print(f"❌ Error en ver_factura: {e}")
        import traceback
        traceback.print_exc()
        return redirect(url_for('ventas'))
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn:
            conn.close()
    
    return render_template('facturas/ver.html', factura=factura)
    
@app.route('/reservas/ver/<int:id>')
def ver_reserva(id):
    """Ver detalles de la reserva"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reservas'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener datos principales de la reserva
        cursor.execute("""
            SELECT r.*, 
                   m.id_mascota, m.nombre as mascota_nombre, m.especie, m.raza,
                   c.id_cliente, c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                   c.telefono as cliente_telefono, c.email as cliente_email,
                   e.id_empleado, e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                   e.especialidad as empleado_especialidad
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            JOIN empleados e ON r.id_empleado = e.id_empleado
            WHERE r.id_reserva = %s
        """, (id,))
        reserva = cursor.fetchone()
        
        if not reserva:
            flash('Reserva no encontrada.', 'danger')
            return redirect(url_for('reservas'))
        
        # ========== ¡CONVERTIR A DICT MUTABLE! ==========
        reserva = dict(reserva)
        # ==============================================
        
        # Calcular tiempo restante si es pendiente
        if reserva['estado'] == 'pendiente' and reserva['fecha_reserva'] > datetime.now():
            tiempo_restante = reserva['fecha_reserva'] - datetime.now()
            dias = tiempo_restante.days
            horas = tiempo_restante.seconds // 3600
            minutos = (tiempo_restante.seconds % 3600) // 60
            reserva['tiempo_restante'] = f"{dias}d {horas}h {minutos}m"
        
        # Obtener servicios de la reserva
        cursor.execute("""
            SELECT rs.*, s.nombre, s.categoria, s.duracion_min
            FROM reserva_servicios rs
            JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE rs.id_reserva = %s
        """, (id,))
        servicios_raw = cursor.fetchall()
        
        # Convertir servicios a dicts mutables
        servicios = []
        total = 0
        for servicio in servicios_raw:
            servicio_dict = dict(servicio)
            
            # Calcular subtotal si no existe
            if 'subtotal' not in servicio_dict or not servicio_dict['subtotal']:
                if 'precio_unitario' in servicio_dict and servicio_dict['precio_unitario']:
                    cantidad = servicio_dict.get('cantidad', 1)
                    servicio_dict['subtotal'] = float(servicio_dict['precio_unitario']) * int(cantidad)
                else:
                    servicio_dict['subtotal'] = 0
            
            total += float(servicio_dict['subtotal'])
            servicios.append(servicio_dict)
        
        # Obtener factura asociada si existe
        cursor.execute("""
            SELECT f.* 
            FROM facturas f 
            WHERE f.id_reserva = %s
            ORDER BY f.fecha_emision DESC 
            LIMIT 1
        """, (id,))
        factura_raw = cursor.fetchone()
        
        # Convertir factura a dict si existe
        factura = None
        if factura_raw:
            factura = dict(factura_raw)
        
        # Ahora podemos agregar nuevas claves
        reserva['servicios'] = servicios
        reserva['total'] = total
        reserva['factura'] = factura
        
    except Error as e:
        flash(f'Error obteniendo datos de la reserva: {e}', 'danger')
        print(f"ERROR en ver_reserva: {e}")  # Para debug
        return redirect(url_for('reservas'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reservas/ver.html', reserva=reserva)
# ================= FUNCIONES PARA CORREO =================

def obtener_correo_admin():
    """Obtiene el correo del administrador del sistema"""
    try:
        # Simplemente devuelve el correo de configuración
        correo = app.config.get('ADMIN_EMAIL')
        if not correo:
            correo = 'ayumu79@gmail.com'  # Correo por defecto
        return correo
    except Exception as e:
        print(f"⚠️ Error en obtener_correo_admin: {e}")
        return 'ayumu79@gmail.com'

# En tu app.py, busca la función cambiar_estado_reserva y actualízala:
@app.route('/reservas/cambiar-estado/<int:id>', methods=['POST'])
@login_required
def cambiar_estado_reserva(id):
    """Cambiar estado de una reserva - VERSIÓN CORREGIDA"""
    try:
        print(f"\n📋 Intentando cambiar estado de reserva ID: {id}")
        
        data = request.get_json()
        nuevo_estado = data.get('estado')
        
        if not nuevo_estado:
            return jsonify({'success': False, 'message': 'Estado no proporcionado'}), 400
        
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'message': 'No hay conexión a la base de datos'}), 500
        
        try:
            cursor = conn.cursor()
            
            # 1. Obtener la reserva actual
            cursor.execute("""
                SELECT r.*, 
                       m.nombre as mascota_nombre, m.especie, m.raza,
                       c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                       c.telefono as cliente_telefono, c.email as cliente_email,
                       e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                       e.email as empleado_email
                FROM reservas r
                JOIN mascotas m ON r.id_mascota = m.id_mascota
                JOIN clientes c ON m.id_cliente = c.id_cliente
                JOIN empleados e ON r.id_empleado = e.id_empleado
                WHERE r.id_reserva = %s
            """, (id,))
            
            reserva = cursor.fetchone()
            
            if not reserva:
                return jsonify({'success': False, 'message': 'Reserva no encontrada'}), 404
            
            print(f"📌 Estado anterior: {reserva['estado']}")
            print(f"🎯 Estado nuevo: {nuevo_estado}")
            
            # Guardar estado anterior
            estado_anterior = reserva['estado']
            
            # 2. Actualizar reserva
            cursor.execute("""
                UPDATE reservas 
                SET estado = %s, fecha_modificacion = NOW()
                WHERE id_reserva = %s
            """, (nuevo_estado, id))
            
            conn.commit()
            
            print(f"✅ Estado cambiado en base de datos")
            
            # 3. Si el nuevo estado es "completada", enviar correo al admin
            if nuevo_estado == 'completada':
                print("🎉 Reserva completada, preparando notificación...")
                
                # Obtener servicios de la reserva para el correo
                cursor.execute("""
                    SELECT s.nombre
                    FROM reserva_servicios rs
                    JOIN servicios s ON rs.id_servicio = s.id_servicio
                    WHERE rs.id_reserva = %s
                """, (id,))
                
                servicios = cursor.fetchall()
                servicios_texto = ", ".join([s['nombre'] for s in servicios]) if servicios else "No especificados"
                
                # 🔥 CALCULAR TOTAL CORRECTAMENTE - VERSIÓN MEJORADA
                print("🔍 Calculando total de la reserva...")
                
                # Primero intentar obtener cantidad de cada servicio
                cursor.execute("""
                    SELECT 
                        rs.id_servicio,
                        rs.precio_unitario,
                        COALESCE(rs.cantidad, 1) as cantidad,  -- Si cantidad es NULL, usar 1
                        rs.subtotal,
                        s.precio as precio_servicio
                    FROM reserva_servicios rs
                    JOIN servicios s ON rs.id_servicio = s.id_servicio
                    WHERE rs.id_reserva = %s
                """, (id,))
                
                servicios_detalle = cursor.fetchall()
                print(f"🔍 Encontrados {len(servicios_detalle)} servicios en la reserva")
                
                # Calcular total de 3 formas diferentes
                total_method1 = 0.0  # Usando subtotal
                total_method2 = 0.0  # Usando precio_unitario * cantidad
                total_method3 = 0.0  # Usando precio_servicio * cantidad
                
                for i, servicio in enumerate(servicios_detalle):
                    print(f"🔍 Servicio {i+1}:")
                    print(f"   precio_unitario: {servicio['precio_unitario']}")
                    print(f"   cantidad: {servicio['cantidad']}")
                    print(f"   subtotal: {servicio['subtotal']}")
                    print(f"   precio_servicio: {servicio['precio_servicio']}")
                    
                    # Método 1: Usar subtotal directamente
                    if servicio['subtotal'] is not None:
                        total_method1 += float(servicio['subtotal'])
                    
                    # Método 2: precio_unitario * cantidad
                    if servicio['precio_unitario'] is not None:
                        cantidad = int(servicio['cantidad']) if servicio['cantidad'] else 1
                        total_method2 += float(servicio['precio_unitario']) * cantidad
                    
                    # Método 3: precio_servicio * cantidad
                    if servicio['precio_servicio'] is not None:
                        cantidad = int(servicio['cantidad']) if servicio['cantidad'] else 1
                        total_method3 += float(servicio['precio_servicio']) * cantidad
                
                print(f"🔍 Totales calculados:")
                print(f"   Método 1 (subtotal): {total_method1:.2f}")
                print(f"   Método 2 (precio_unitario * cantidad): {total_method2:.2f}")
                print(f"   Método 3 (precio_servicio * cantidad): {total_method3:.2f}")
                
                # Elegir el total que no sea 0
                if total_method1 > 0:
                    total = total_method1
                    print(f"✅ Usando total_method1: {total:.2f}")
                elif total_method2 > 0:
                    total = total_method2
                    print(f"✅ Usando total_method2: {total:.2f}")
                elif total_method3 > 0:
                    total = total_method3
                    print(f"✅ Usando total_method3: {total:.2f}")
                else:
                    total = 0.0
                    print("⚠️ Todos los métodos devolvieron 0")
                
                # Si todo falla, hacer consulta directa
                if total == 0:
                    print("🔍 Intentando consulta directa de total...")
                    cursor.execute("""
                        SELECT COALESCE(SUM(subtotal), 0) as total_directo
                        FROM reserva_servicios
                        WHERE id_reserva = %s
                    """, (id,))
                    
                    total_directo = cursor.fetchone()
                    if total_directo and total_directo['total_directo']:
                        total = float(total_directo['total_directo'])
                        print(f"✅ Total directo de BD: {total:.2f}")
                
                print(f"🎯 Total final para correo: S/ {total:.2f}")
                
                # 🔥 CREAR DICCIONARIO CON LOS DATOS - VERSIÓN MEJORADA
                reserva_dict = {
                    'codigo_reserva': reserva['codigo_reserva'],
                    'fecha_reserva': reserva['fecha_reserva'],
                    'mascota_nombre': reserva['mascota_nombre'],
                    'especie': reserva['especie'],
                    'raza': reserva['raza'],
                    'cliente_nombre': reserva['cliente_nombre'],
                    'cliente_apellido': reserva['cliente_apellido'],
                    'cliente_telefono': reserva['cliente_telefono'],
                    'cliente_email': reserva['cliente_email'],
                    'empleado_nombre': reserva['empleado_nombre'],
                    'empleado_apellido': reserva['empleado_apellido'],
                    'empleado_especialidad': 'No especificada',  # Esto no está en tu consulta
                    'notas': reserva.get('notas', 'No hay notas'),
                    'total': float(total),
                    'servicios_texto': servicios_texto,
                    'id_reserva': id  # Añadir ID para debug
                }
                
                print(f"🔍 Enviando reserva_dict a función de correo:")
                for key, value in reserva_dict.items():
                    print(f"   {key}: {value}")

                resultado_correo = enviar_correo_reserva_completada(reserva_dict)
                
                if resultado_correo:
                    print("✅ Notificación enviada exitosamente")
                else:
                    print("⚠️ Notificación no pudo ser enviada, pero la reserva se actualizó")
            
            cursor.close()
            conn.close()
            
            return jsonify({
                'success': True, 
                'message': f'Reserva {reserva["codigo_reserva"]} actualizada a {nuevo_estado}',
                'estado': nuevo_estado
            })
            
        except Error as e:
            if conn:
                conn.rollback()
            print(f"❌ Error cambiando estado: {str(e)}")
            return jsonify({'success': False, 'message': f'Error de base de datos: {str(e)}'}), 500
            
    except Exception as e:
        print(f"❌ Error general: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': str(e)}), 500

        # ================= RUTA DE PRUEBA =================

@app.route('/test-correo/<int:id>')
@login_required
def test_correo(id):
    """Ruta para probar el envío de correo sin cambiar estado"""
    try:
        conn = get_db_connection()
        if not conn:
            flash('No hay conexión a la base de datos.', 'danger')
            return redirect(url_for('reservas'))
        
        cursor = conn.cursor()
        
        # Obtener datos de la reserva
        cursor.execute("""
            SELECT r.*, 
                   m.nombre as mascota_nombre, m.especie, m.raza,
                   c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                   c.telefono as cliente_telefono, c.email as cliente_email,
                   e.nombre as empleado_nombre, e.apellido as empleado_apellido,
                   e.email as empleado_email
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            JOIN empleados e ON r.id_empleado = e.id_empleado
            WHERE r.id_reserva = %s
        """, (id,))
        
        reserva = cursor.fetchone()
        
        if not reserva:
            flash('Reserva no encontrada.', 'danger')
            return redirect(url_for('reservas'))
        
        # Obtener total
        cursor.execute("SELECT SUM(subtotal) as total FROM reserva_servicios WHERE id_reserva = %s", (id,))
        total_result = cursor.fetchone()
        total = total_result['total'] if total_result and total_result['total'] else 0
        
        cursor.close()
        conn.close()
        
        # Crear objeto para el correo
        # Crear diccionario para el correo
        reserva_dict = {
    'codigo_reserva': reserva['codigo_reserva'],
    'fecha_reserva': reserva['fecha_reserva'],
    'mascota_nombre': reserva['mascota_nombre'],
    'especie': reserva['especie'],
    'raza': reserva['raza'],
    'cliente_nombre': reserva['cliente_nombre'],
    'cliente_apellido': reserva['cliente_apellido'],
    'cliente_telefono': reserva['cliente_telefono'],
    'cliente_email': reserva['cliente_email'],
    'empleado_nombre': reserva['empleado_nombre'],
    'empleado_apellido': reserva['empleado_apellido'],
    'empleado_especialidad': '',
    'notas': reserva.get('notas', ''),
    'total': float(total)
        }

        resultado = enviar_correo_reserva_completada(reserva_dict)
        
        if resultado:
            flash(f'Correo de prueba generado para reserva {reserva["codigo_reserva"]}', 'success')
        else:
            flash('Error al generar correo de prueba', 'danger')
            
        return redirect(url_for('ver_reserva', id=id))
        
    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('reservas'))
# ================= RUTAS DE FACTURACIÓN =================

@app.route('/reservas/<int:id>/facturar', methods=['GET', 'POST'])
def facturar_reserva(id):
    """Crear factura a partir de una reserva"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('ver_reserva', id=id))
    
    try:
        cursor = conn.cursor()
        
        # 1. Verificar que la reserva existe y está completada
        cursor.execute("""
            SELECT r.*, 
                   m.id_mascota, m.nombre as mascota_nombre,
                   c.id_cliente, c.nombre as cliente_nombre, c.apellido as cliente_apellido,
                   c.dni as cliente_dni, c.direccion as cliente_direccion
            FROM reservas r
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            JOIN clientes c ON m.id_cliente = c.id_cliente
            WHERE r.id_reserva = %s
        """, (id,))
        reserva = cursor.fetchone()
        
        if not reserva:
            flash('Reserva no encontrada.', 'danger')
            return redirect(url_for('reservas'))
        
        # Convertir a dict mutable
        reserva = dict(reserva)
        
        print(f"🔍 Reserva ID {id} encontrada. Estado: {reserva['estado']}")
        
        if reserva['estado'] != 'completada':
            flash('Solo se pueden facturar reservas completadas.', 'warning')
            return redirect(url_for('ver_reserva', id=id))
        
        # 2. Verificar que no tenga factura ya
        cursor.execute("SELECT * FROM facturas WHERE id_reserva = %s", (id,))
        factura_existente = cursor.fetchone()
        
        if factura_existente:
            flash('Esta reserva ya tiene una factura.', 'info')
            return redirect(url_for('ver_factura', id=factura_existente['id_factura']))
        
        if request.method == 'POST':
            print(f"🔍 POST recibido para facturar reserva {id}")
            
            # 3. Obtener datos del formulario
            tipo_comprobante = request.form.get('tipo_comprobante', 'boleta')
            metodo_pago = request.form.get('metodo_pago', 'efectivo')
            notas = request.form.get('notas', '').strip()
            
            print(f"🔍 Tipo comprobante: {tipo_comprobante}")
            print(f"🔍 Método pago: {metodo_pago}")
            
            # Validar tipo de comprobante según DNI
            if tipo_comprobante == 'factura' and not reserva['cliente_dni']:
                flash('Para emitir factura el cliente debe tener DNI registrado.', 'danger')
                return redirect(url_for('facturar_reserva', id=id))
            
            # 4. Obtener servicios para calcular totales (MANEJO DE NULLS)
            cursor.execute("""
                SELECT 
                    COALESCE(SUM(subtotal), 0) as total_subtotal,
                    COALESCE(SUM(precio_unitario * COALESCE(cantidad, 1)), 0) as total_calculado
                FROM reserva_servicios
                WHERE id_reserva = %s
            """, (id,))
            
            total_result = cursor.fetchone()
            print(f"🔍 Resultados total: {total_result}")
            
            # Usar el primer total no-cero
            if total_result['total_subtotal'] and float(total_result['total_subtotal']) > 0:
                total_servicios = float(total_result['total_subtotal'])
            elif total_result['total_calculado'] and float(total_result['total_calculado']) > 0:
                total_servicios = float(total_result['total_calculado'])
            else:
                total_servicios = 0.0
            
            print(f"🔍 Total servicios: {total_servicios}")
            
            # 5. Calcular según tipo de comprobante
            if tipo_comprobante == 'factura':
                # Para factura: calcular IGV (18%)
                if total_servicios > 0:
                    subtotal = total_servicios / 1.18  # Base imponible
                    igv = subtotal * 0.18  # IGV 18%
                    total = subtotal + igv
                else:
                    subtotal = 0.0
                    igv = 0.0
                    total = 0.0
                    
                print(f"🔍 FACTURA - Subtotal: {subtotal:.2f}, IGV: {igv:.2f}, Total: {total:.2f}")
            else:
                # Para boleta: no hay IGV (para régimen especial)
                subtotal = total_servicios
                igv = 0.00
                total = total_servicios
                print(f"🔍 BOLETA - Subtotal: {subtotal:.2f}, IGV: {igv:.2f}, Total: {total:.2f}")
            
            # 6. Generar número de factura/boleta
            cursor.execute("SELECT COALESCE(MAX(id_factura), 0) + 1 as next_id FROM facturas")
            next_id_result = cursor.fetchone()
            next_id = next_id_result['next_id']
            
            serie = 'B001' if tipo_comprobante == 'boleta' else 'F001'
            numero = f"{next_id:04d}"
            
            print(f"🔍 Creando {tipo_comprobante.upper()} {serie}-{numero}")
            
            # 7. Crear factura con RETURNING para obtener ID (POSTGRESQL)
            cursor.execute("""
                INSERT INTO facturas (
                    serie, numero, tipo_comprobante, id_cliente, id_reserva,
                    subtotal, igv, total, metodo_pago, notas, 
                    id_empleado_cajero, estado, fecha_emision
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                RETURNING id_factura
            """, (
                serie, numero, tipo_comprobante, reserva['id_cliente'], id,
                round(subtotal, 2), round(igv, 2), round(total, 2),
                metodo_pago, notas or None, 
                1, 'pendiente'
            ))
            
            # Obtener el ID de la factura creada
            result = cursor.fetchone()
            if not result or 'id_factura' not in result:
                raise Exception("No se pudo obtener el ID de la factura creada")
            
            id_factura = result['id_factura']
            print(f"✅ Factura creada con ID: {id_factura}")
            
            # 8. Copiar servicios de reserva a factura_servicios (MANEJO DE NULLS)
            cursor.execute("""
                INSERT INTO factura_servicios (
                    id_factura, id_servicio, descripcion, 
                    precio_unitario, cantidad, subtotal
                )
                SELECT 
                    %s,
                    rs.id_servicio,
                    s.nombre,
                    COALESCE(rs.precio_unitario, s.precio, 0),
                    COALESCE(rs.cantidad, 1),
                    COALESCE(
                        rs.subtotal, 
                        COALESCE(rs.precio_unitario, s.precio, 0) * COALESCE(rs.cantidad, 1),
                        0
                    )
                FROM reserva_servicios rs
                JOIN servicios s ON rs.id_servicio = s.id_servicio
                WHERE rs.id_reserva = %s
            """, (id_factura, id))
            
            # 9. Actualizar estado de la reserva a "facturada" (opcional)
            cursor.execute("""
                UPDATE reservas 
                SET estado = 'facturada' 
                WHERE id_reserva = %s
            """, (id,))
            
            conn.commit()
            
            flash(f'{tipo_comprobante.capitalize()} {serie}-{numero} creada exitosamente.', 'success')
            return redirect(url_for('ver_factura', id=id_factura))
        
        # ========== GET: Mostrar formulario de facturación ==========
        print(f"🔍 GET: Mostrando formulario para reserva {id}")
        
        # Obtener servicios de la reserva para mostrar resumen (MANEJO DE NULLS)
        cursor.execute("""
            SELECT 
                s.id_servicio,
                s.nombre,
                s.categoria,
                COALESCE(rs.precio_unitario, s.precio, 0) as precio_unitario,
                COALESCE(rs.cantidad, 1) as cantidad,
                COALESCE(
                    rs.subtotal, 
                    COALESCE(rs.precio_unitario, s.precio, 0) * COALESCE(rs.cantidad, 1),
                    0
                ) as subtotal
            FROM reserva_servicios rs
            JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE rs.id_reserva = %s
            ORDER BY s.nombre
        """, (id,))
        
        servicios_raw = cursor.fetchall()
        servicios = []
        total = 0.0
        
        print(f"🔍 Encontrados {len(servicios_raw)} servicios")
        
        for i, servicio in enumerate(servicios_raw):
            servicio_dict = dict(servicio)
            
            # Asegurar que todos los valores sean numéricos
            precio = float(servicio_dict.get('precio_unitario', 0))
            cantidad = int(servicio_dict.get('cantidad', 1))
            subtotal_serv = float(servicio_dict.get('subtotal', 0))
            
            # Si subtotal es 0, calcularlo
            if subtotal_serv == 0 and precio > 0:
                subtotal_serv = precio * cantidad
                servicio_dict['subtotal'] = subtotal_serv
            
            servicio_dict['precio_unitario'] = precio
            servicio_dict['cantidad'] = cantidad
            servicio_dict['subtotal'] = subtotal_serv
            
            total += subtotal_serv
            servicios.append(servicio_dict)
            
            print(f"🔍 Servicio {i+1}: {servicio_dict['nombre']} - Precio: {precio}, Cant: {cantidad}, Subtotal: {subtotal_serv}")
        
        print(f"🔍 Total general: {total:.2f}")
        
        # Calcular para mostrar en formulario
        if total > 0:
            subtotal_display = total / 1.18  # Perú tiene 18% IGV
            igv_display = total - subtotal_display
        else:
            subtotal_display = 0.0
            igv_display = 0.0
        
        # Verificar si cliente tiene DNI para factura
        puede_factura = bool(reserva.get('cliente_dni'))
        print(f"🔍 Cliente tiene DNI: {puede_factura} ({reserva.get('cliente_dni')})")
        
        # Obtener datos del cliente para mostrar
        cliente_info = {
            'nombre_completo': f"{reserva.get('cliente_nombre', '')} {reserva.get('cliente_apellido', '')}".strip(),
            'dni': reserva.get('cliente_dni', 'No registrado'),
            'direccion': reserva.get('cliente_direccion', 'No registrada')
        }
        
    except Error as e:
        if conn:
            conn.rollback()
        flash(f'Error creando factura: {e}', 'danger')
        print(f"❌ Error en facturar_reserva: {e}")
        import traceback
        traceback.print_exc()
        return redirect(url_for('ver_reserva', id=id))
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'Error inesperado: {e}', 'danger')
        print(f"❌ Error general en facturar_reserva: {e}")
        import traceback
        traceback.print_exc()
        return redirect(url_for('ver_reserva', id=id))
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn:
            conn.close()
    
    return render_template('facturas/crear.html', 
                         reserva=reserva, 
                         servicios=servicios,
                         total=round(total, 2),
                         subtotal=round(subtotal_display, 2),
                         igv=round(igv_display, 2),
                         puede_factura=puede_factura,
                         cliente_info=cliente_info)
# ================= RUTAS DE REPORTES Y CONFIGURACIÓN =================

# Ruta para configuración PRINCIPAL
@app.route('/config')
def config():  # <- NOMBRE ORIGINAL
    """Página principal de configuración"""
    if 'id_usuario' not in session:
        return redirect(url_for('login'))
    
    return render_template('config.html')


# ================= MANEJO DE ERRORES =================

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500

# ================= CONTEXT PROCESSORS =================

@app.context_processor
def inject_now():
    """Inyectar fecha actual en todas las plantillas"""
    from datetime import datetime
    return {'now': datetime.now()}

@app.context_processor
def inject_user_data():
    """Inyectar datos del usuario (ahora siempre con usuario demo)"""
    return {
        'current_user': {
            'id': 1,
            'username': 'admin',
            'rol': 'admin',
            'nombre': 'Administrador'
        },
        'user_role': 'admin'
    }
# Ruta para configurar la contraseña
@app.route('/configurar_password', methods=['POST'])
def configurar_password():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'})
    
    data = request.get_json()
    tipo = data.get('tipo', 'reportes')
    password = data.get('password', '')
    
    config_path = 'config_reportes.json'
    
    try:
        # Leer configuración existente
        if os.path.exists(config_path):
            with open(config_path, 'r') as f:
                config = json.load(f)
        else:
            config = {}
        
        # Actualizar contraseña
        config[f'password_{tipo}'] = password
        
        # Guardar
        with open(config_path, 'w') as f:
            json.dump(config, f, indent=4)
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error guardando configuración: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/check_session')
def check_session():
    """Verificar sesión para el modal"""
    return jsonify({'logged_in': 'id_usuario' in session})


# ============================================
# RUTAS PARA MANEJO DE CAJA
# ============================================


@app.route('/caja/apertura', methods=['GET', 'POST'])
def apertura_caja():
    from datetime import datetime
    """Abrir caja del día"""
    if request.method == 'POST':
        try:
            monto_apertura = float(request.form.get('monto_apertura', 0))
            id_empleado = session.get('id_empleado', 1)  # Usa el ID de la sesión
            
            if monto_apertura <= 0:
                flash('El monto de apertura debe ser mayor a 0.', 'danger')
                return redirect(url_for('apertura_caja'))
            
            conn = get_db_connection()
            if conn:
                cursor = conn.cursor()
                
                # Verificar si ya hay caja abierta hoy para este empleado
                cursor.execute("""
                    SELECT id_caja FROM caja_diaria 
                    WHERE fecha = CURRENT_DATE 
                    AND id_empleado_cajero = %s
                    AND estado = 'abierta'
                """, (id_empleado,))
                
                caja_existente = cursor.fetchone()
                
                if caja_existente:
                    flash('Ya tienes una caja abierta para hoy.', 'warning')
                else:
                    # Crear nueva caja
                    cursor.execute("""
                        INSERT INTO caja_diaria 
                        (fecha, id_empleado_cajero, monto_apertura, estado, hora_apertura)
                        VALUES (CURRENT_DATE, %s, %s, 'abierta', NOW())
                    """, (id_empleado, monto_apertura))
                    
                    conn.commit()
                    flash(f'✅ Caja abierta con S/ {monto_apertura:.2f} exitosamente.', 'success')
                    return redirect(url_for('dashboard'))
                
                cursor.close()
                conn.close()
            else:
                flash('❌ Error de conexión a la base de datos.', 'danger')
            
        except Exception as e:
            flash(f'❌ Error al abrir caja: {str(e)}', 'danger')
    
    # GET request: mostrar formulario
    date_now = datetime.now()
    return render_template('caja/apertura.html', date_now=date_now)

def safe_float(value, default=0.0):
    """Convertir cualquier valor a float de forma segura"""
    if value is None:
        return default
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


@app.route('/caja/cierre', methods=['GET', 'POST'])
def cierre_caja():
    """Cerrar caja del día - VERSIÓN CORREGIDA"""
    conn = get_db_connection()
    
    if not conn:
        flash('❌ Error de conexión a la base de datos.', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        cursor = conn.cursor()
        id_empleado = session.get('id_empleado', 1)
        
        # Obtener caja abierta actual
        cursor.execute("""
            SELECT c.*, CONCAT(e.nombre, ' ', e.apellido) as cajero
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            WHERE fecha = CURRENT_DATE 
            AND estado = 'abierta'
            AND c.id_empleado_cajero = %s
        """, (id_empleado,))
        
        caja_actual = cursor.fetchone()
        
        if not caja_actual:
            flash('ℹ️ No tienes caja abierta para hoy.', 'info')
            return redirect(url_for('dashboard'))
        
        if request.method == 'POST':
            monto_cierre = float(request.form.get('monto_cierre', 0))
            observaciones = request.form.get('observaciones', '')
            
            # Convertir todos los valores a float
            monto_apertura = safe_float(caja_actual['monto_apertura'])
            venta_efectivo = safe_float(caja_actual['venta_efectivo'])
            venta_tarjeta = safe_float(caja_actual['venta_tarjeta'])
            venta_digital = safe_float(caja_actual['venta_digital'])
            total_ventas = safe_float(caja_actual['total_ventas'])
            
            # Calcular diferencia
            efectivo_esperado = monto_apertura + venta_efectivo
            diferencia = monto_cierre - efectivo_esperado
            
            print(f"🔢 Valores calculados:")
            print(f"   Monto cierre: {monto_cierre}")
            print(f"   Monto apertura: {monto_apertura}")
            print(f"   Venta efectivo: {venta_efectivo}")
            print(f"   Efectivo esperado: {efectivo_esperado}")
            print(f"   Diferencia: {diferencia}")
            
            # Actualizar caja
            cursor.execute("""
                UPDATE caja_diaria 
                SET monto_cierre = %s,
                    diferencia = %s,
                    observaciones = %s,
                    estado = 'cerrada',
                    hora_cierre = NOW()
                WHERE id_caja = %s
            """, (monto_cierre, diferencia, observaciones, caja_actual['id_caja']))
            
            conn.commit()
            
            # Determinar mensaje según diferencia
            if abs(diferencia) < 0.01:  # Menos de 1 céntimo
                mensaje = f'✅ Caja cerrada perfectamente. Sin diferencia.'
            elif diferencia > 0:
                mensaje = f'✅ Caja cerrada. Sobrante: S/ {diferencia:.2f}'
            else:
                mensaje = f'✅ Caja cerrada. Faltante: S/ {abs(diferencia):.2f}'
            
            flash(mensaje, 'success')
            return redirect(url_for('dashboard'))
        
        # GET request: preparar datos para la plantilla
        # Convertir todos los valores Decimal a float para la plantilla
        for key in ['monto_apertura', 'venta_efectivo', 'venta_tarjeta', 
                   'venta_digital', 'total_ventas', 'monto_cierre', 'diferencia']:
            if key in caja_actual:
                caja_actual[f'{key}_float'] = safe_float(caja_actual[key])
        
        # Calcular efectivo esperado para mostrar
        efectivo_esperado = safe_float(caja_actual['monto_apertura']) + safe_float(caja_actual['venta_efectivo'])
        caja_actual['efectivo_esperado'] = efectivo_esperado
        
        return render_template('caja/cierre.html', caja=caja_actual)
        
    except Exception as e:
        flash(f'❌ Error al cerrar caja: {str(e)}', 'danger')
        import traceback
        print(f"🔍 Error completo: {traceback.format_exc()}")
        return redirect(url_for('dashboard'))
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/api/caja/estado')
def estado_caja():
    """API para verificar estado de caja (para AJAX)"""
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'caja_abierta': False, 'message': 'Sin conexión a BD'})
    
    try:
        cursor = conn.cursor()
        id_empleado = session.get('id_empleado', 1)
        
        cursor.execute("""
            SELECT c.*, CONCAT(e.nombre, ' ', e.apellido) as cajero
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            WHERE fecha = CURRENT_DATE 
            AND estado = 'abierta'
            AND c.id_empleado_cajero = %s
        """, (id_empleado,))
        
        caja = cursor.fetchone()
        
        return jsonify({
            'caja_abierta': caja is not None,
            'caja': caja
        })
        
    except Exception as e:
        return jsonify({'caja_abierta': False, 'message': str(e)})
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/caja/historial')
def historial_caja():
    """Ver historial de cajas"""
    conn = get_db_connection()
    
    if not conn:
        flash('❌ Error de conexión a la base de datos.', 'danger')
        return redirect(url_for('dashboard'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener últimas 30 cajas
        cursor.execute("""
            SELECT c.*, CONCAT(e.nombre, ' ', e.apellido) as cajero
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            ORDER BY c.fecha DESC, c.hora_apertura DESC
            LIMIT 30
        """)
        
        cajas = cursor.fetchall()
        
        return render_template('caja/historial.html', cajas=cajas)
        
    except Exception as e:
        flash(f'❌ Error obteniendo historial: {str(e)}', 'danger')
        return redirect(url_for('dashboard'))
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()
            
@app.template_filter('igv_subtotal')
def igv_subtotal_filter(total):
    """Calcular subtotal sin IGV (18%)"""
    try:
        return float(total) / 1.18
    except (ValueError, TypeError):
        return 0.0

@app.template_filter('igv_amount')
def igv_amount_filter(total):
    """Calcular monto de IGV (18%)"""
    try:
        return float(total) - (float(total) / 1.18)
    except (ValueError, TypeError):
        return 0.0

@app.route('/debug/routes')
def debug_routes():
    """Mostrar todas las rutas registradas"""
    routes = []
    for rule in app.url_map.iter_rules():
        routes.append({
            'endpoint': rule.endpoint,
            'methods': list(rule.methods),
            'rule': str(rule)
        })
    routes.sort(key=lambda x: x['endpoint'])
    
    # Verificar específicamente las rutas de caja
    caja_routes = [r for r in routes if 'caja' in r['endpoint'].lower()]
    
    return jsonify({
        'total_routes': len(routes),
        'caja_routes': caja_routes,
        'tiene_apertura_caja': any(r['endpoint'] == 'apertura_caja' for r in routes)
    })
@app.route('/ventas/crear', methods=['GET', 'POST'])  # Añade POST
def crear_venta():
    """Crear nueva venta directa (sin reserva) - SOLO SERVICIOS"""
    conn = get_db_connection()
    
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('ventas'))
    
    try:
        cursor = conn.cursor()
        
        if request.method == 'POST':
            # ========== PROCESAR VENTA ==========
            print("🔍 Procesando nueva venta directa...")
            
            id_cliente = request.form.get('id_cliente', '').strip()
            tipo_comprobante = request.form.get('tipo_comprobante', 'boleta')
            metodo_pago = request.form.get('metodo_pago', 'efectivo')
            servicios = request.form.getlist('servicios[]')
            cantidades = request.form.getlist('cantidades[]')
            notas = request.form.get('notas', '').strip()
            
            print(f"🔍 Datos recibidos:")
            print(f"   Cliente: {id_cliente}")
            print(f"   Tipo: {tipo_comprobante}")
            print(f"   Método pago: {metodo_pago}")
            print(f"   Servicios: {servicios}")
            print(f"   Cantidades: {cantidades}")
            
            # Validaciones
            if not servicios:
                flash('Debe seleccionar al menos un servicio.', 'danger')
                return redirect(url_for('crear_venta'))
            
            if tipo_comprobante == 'factura' and (not id_cliente or id_cliente == '0'):
                flash('Para factura debe seleccionar un cliente.', 'danger')
                return redirect(url_for('crear_venta'))
            
            # Validar DNI para factura
            if tipo_comprobante == 'factura' and id_cliente and id_cliente != '0':
                cursor.execute("SELECT dni FROM clientes WHERE id_cliente = %s", (id_cliente,))
                cliente = cursor.fetchone()
                if not cliente or not cliente['dni']:
                    flash('Para factura, el cliente debe tener DNI registrado.', 'danger')
                    return redirect(url_for('crear_venta'))
            
            # Calcular total
            total = 0.0
            servicios_detalle = []
            
            for i, servicio_id in enumerate(servicios):
                cursor.execute("""
                    SELECT id_servicio, nombre, precio 
                    FROM servicios 
                    WHERE id_servicio = %s AND activo = TRUE
                """, (servicio_id,))
                
                servicio = cursor.fetchone()
                
                if servicio:
                    cantidad = int(cantidades[i]) if i < len(cantidades) and cantidades[i] else 1
                    precio_unit = float(servicio['precio'])
                    subtotal = precio_unit * cantidad
                    total += subtotal
                    
                    servicios_detalle.append({
                        'id_servicio': servicio_id,
                        'nombre': servicio['nombre'],
                        'precio_unitario': precio_unit,
                        'cantidad': cantidad,
                        'subtotal': subtotal
                    })
                    
                    print(f"🔍 Servicio {i+1}: {servicio['nombre']} x{cantidad} = S/ {subtotal:.2f}")
            
            print(f"🔍 Total calculado: S/ {total:.2f}")
            
            # Calcular IGV según tipo de comprobante
            if tipo_comprobante == 'factura' and total > 0:
                subtotal_val = total / 1.18  # Base imponible
                igv = subtotal_val * 0.18    # IGV 18%
                print(f"🔍 FACTURA - Base: {subtotal_val:.2f}, IGV: {igv:.2f}")
            else:
                subtotal_val = total
                igv = 0.0
                print(f"🔍 BOLETA - Total: {total:.2f}, IGV: 0.00")
            
            # Generar número de comprobante
            cursor.execute("SELECT COALESCE(MAX(id_factura), 0) + 1 as next_id FROM facturas")
            next_id = cursor.fetchone()['next_id']
            
            serie = 'B001' if tipo_comprobante == 'boleta' else 'F001'
            numero = f"{next_id:04d}"
            
            print(f"🔍 Creando {tipo_comprobante.upper()} {serie}-{numero}")
            
            # Preparar datos del cliente
            cliente_id = int(id_cliente) if id_cliente and id_cliente != '0' else None
            
            # Crear factura
            cursor.execute("""
                INSERT INTO facturas (
                    serie, numero, tipo_comprobante, id_cliente,
                    subtotal, igv, total, metodo_pago, notas, 
                    id_empleado_cajero, estado, fecha_emision
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                RETURNING id_factura
            """, (
                serie, numero, tipo_comprobante, cliente_id,
                round(subtotal_val, 2), round(igv, 2), round(total, 2),
                metodo_pago, notas or None, 
                session.get('id_usuario', 1), 'pagada'  # Venta directa = pagada
            ))
            
            result = cursor.fetchone()
            id_factura = result['id_factura'] if result else None
            
            if not id_factura:
                raise Exception("No se pudo obtener el ID de la factura creada")
            
            print(f"✅ Factura creada con ID: {id_factura}")
            
            # Agregar servicios a la factura
            for servicio in servicios_detalle:
                cursor.execute("""
                    INSERT INTO factura_servicios (
                        id_factura, id_servicio, descripcion, 
                        precio_unitario, cantidad, subtotal
                    ) VALUES (%s, %s, %s, %s, %s, %s)
                """, (
                    id_factura, servicio['id_servicio'], servicio['nombre'],
                    servicio['precio_unitario'], servicio['cantidad'], servicio['subtotal']
                ))
            
            conn.commit()
            
            flash(f'Venta {serie}-{numero} creada exitosamente. Total: S/ {total:.2f}', 'success')
            return redirect(url_for('ver_factura', id=id_factura))
        
        # ========== GET: Mostrar formulario ==========
        print("🔍 Mostrando formulario de venta...")
        
        # Obtener clientes para el select
        cursor.execute("""
            SELECT id_cliente, nombre, apellido, dni 
            FROM clientes 
            ORDER BY apellido, nombre
        """)
        clientes = cursor.fetchall()
        
        # Obtener servicios activos
        cursor.execute("""
            SELECT id_servicio, nombre, precio, categoria, descripcion 
            FROM servicios 
            WHERE activo = TRUE 
            ORDER BY categoria, nombre
        """)
        servicios = cursor.fetchall()
        
        # Agrupar servicios por categoría
        servicios_por_categoria = {}
        for servicio in servicios:
            categoria = servicio['categoria']
            if categoria not in servicios_por_categoria:
                servicios_por_categoria[categoria] = []
            servicios_por_categoria[categoria].append(servicio)
        
        print(f"🔍 Encontrados {len(servicios)} servicios en {len(servicios_por_categoria)} categorías")
        
        return render_template('ventas/crear.html', 
                             clientes=clientes, 
                             servicios_por_categoria=servicios_por_categoria,
                             productos=[])  # Lista vacía ya que no manejas productos
        
    except Error as e:
        if conn:
            conn.rollback()
        flash(f'Error creando venta: {e}', 'danger')
        print(f"❌ Error en crear_venta: {e}")
        import traceback
        traceback.print_exc()
        return redirect(url_for('ventas'))
    except Exception as e:
        if conn:
            conn.rollback()
        flash(f'Error inesperado: {e}', 'danger')
        print(f"❌ Error general en crear_venta: {e}")
        import traceback
        traceback.print_exc()
        return redirect(url_for('ventas'))
    finally:
        if 'cursor' in locals():
            cursor.close()
        if conn:
            conn.close()     


@app.route('/facturas/<int:id>/pagar', methods=['POST'])
def pagar_factura(id):
    """Registrar pago de factura - VERSIÓN CORREGIDA"""
    print(f"🔍 DEBUG: Pago para factura {id}")
    
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        # Verificar que recibimos JSON
        if not request.is_json:
            return jsonify({'success': False, 'message': 'Solo se aceptan solicitudes JSON.'}), 400
        
        data = request.get_json()
        print(f"📦 Datos recibidos: {data}")
        
        if not data:
            return jsonify({'success': False, 'message': 'No se recibieron datos.'}), 400
        
        monto_pagado = data.get('monto', 0)
        metodo_pago = data.get('metodo_pago', 'efectivo')
        es_parcial = data.get('es_parcial', False)
        
        # Convertir y validar monto
        try:
            monto_pagado = float(monto_pagado)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Monto inválido.'}), 400
        
        if monto_pagado <= 0:
            return jsonify({'success': False, 'message': 'El monto debe ser mayor a 0.'}), 400
        
        cursor = conn.cursor()
        
        # 1. Obtener factura
        cursor.execute("""
            SELECT f.*, c.nombre, c.apellido, c.dni
            FROM facturas f
            LEFT JOIN clientes c ON f.id_cliente = c.id_cliente
            WHERE f.id_factura = %s
        """, (id,))
        
        factura = cursor.fetchone()
        print(f"📄 Factura obtenida: {factura}")
        
        if not factura:
            return jsonify({'success': False, 'message': 'Factura no encontrada.'}), 404
        
        if factura['estado'] == 'pagada':
            return jsonify({'success': False, 'message': 'La factura ya está pagada.'}), 400
        
        if factura['estado'] == 'anulada':
            return jsonify({'success': False, 'message': 'No se puede pagar una factura anulada.'}), 400
        
        # 2. Verificar monto
        total_factura = safe_float(factura['total'])
        if monto_pagado > total_factura:
            return jsonify({'success': False, 'message': f'El monto excede el total. Total: S/ {total_factura:.2f}'}), 400
        
        # 3. Determinar nuevo estado
        if es_parcial and monto_pagado < total_factura:
            nuevo_estado = 'credito'
            saldo_pendiente = total_factura - monto_pagado
            mensaje_estado = f'Pago parcial de S/ {monto_pagado:.2f} registrado. Pendiente: S/ {saldo_pendiente:.2f}'
            
            # Verificar si existe columna saldo_pendiente
            cursor.execute("SHOW COLUMNS FROM facturas LIKE 'saldo_pendiente'")
            tiene_saldo_pendiente = cursor.fetchone()
            
            if tiene_saldo_pendiente:
                cursor.execute("""
                    UPDATE facturas 
                    SET estado = 'credito', 
                        metodo_pago = %s, 
                        fecha_pago = NOW(),
                        saldo_pendiente = %s
                    WHERE id_factura = %s
                """, (metodo_pago, saldo_pendiente, id))
            else:
                cursor.execute("""
                    UPDATE facturas 
                    SET estado = 'credito', 
                        metodo_pago = %s, 
                        fecha_pago = NOW()
                    WHERE id_factura = %s
                """, (metodo_pago, id))
        else:
            # Pago completo
            nuevo_estado = 'pagada'
            monto_pagado = total_factura
            mensaje_estado = f'Pago completo de S/ {monto_pagado:.2f} registrado.'
            
            cursor.execute("""
                UPDATE facturas 
                SET estado = 'pagada', 
                    metodo_pago = %s, 
                    fecha_pago = NOW()
                WHERE id_factura = %s
            """, (metodo_pago, id))
        
        print(f"🔄 Actualizando factura {id} a estado: {nuevo_estado}")
        
        # 4. Registrar movimiento en caja
        id_empleado = session.get('id_empleado', 1)
        cursor.execute("""
            SELECT id_caja 
            FROM caja_diaria 
            WHERE fecha = CURRENT_DATE 
            AND estado = 'abierta'
            AND id_empleado_cajero = %s
            LIMIT 1
        """, (id_empleado,))
        
        caja_abierta = cursor.fetchone()
        
        if caja_abierta:
            id_caja = caja_abierta['id_caja']
            print(f"📦 Caja abierta encontrada: ID {id_caja}")
            
            # Registrar movimiento en caja
            try:
                cursor.execute("""
                    INSERT INTO movimientos_caja 
                    (id_caja, id_factura, tipo, metodo_pago, concepto, monto, fecha_movimiento, id_empleado)
                    VALUES (%s, %s, 'ingreso', %s, %s, %s, NOW(), %s)
                """, (id_caja, id, metodo_pago, 
                      f'Pago de {factura["tipo_comprobante"]} {factura["numero"]}', 
                      monto_pagado, id_empleado))
                
                print("✅ Movimiento en caja registrado")
                
                # Actualizar totales en caja
                campo_venta = ''
                if metodo_pago == 'efectivo':
                    campo_venta = 'venta_efectivo'
                elif metodo_pago == 'tarjeta':
                    campo_venta = 'venta_tarjeta'
                else:
                    campo_venta = 'venta_digital'
                
                cursor.execute(f"""
                    UPDATE caja_diaria 
                    SET {campo_venta} = COALESCE({campo_venta}, 0) + %s,
                        total_ventas = COALESCE(total_ventas, 0) + %s
                    WHERE id_caja = %s
                """, (monto_pagado, monto_pagado, id_caja))
                
            except Exception as e:
                print(f"⚠️  Error registrando en caja: {e}")
        else:
            print("⚠️  No hay caja abierta")
            mensaje_estado += " (Sin registro en caja)"
        
        conn.commit()
        
        print(f"✅ Pago registrado para factura {id}")
        
        return jsonify({
            'success': True, 
            'message': mensaje_estado,
            'nuevo_estado': nuevo_estado,
            'monto_pagado': monto_pagado,
            'caja_registrada': caja_abierta is not None
        })
            
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"❌ Error en pagar_factura: {e}")
        import traceback
        print(f"🔍 Traceback: {traceback.format_exc()}")
        return jsonify({'success': False, 'message': f'Error registrando pago: {str(e)}'}), 500
    finally:
        if 'cursor' in locals() and cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/facturas/<int:id>/anular', methods=['POST'])
def anular_factura(id):
    """Anular factura"""
    print(f"🔍 DEBUG: Anulando factura {id}")
    
    conn = get_db_connection()
    
    if not conn:
        return jsonify({'success': False, 'message': 'No hay conexión a la base de datos.'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar que la factura existe
        cursor.execute("SELECT estado, numero FROM facturas WHERE id_factura = %s", (id,))
        factura = cursor.fetchone()
        
        if not factura:
            return jsonify({'success': False, 'message': 'Factura no encontrada.'}), 404
        
        if factura['estado'] == 'anulada':
            return jsonify({'success': False, 'message': 'La factura ya está anulada.'})
        
        # No permitir anular facturas ya pagadas sin confirmación especial
        if factura['estado'] == 'pagada':
            return jsonify({
                'success': False, 
                'message': 'No se puede anular una factura pagada automáticamente. Contacte al administrador.'
            })
        
        # Anular factura
        cursor.execute("UPDATE facturas SET estado = 'anulada' WHERE id_factura = %s", (id,))
        conn.commit()
        
        print(f"✅ Factura {factura['numero']} anulada exitosamente")
        
        return jsonify({'success': True, 'message': 'Factura anulada exitosamente.'})
            
    except Error as e:
        if conn:
            conn.rollback()
        print(f"❌ Error anulando factura: {e}")
        return jsonify({'success': False, 'message': f'Error anulando factura: {str(e)}'}), 500
    finally:
        if 'cursor' in locals() and cursor:
            cursor.close()
        if conn:
            conn.close()        


# Añadir estas rutas a tu app.py

@app.route('/empleados')
@login_required
def empleados():
    """Página principal de gestión de empleados"""
    return render_template('empleados/listar.html')

# ==================== API PARA EMPLEADOS ====================

@app.route('/api/empleados', methods=['GET'])
def api_get_empleados():
    """Obtener todos los empleados"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id_empleado, dni, nombre, apellido, telefono, email, 
                   especialidad, fecha_contratacion, activo
            FROM empleados
            ORDER BY nombre, apellido
        """)
        
        # Convertir resultados a lista de diccionarios
        empleados_raw = cursor.fetchall()
        empleados = []
        
        for empleado in empleados_raw:
            empleado_dict = dict(empleado)
            
            # Formatear fechas para JSON
            if empleado_dict['fecha_contratacion']:
                empleado_dict['fecha_contratacion'] = empleado_dict['fecha_contratacion'].strftime('%Y-%m-%d')
            
            # Asegurar que los campos sean strings
            empleado_dict['telefono'] = str(empleado_dict['telefono']) if empleado_dict['telefono'] else ''
            empleado_dict['email'] = str(empleado_dict['email']) if empleado_dict['email'] else ''
            empleado_dict['especialidad'] = str(empleado_dict['especialidad']) if empleado_dict['especialidad'] else ''
            
            empleados.append(empleado_dict)
        
        return jsonify(empleados)
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados', methods=['POST'])
def api_create_empleado():
    """Crear nuevo empleado"""
    data = request.get_json()
    
    # Validar datos requeridos
    required_fields = ['dni', 'nombre', 'apellido', 'email', 'especialidad']
    for field in required_fields:
        if field not in data or not data[field]:
            return jsonify({'error': f'El campo {field} es requerido'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el DNI ya existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE dni = %s", (data['dni'],))
        if cursor.fetchone():
            return jsonify({'error': 'El DNI ya está registrado'}), 400
        
        # Verificar si el email ya existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE email = %s", (data['email'],))
        if cursor.fetchone():
            return jsonify({'error': 'El email ya está registrado'}), 400
        
        # Insertar empleado
        cursor.execute("""
            INSERT INTO empleados (dni, nombre, apellido, telefono, email, 
                                  especialidad, fecha_contratacion, activo)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            data['dni'],
            data['nombre'],
            data['apellido'],
            data.get('telefono'),
            data['email'],
            data['especialidad'],
            data.get('fecha_contratacion'),
            data.get('activo', True)
        ))
        
        id_empleado = cursor.lastrowid
        
        # Crear usuario si se solicitó
        if 'usuario' in data and data['usuario']:
            usuario_data = data['usuario']
            
            # Validar datos del usuario
            if not usuario_data.get('username') or not usuario_data.get('password') or not usuario_data.get('rol'):
                return jsonify({'error': 'Faltan datos para crear el usuario'}), 400
            
            # Crear hash de la contraseña
            password_hash = hashlib.sha256(usuario_data['password'].encode()).hexdigest()
            
            cursor.execute("""
                INSERT INTO usuarios (id_empleado, username, password_hash, rol)
                VALUES (%s, %s, %s, %s)
            """, (
                id_empleado,
                usuario_data['username'],
                password_hash,
                usuario_data['rol']
            ))
        
        conn.commit()
        
        # Obtener el empleado creado
        cursor.execute("""
            SELECT id_empleado, dni, nombre, apellido, telefono, email, 
                   especialidad, fecha_contratacion, activo
            FROM empleados
            WHERE id_empleado = %s
        """, (id_empleado,))
        empleado_raw = cursor.fetchone()
        
        # Convertir a diccionario
        empleado = dict(empleado_raw)
        
        # Formatear fecha
        if empleado['fecha_contratacion']:
            empleado['fecha_contratacion'] = empleado['fecha_contratacion'].strftime('%Y-%m-%d')
        
        return jsonify({'success': True, 'empleado': empleado})
        
    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados/<int:id>', methods=['GET'])
def api_get_empleado(id):
    """Obtener un empleado específico"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id_empleado, dni, nombre, apellido, telefono, email, 
                   especialidad, fecha_contratacion, activo
            FROM empleados
            WHERE id_empleado = %s
        """, (id,))
        empleado_raw = cursor.fetchone()
        
        if not empleado_raw:
            return jsonify({'error': 'Empleado no encontrado'}), 404
        
        # Convertir a diccionario
        empleado = dict(empleado_raw)
        
        # Formatear fechas
        if empleado['fecha_contratacion']:
            empleado['fecha_contratacion'] = empleado['fecha_contratacion'].strftime('%Y-%m-%d')
        
        # Asegurar tipos de datos
        empleado['telefono'] = str(empleado['telefono']) if empleado['telefono'] else ''
        empleado['email'] = str(empleado['email']) if empleado['email'] else ''
        empleado['especialidad'] = str(empleado['especialidad']) if empleado['especialidad'] else ''
        empleado['activo'] = bool(empleado['activo'])
        
        return jsonify(empleado)
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados/<int:id>', methods=['PUT'])
def api_update_empleado(id):
    """Actualizar empleado"""
    data = request.get_json()
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el empleado existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE id_empleado = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Empleado no encontrado'}), 404
        
        # Verificar duplicados de DNI (excluyendo el actual)
        if 'dni' in data:
            cursor.execute("SELECT id_empleado FROM empleados WHERE dni = %s AND id_empleado != %s", 
                          (data['dni'], id))
            if cursor.fetchone():
                return jsonify({'error': 'El DNI ya está registrado por otro empleado'}), 400
        
        # Verificar duplicados de email (excluyendo el actual)
        if 'email' in data:
            cursor.execute("SELECT id_empleado FROM empleados WHERE email = %s AND id_empleado != %s", 
                          (data['email'], id))
            if cursor.fetchone():
                return jsonify({'error': 'El email ya está registrado por otro empleado'}), 400
        
        # Construir query de actualización dinámica
        update_fields = []
        params = []
        
        field_mapping = {
            'dni': 'dni',
            'nombre': 'nombre',
            'apellido': 'apellido',
            'telefono': 'telefono',
            'email': 'email',
            'especialidad': 'especialidad',
            'fecha_contratacion': 'fecha_contratacion',
            'activo': 'activo'
        }
        
        for key, db_field in field_mapping.items():
            if key in data:
                update_fields.append(f"{db_field} = %s")
                params.append(data[key])
        
        if not update_fields:
            return jsonify({'error': 'No hay datos para actualizar'}), 400
        
        params.append(id)  # Para el WHERE
        
        query = f"""
            UPDATE empleados 
            SET {', '.join(update_fields)}
            WHERE id_empleado = %s
        """
        
        cursor.execute(query, params)
        conn.commit()
        
        return jsonify({'success': True})
        
    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados/<int:id>', methods=['DELETE'])
def api_delete_empleado(id):
    """Desactivar empleado (no eliminar permanentemente)"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el empleado existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE id_empleado = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Empleado no encontrado'}), 404
        
        # Desactivar empleado (no eliminar)
        cursor.execute("UPDATE empleados SET activo = FALSE WHERE id_empleado = %s", (id,))
        
        # También desactivar el usuario si existe
        cursor.execute("UPDATE usuarios SET activo = FALSE WHERE id_empleado = %s", (id,))
        
        conn.commit()
        return jsonify({'success': True})
        
    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados/<int:id>/usuario', methods=['GET'])
def api_get_usuario_empleado(id):
    """Obtener información del usuario de un empleado"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el empleado existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE id_empleado = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Empleado no encontrado'}), 404
        
        # Obtener usuario
        cursor.execute("""
            SELECT id_usuario, username, rol, fecha_creacion, ultimo_login, activo
            FROM usuarios
            WHERE id_empleado = %s
        """, (id,))
        usuario_raw = cursor.fetchone()
        
        if usuario_raw:
            usuario = dict(usuario_raw)
            
            # Formatear fechas
            if usuario['fecha_creacion']:
                usuario['fecha_creacion'] = usuario['fecha_creacion'].strftime('%d/%m/%Y %H:%M')
            if usuario['ultimo_login']:
                usuario['ultimo_login'] = usuario['ultimo_login'].strftime('%d/%m/%Y %H:%M')
            
            return jsonify({'existe': True, 'usuario': usuario})
        else:
            return jsonify({'existe': False})
        
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/empleados/<int:id>/usuario', methods=['POST'])
def api_create_usuario_empleado(id):
    """Crear usuario para un empleado"""
    data = request.get_json()
    
    # Validar datos
    required_fields = ['username', 'password', 'rol']
    for field in required_fields:
        if field not in data or not data[field]:
            return jsonify({'error': f'El campo {field} es requerido'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el empleado existe
        cursor.execute("SELECT id_empleado FROM empleados WHERE id_empleado = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Empleado no encontrado'}), 404
        
        # Verificar si ya tiene usuario
        cursor.execute("SELECT id_usuario FROM usuarios WHERE id_empleado = %s", (id,))
        if cursor.fetchone():
            return jsonify({'error': 'El empleado ya tiene un usuario'}), 400
        
        # Verificar si el username ya existe
        cursor.execute("SELECT id_usuario FROM usuarios WHERE username = %s", (data['username'],))
        if cursor.fetchone():
            return jsonify({'error': 'El username ya está en uso'}), 400
        
        # Crear hash de la contraseña
        password_hash = hashlib.sha256(data['password'].encode()).hexdigest()
        
        # Crear usuario
        cursor.execute("""
            INSERT INTO usuarios (id_empleado, username, password_hash, rol)
            VALUES (%s, %s, %s, %s)
        """, (id, data['username'], password_hash, data['rol']))
        
        conn.commit()
        return jsonify({'success': True})
        
    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/usuarios/<int:id>', methods=['PUT'])
@login_required
def api_update_usuario(id):
    """Actualizar usuario"""
    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No hay datos para actualizar'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el usuario existe
        cursor.execute("SELECT id_usuario FROM usuarios WHERE id_usuario = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Usuario no encontrado'}), 404
        
        # Construir query de actualización dinámica
        update_fields = []
        params = []
        
        if 'rol' in data:
            update_fields.append("rol = %s")
            params.append(data['rol'])
        
        if 'password' in data and data['password']:
            password_hash = hashlib.sha256(data['password'].encode()).hexdigest()
            update_fields.append("password_hash = %s")
            params.append(password_hash)
        
        if 'activo' in data:
            update_fields.append("activo = %s")
            params.append(data['activo'])
        
        if not update_fields:
            return jsonify({'error': 'No hay datos válidos para actualizar'}), 400
        
        params.append(id)  # Para el WHERE
        
        query = f"""
            UPDATE usuarios 
            SET {', '.join(update_fields)}
            WHERE id_usuario = %s
        """
        
        cursor.execute(query, params)
        conn.commit()
        
        return jsonify({'success': True})
        
    except Error as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# ==================== FUNCIONES DE REPORTES ====================

@app.route('/api/empleados/estadisticas')
def api_get_estadisticas_empleados():
    """Obtener estadísticas de empleados"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Estadísticas básicas
        cursor.execute("""
            SELECT 
                COUNT(*) as total,
                SUM(CASE WHEN activo = TRUE THEN 1 ELSE 0 END) as activos,
                SUM(CASE WHEN activo = FALSE THEN 1 ELSE 0 END) as inactivos,
                COUNT(DISTINCT especialidad) as especialidades_diferentes
            FROM empleados
        """)
        estadisticas = cursor.fetchone()
        
        # Conteo por especialidad
        cursor.execute("""
            SELECT especialidad, COUNT(*) as cantidad
            FROM empleados
            WHERE activo = TRUE
            GROUP BY especialidad
            ORDER BY cantidad DESC
        """)
        por_especialidad = cursor.fetchall()
        
        # Últimos empleados contratados
        cursor.execute("""
            SELECT nombre, apellido, fecha_contratacion, especialidad
            FROM empleados
            WHERE fecha_contratacion IS NOT NULL
            ORDER BY fecha_contratacion DESC
            LIMIT 5
        """)
        ultimos_contratados = cursor.fetchall()
        
        # Formatear fechas
        for empleado in ultimos_contratados:
            if empleado['fecha_contratacion']:
                empleado['fecha_contratacion'] = empleado['fecha_contratacion'].strftime('%d/%m/%Y')
        
        return jsonify({
            'estadisticas': estadisticas,
            'por_especialidad': por_especialidad,
            'ultimos_contratados': ultimos_contratados
        })
        
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/empleados/exportar')
def exportar_empleados():
    """Exportar empleados a Excel"""
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('empleados'))
    
    try:
        cursor = conn.cursor()
        cursor.execute("""
            SELECT dni, nombre, apellido, telefono, email, especialidad,
                   fecha_contratacion, 
                   CASE WHEN activo = TRUE THEN 'Activo' ELSE 'Inactivo' END as estado
            FROM empleados
            ORDER BY nombre, apellido
        """)
        empleados = cursor.fetchall()
        
        # Crear DataFrame
        import pandas as pd
        df = pd.DataFrame(empleados)
        
        # Crear respuesta Excel
        from io import BytesIO
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Empleados', index=False)
            
            # Ajustar ancho de columnas
            worksheet = writer.sheets['Empleados']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        # Enviar archivo
        from flask import send_file
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'empleados_petglow_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
        
    except Error as e:
        flash(f'Error exportando empleados: {str(e)}', 'danger')
        return redirect(url_for('empleados'))
    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('empleados'))
    finally:
        cursor.close()
        conn.close()
@app.route('/api/empleado/<int:id>/disponibilidad')
def verificar_disponibilidad_empleado(id):
    """Verificar disponibilidad de un empleado en fecha y hora específicas"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'disponible': False, 'mensaje': 'Error de conexión'}), 500
    
    try:
        cursor = conn.cursor()
        
        fecha = request.args.get('fecha')
        hora = request.args.get('hora')
        
        if not fecha or not hora:
            return jsonify({'disponible': False, 'mensaje': 'Fecha y hora requeridas'}), 400
        
        # Obtener información del empleado
        cursor.execute("SELECT nombre, apellido FROM empleados WHERE id_empleado = %s", (id,))
        empleado = cursor.fetchone()
        
        if not empleado:
            return jsonify({'disponible': False, 'mensaje': 'Empleado no encontrado'}), 404
        
        # Verificar si es administrador
        nombre_completo = f"{empleado['nombre']} {empleado['apellido']}".lower()
        es_administrador = any(keyword in nombre_completo for keyword in ['admin', 'sistema', 'administrador'])
        
        # Si es administrador, siempre está disponible
        if es_administrador:
            return jsonify({
                'disponible': True,
                'mensaje': 'Administrador - puede múltiples reservas'
            })
        
        # Si no es administrador, verificar disponibilidad normal
        fecha_hora_str = f"{fecha} {hora}"
        try:
            fecha_hora_dt = datetime.strptime(fecha_hora_str, '%Y-%m-%d %H:%M')
        except ValueError:
            return jsonify({'disponible': False, 'mensaje': 'Formato de fecha/hora inválido'}), 400
        
        # Verificar reservas existentes
        cursor.execute("""
            SELECT r.id_reserva, r.fecha_reserva, s.duracion_min
            FROM reservas r
            JOIN reserva_servicios rs ON r.id_reserva = rs.id_reserva
            JOIN servicios s ON rs.id_servicio = s.id_servicio
            WHERE r.id_empleado = %s 
            AND r.estado NOT IN ('cancelada', 'no_show')
            AND DATE(r.fecha_reserva) = DATE(%s)
        """, (id, fecha))
        
        reservas_existentes = cursor.fetchall()
        
        # Para simplificar, asumimos duración de 60 minutos
        duracion_estimada = 60
        nueva_inicio = fecha_hora_dt
        nueva_fin = fecha_hora_dt + timedelta(minutes=duracion_estimada)
        
        for reserva in reservas_existentes:
            reserva_inicio = reserva['fecha_reserva']
            reserva_duracion = int(reserva['duracion_min']) if reserva['duracion_min'] else 60
            reserva_fin = reserva_inicio + timedelta(minutes=int(reserva_duracion))
            
            if (nueva_inicio < reserva_fin and nueva_fin > reserva_inicio):
                return jsonify({
                    'disponible': False,
                    'mensaje': f'Ya tiene reserva de {reserva_inicio.strftime("%H:%M")} a {reserva_fin.strftime("%H:%M")}'
                })
        
        return jsonify({'disponible': True, 'mensaje': 'Empleado disponible'})
        
    except Error as e:
        return jsonify({'disponible': False, 'mensaje': f'Error: {str(e)}'}), 500
    finally:
        cursor.close()
        conn.close()
# ==================== RUTAS PARA USUARIOS ====================

@app.route('/usuarios')
@login_required  # Agregar este decorador
@admin_required  # Solo administradores pueden ver usuarios
def usuarios():
    """Listar usuarios del sistema"""
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('index'))
    
    try:
        cursor = conn.cursor()
        
        # Obtener todos los usuarios
        cursor.execute("""
            SELECT u.*, 
                   e.nombre as empleado_nombre, 
                   e.apellido as empleado_apellido, 
                   e.email as empleado_email,
                   e.dni as empleado_dni
            FROM usuarios u
            LEFT JOIN empleados e ON u.id_empleado = e.id_empleado
            ORDER BY u.id_usuario DESC
        """)
        
        usuarios_raw = cursor.fetchall()
        
        # Convertir a diccionarios y formatear
        usuarios = []
        for usuario in usuarios_raw:
            user_dict = dict(usuario)
            
            # Convertir booleano
            user_dict['activo'] = bool(user_dict['activo'])
            
            # Formatear fechas si existen
            if user_dict['ultimo_login']:
                # Mantener como datetime para que la plantilla pueda usar .strftime()
                pass  # Se queda como datetime
            else:
                user_dict['ultimo_login'] = None
            
            if user_dict['fecha_creacion']:
                pass  # Se queda como datetime
            
            usuarios.append(user_dict)
        
        # Estadísticas
        total_usuarios = len(usuarios)
        usuarios_activos = sum(1 for u in usuarios if u['activo'])
        administradores = sum(1 for u in usuarios if u['rol'] == 'admin')
        
        # Obtener empleados sin usuario
        cursor.execute("""
            SELECT e.* FROM empleados e
            LEFT JOIN usuarios u ON e.id_empleado = u.id_empleado
            WHERE u.id_usuario IS NULL AND e.activo = TRUE
            ORDER BY e.nombre
        """)
        empleados_sin_usuario_raw = cursor.fetchall()
        
        empleados_sin_usuario = []
        for emp in empleados_sin_usuario_raw:
            empleados_sin_usuario.append(dict(emp))
        
        # Último registro
        ultimo_registro = None
        if usuarios and usuarios[0].get('fecha_creacion'):
            ultimo_registro = usuarios[0]['fecha_creacion'].strftime('%d/%m/%Y')
        
    except Error as e:
        flash(f'Error cargando usuarios: {e}', 'danger')
        return redirect(url_for('index'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('usuarios/listar.html',
                         usuarios=usuarios,
                         total_usuarios=total_usuarios,
                         usuarios_activos=usuarios_activos,
                         administradores=administradores,
                         ultimo_registro=ultimo_registro,
                         empleados_sin_usuario=empleados_sin_usuario)

# ==================== API PARA USUARIOS ====================

@app.route('/api/usuarios', methods=['POST'])
def api_crear_usuario():
    """Crear un nuevo usuario (API)"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'error': 'No autorizado'}), 401
    
    data = request.json
    required_fields = ['username', 'password', 'rol']
    
    # Validar campos requeridos
    for field in required_fields:
        if field not in data or not data[field]:
            return jsonify({'success': False, 'error': f'El campo {field} es requerido'}), 400
    
    # Validar longitud de contraseña
    if len(data['password']) < 6:
        return jsonify({'success': False, 'error': 'La contraseña debe tener al menos 6 caracteres'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el usuario ya existe
        cursor.execute("SELECT id_usuario FROM usuarios WHERE username = %s", (data['username'],))
        if cursor.fetchone():
            return jsonify({'success': False, 'error': 'El nombre de usuario ya existe'}), 400
        
        # Crear usuario
        password_hash = hash_password(data['password'])
        cursor.execute("""
            INSERT INTO usuarios (username, password_hash, id_empleado, rol, activo)
            VALUES (%s, %s, %s, %s, %s)
        """, (data['username'], password_hash, 
            data.get('id_empleado'), data['rol'], 
            data.get('activo', True)))

        conn.commit()
        
        return jsonify({'success': True, 'message': 'Usuario creado exitosamente'})
        
    except Error as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/usuarios/<int:id>', methods=['GET'])
def api_obtener_usuario(id):
    """Obtener información de un usuario (API)"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'error': 'No autorizado'}), 401
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Obtener usuario - USAR COALESCE para manejar NULLs
        cursor.execute("""
            SELECT 
                u.id_usuario,
                u.username,
                u.id_empleado,
                u.rol,
                u.activo,
                u.fecha_creacion,
                u.ultimo_login,
                COALESCE(e.nombre, '') as nombre,
                COALESCE(e.apellido, '') as apellido,
                COALESCE(e.dni, '') as dni,
                COALESCE(e.email, '') as email
            FROM usuarios u
            LEFT JOIN empleados e ON u.id_empleado = e.id_empleado
            WHERE u.id_usuario = %s
        """, (id,))
        
        usuario_raw = cursor.fetchone()
        
        if not usuario_raw:
            return jsonify({'success': False, 'error': 'Usuario no encontrado'}), 404
        
        # Convertir a diccionario
        usuario = dict(usuario_raw)
        
        # Remover password_hash por seguridad
        if 'password_hash' in usuario:
            del usuario['password_hash']
        
        # Manejar valores booleanos de PostgreSQL
        if 'activo' in usuario:
            usuario['activo'] = bool(usuario['activo'])
        
        # Formatear fechas para JSON
        for field in ['fecha_creacion', 'ultimo_login']:
            if field in usuario and usuario[field]:
                usuario[field] = usuario[field].isoformat()
            else:
                usuario[field] = None
        
        # Intentar obtener historial de login
        try:
            cursor.execute("""
                SELECT fecha_login, ip_address, user_agent
                FROM login_history
                WHERE id_usuario = %s
                ORDER BY fecha_login DESC
                LIMIT 5
            """, (id,))
            
            historial_raw = cursor.fetchall()
            # Convertir historial a lista de diccionarios y formatear fechas
            historial = []
            for h in historial_raw:
                hist_item = dict(h)
                if hist_item.get('fecha_login'):
                    hist_item['fecha_login'] = hist_item['fecha_login'].isoformat()
                historial.append(hist_item)
            
            usuario['historial_login'] = historial
        except Exception as e:
            print(f"Advertencia: No se pudo obtener historial de login: {e}")
            usuario['historial_login'] = []
        
        return jsonify({'success': True, 'usuario': usuario})
        
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()
        
@app.route('/api/usuarios/<int:id>', methods=['PUT'])
def api_actualizar_usuario(id):
    """Actualizar un usuario (API)"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'error': 'No autorizado'}), 401
    
    data = request.json
    
    # Validar campos requeridos
    if 'username' not in data or not data['username']:
        return jsonify({'success': False, 'error': 'El nombre de usuario es requerido'}), 400
    
    if 'rol' not in data or not data['rol']:
        return jsonify({'success': False, 'error': 'El rol es requerido'}), 400
    
    # Validar contraseña si se proporciona
    if 'password' in data and data['password']:
        if len(data['password']) < 6:
            return jsonify({'success': False, 'error': 'La contraseña debe tener al menos 6 caracteres'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el usuario existe
        cursor.execute("SELECT id_usuario FROM usuarios WHERE id_usuario = %s", (id,))
        if not cursor.fetchone():
            return jsonify({'success': False, 'error': 'Usuario no encontrado'}), 404
        
        # Verificar si el username ya existe (excluyendo el actual)
        cursor.execute("SELECT id_usuario FROM usuarios WHERE username = %s AND id_usuario != %s", 
                      (data['username'], id))
        if cursor.fetchone():
            return jsonify({'success': False, 'error': 'El nombre de usuario ya existe'}), 400
        
        # Actualizar usuario
        if 'password' in data and data['password']:
            password_hash = hash_password(data['password'])
            cursor.execute("""
                UPDATE usuarios 
                SET username = %s, password_hash = %s, rol = %s, activo = %s
                WHERE id_usuario = %s
            """, (data['username'], password_hash, data['rol'], 
                data.get('activo', True), id))
        else:
            cursor.execute("""
                UPDATE usuarios 
                SET username = %s, rol = %s, activo = %s
                WHERE id_usuario = %s
            """, (data['username'], data['rol'], 
                data.get('activo', True), id))
        
        conn.commit()
        
        return jsonify({'success': True, 'message': 'Usuario actualizado exitosamente'})
        
    except Error as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/usuarios/<int:id>', methods=['DELETE'])
def api_eliminar_usuario(id):
    """Eliminar un usuario (API)"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'error': 'No autorizado'}), 401
    
    # Prevenir que un usuario se elimine a sí mismo
    if id == session.get('id_usuario'):
        return jsonify({'success': False, 'error': 'No puedes eliminar tu propio usuario'}), 400
    
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'Error de conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Verificar si el usuario existe
        cursor.execute("SELECT id_usuario, rol FROM usuarios WHERE id_usuario = %s", (id,))
        usuario = cursor.fetchone()
        
        if not usuario:
            return jsonify({'success': False, 'error': 'Usuario no encontrado'}), 404
        
        # Prevenir eliminación del último administrador
        if usuario['rol'] == 'admin':
            cursor.execute("SELECT COUNT(*) as total_admins FROM usuarios WHERE rol = 'admin'")
            total_admins = cursor.fetchone()['total_admins']
            
            if total_admins <= 1:
                return jsonify({'success': False, 'error': 'No se puede eliminar el único administrador'}), 400
        
        # Eliminar usuario
        cursor.execute("DELETE FROM usuarios WHERE id_usuario = %s", (id,))
        
        conn.commit()
        
        return jsonify({'success': True, 'message': 'Usuario eliminado exitosamente'})
        
    except Error as e:
        conn.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# ==================== CONFIGURACIÓN SIMPLE DE CONTRASEÑA ====================

def obtener_contraseña_reportes():
    """Obtener la contraseña desde un archivo simple"""
    try:
        if os.path.exists('clave_reportes.txt'):
            with open('clave_reportes.txt', 'r', encoding='utf-8') as f:
                return f.read().strip()
        else:
            # Contraseña por defecto
            contraseña_default = "medina123"
            with open('clave_reportes.txt', 'w', encoding='utf-8') as f:
                f.write(contraseña_default)
            return contraseña_default
    except Exception as e:
        print(f"Error al obtener contraseña: {e}")
        return "medina123"  # Fallback

def cambiar_contraseña_reportes(nueva_contraseña):
    """Cambiar la contraseña"""
    try:
        with open('clave_reportes.txt', 'w', encoding='utf-8') as f:
            f.write(nueva_contraseña.strip())
        return True
    except Exception as e:
        print(f"Error al cambiar contraseña: {e}")
        return False

@app.route('/verificar_contraseña_reportes', methods=['POST'])
def verificar_contraseña_reportes():
    """Verificar contraseña desde el modal (AJAX)"""
    if 'id_usuario' not in session:
        return jsonify({'success': False, 'message': 'No autorizado'})
    
    data = request.get_json()
    password_ingresada = data.get('password', '')
    
    contraseña_correcta = obtener_contraseña_reportes()
    
    if password_ingresada == contraseña_correcta:
        session['reportes_acceso'] = True
        return jsonify({'success': True, 'message': 'Contraseña correcta'})
    
    return jsonify({'success': False, 'message': 'Contraseña incorrecta'})

@app.route('/cambiar_contraseña_reportes', methods=['POST'])
def cambiar_contraseña_reportes_route():
    """Cambiar la contraseña desde la configuración"""
    if 'id_usuario' not in session or session.get('rol') != 'admin':
        flash('No autorizado', 'danger')
        return redirect(url_for('config'))
    
    contraseña_actual = request.form.get('password_actual', '').strip()
    nueva_contraseña = request.form.get('nueva_password', '').strip()
    confirmar_contraseña = request.form.get('confirmar_password', '').strip()
    
    # Verificar que todos los campos estén completos
    if not contraseña_actual or not nueva_contraseña or not confirmar_contraseña:
        flash('Todos los campos son obligatorios', 'danger')
        return redirect(url_for('config'))
    
    # Verificar contraseña actual
    contraseña_guardada = obtener_contraseña_reportes()
    
    if contraseña_actual != contraseña_guardada:
        flash('La contraseña actual es incorrecta', 'danger')
        return redirect(url_for('config'))
    
    if nueva_contraseña != confirmar_contraseña:
        flash('Las nuevas contraseñas no coinciden', 'danger')
        return redirect(url_for('config'))
    
    if len(nueva_contraseña) < 4:
        flash('La contraseña debe tener al menos 4 caracteres', 'danger')
        return redirect(url_for('config'))
    
    # Cambiar contraseña
    if cambiar_contraseña_reportes(nueva_contraseña):
        flash('Contraseña actualizada correctamente', 'success')
    else:
        flash('Error al guardar la contraseña', 'danger')
    
    return redirect(url_for('config'))

# ==================== RUTAS DE REPORTES CON VERIFICACIÓN SIMPLE ====================

def verificar_acceso_reportes(f):
    """Decorador simple para verificar acceso a reportes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 1. Verificar que sea admin
        if session.get('rol') != 'admin':
            flash('Solo los administradores pueden acceder a reportes', 'danger')
            return redirect(url_for('dashboard'))
        
        # 2. Verificar si ya ingresó la contraseña en esta sesión
        if not session.get('reportes_acceso'):
            flash('Debe ingresar la contraseña para acceder a reportes', 'warning')
            return redirect(url_for('dashboard'))
        
        return f(*args, **kwargs)
    return decorated_function


# ==================== RUTAS DE REPORTES ====================

@app.route('/reportes')
@login_required
@verificar_acceso_reportes
def reportes():
    """Página principal de reportes"""
    return render_template('reportes/generar.html')

@app.route('/reportes/ventas')
@login_required
@verificar_acceso_reportes
def reporte_ventas():
    """Reporte de ventas"""
    fecha_inicio = request.args.get('fecha_inicio', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_fin = request.args.get('fecha_fin', datetime.now().strftime('%Y-%m-%d'))
    
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reportes'))
    
    try:
        cursor = conn.cursor()
        
        # 1. Estadísticas generales del período
        cursor.execute("""
            SELECT 
                COUNT(*) as total_facturas,
                SUM(CASE WHEN estado = 'pagada' THEN total ELSE 0 END) as total_ingresos,
                SUM(CASE WHEN estado = 'pendiente' THEN total ELSE 0 END) as total_pendiente,
                SUM(CASE WHEN estado = 'credito' THEN saldo_pendiente ELSE 0 END) as total_credito,
                AVG(total) as promedio_venta,
                MAX(total) as venta_maxima,
                MIN(total) as venta_minima
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
        """, (fecha_inicio, fecha_fin))
        
        estadisticas = cursor.fetchone()
        
        # 2. Ventas por día
        cursor.execute("""
            SELECT 
                DATE(fecha_emision) as fecha,
                COUNT(*) as cantidad_facturas,
                SUM(total) as total_dia,
                AVG(total) as promedio_dia
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
            GROUP BY DATE(fecha_emision)
            ORDER BY fecha
        """, (fecha_inicio, fecha_fin))
        
        ventas_por_dia = cursor.fetchall()
        
        # 3. Ventas por método de pago
        cursor.execute("""
            SELECT 
                COALESCE(metodo_pago, 'no especificado') as metodo_pago,
                COUNT(*) as cantidad,
                SUM(total) as total_metodo
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
            GROUP BY metodo_pago
            ORDER BY total_metodo DESC
        """, (fecha_inicio, fecha_fin))
        
        ventas_por_metodo = cursor.fetchall()
        
        # 4. Top 10 clientes
        cursor.execute("""
            SELECT 
                c.id_cliente,
                c.nombre,
                c.apellido,
                COUNT(f.id_factura) as cantidad_facturas,
                SUM(f.total) as total_gastado
            FROM facturas f
            JOIN clientes c ON f.id_cliente = c.id_cliente
            WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                AND f.estado = 'pagada'
            GROUP BY c.id_cliente, c.nombre, c.apellido
            ORDER BY total_gastado DESC
            LIMIT 10
        """, (fecha_inicio, fecha_fin))
        
        top_clientes = cursor.fetchall()
        
        # 5. Servicios más vendidos
        cursor.execute("""
            SELECT 
                s.id_servicio,
                s.nombre,
                s.categoria,
                COUNT(fs.id_detalle) as veces_vendido,
                SUM(fs.cantidad) as cantidad_total,
                SUM(fs.subtotal) as ingresos_servicio,
                AVG(fs.precio_unitario) as precio_promedio
            FROM factura_servicios fs
            JOIN servicios s ON fs.id_servicio = s.id_servicio
            JOIN facturas f ON fs.id_factura = f.id_factura
            WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                AND f.estado = 'pagada'
            GROUP BY s.id_servicio, s.nombre, s.categoria
            ORDER BY ingresos_servicio DESC
            LIMIT 15
        """, (fecha_inicio, fecha_fin))
        
        top_servicios = cursor.fetchall()
        
        # Calcular totales
        total_ingresos = estadisticas['total_ingresos'] or 0
        total_facturas = estadisticas['total_facturas'] or 0
        
        # Preparar datos para gráficos
        dias = []
        ventas_diarias = []
        
        for v in ventas_por_dia:
            if v['fecha']:
                dias.append(v['fecha'].strftime('%d/%m'))
                ventas_diarias.append(float(v['total_dia'] or 0))
        
        # Métodos de pago
        metodos = []
        totales_metodos = []
        
        for v in ventas_por_metodo:
            metodos.append(v['metodo_pago'])
            totales_metodos.append(float(v['total_metodo'] or 0))
        
    except Error as e:
        flash(f'Error generando reporte: {e}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reportes/ventas.html',
                         fecha_inicio=fecha_inicio,
                         fecha_fin=fecha_fin,
                         estadisticas=estadisticas,
                         top_clientes=top_clientes,
                         top_servicios=top_servicios,
                         total_ingresos=total_ingresos,
                         total_facturas=total_facturas,
                         dias=json.dumps(dias) if dias else '[]',
                         ventas_diarias=json.dumps(ventas_diarias) if ventas_diarias else '[]',
                         metodos=json.dumps(metodos) if metodos else '[]',
                         totales_metodos=json.dumps(totales_metodos) if totales_metodos else '[]')

@app.route('/reportes/caja')
@login_required
@verificar_acceso_reportes
def reporte_caja():
    """Reporte de caja"""
    fecha_inicio = request.args.get('fecha_inicio', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_fin = request.args.get('fecha_fin', datetime.now().strftime('%Y-%m-%d'))
    
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reportes'))
    
    try:
        cursor = conn.cursor()
        
        # 1. Cierres de caja en el período
        cursor.execute("""
            SELECT 
                c.*,
                CONCAT(e.nombre, ' ', e.apellido) as cajero
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            WHERE c.fecha BETWEEN %s AND %s
            ORDER BY c.fecha DESC
        """, (fecha_inicio, fecha_fin))
        
        cierres_caja = cursor.fetchall()
        
        # 2. Estadísticas de caja
        cursor.execute("""
            SELECT 
                COUNT(*) as total_dias,
                SUM(c.monto_apertura) as total_aperturas,
                SUM(c.monto_cierre) as total_cierres,
                SUM(c.venta_efectivo) as total_efectivo,
                SUM(c.venta_tarjeta) as total_tarjeta,
                SUM(c.venta_digital) as total_digital,
                SUM(c.total_ventas) as total_ventas,
                AVG(c.diferencia) as diferencia_promedio,
                SUM(CASE WHEN c.diferencia > 0 THEN c.diferencia ELSE 0 END) as total_sobrantes,
                SUM(CASE WHEN c.diferencia < 0 THEN ABS(c.diferencia) ELSE 0 END) as total_faltantes
            FROM caja_diaria c
            WHERE c.fecha BETWEEN %s AND %s
                AND c.estado = 'cerrada'
        """, (fecha_inicio, fecha_fin))
        
        estadisticas = cursor.fetchone()
        
        # 3. Resumen por día
        cursor.execute("""
            SELECT 
                c.fecha,
                SUM(c.venta_efectivo) as efectivo,
                SUM(c.venta_tarjeta) as tarjeta,
                SUM(c.venta_digital) as digital,
                SUM(c.total_ventas) as total
            FROM caja_diaria c
            WHERE c.fecha BETWEEN %s AND %s
                AND c.estado = 'cerrada'
            GROUP BY c.fecha
            ORDER BY c.fecha
        """, (fecha_inicio, fecha_fin))
        
        resumen_diario = cursor.fetchall()
        
        # 4. Top cajeros
        cursor.execute("""
            SELECT 
                e.id_empleado,
                CONCAT(e.nombre, ' ', e.apellido) as cajero,
                COUNT(c.id_caja) as dias_trabajados,
                SUM(c.total_ventas) as total_manejado,
                AVG(c.diferencia) as diferencia_promedio
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            WHERE c.fecha BETWEEN %s AND %s
                AND c.estado = 'cerrada'
            GROUP BY e.id_empleado, e.nombre, e.apellido
            ORDER BY total_manejado DESC
        """, (fecha_inicio, fecha_fin))
        
        top_cajeros = cursor.fetchall()
        
        # Preparar datos para gráficos
        fechas = []
        ventas_diarias = []
        
        for dia in resumen_diario:
            if dia['fecha']:
                fechas.append(dia['fecha'].strftime('%d/%m'))
                ventas_diarias.append(float(dia['total'] or 0))
        
    except Error as e:
        flash(f'Error generando reporte: {e}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reportes/caja.html',
                         fecha_inicio=fecha_inicio,
                         fecha_fin=fecha_fin,
                         cierres_caja=cierres_caja,
                         estadisticas=estadisticas,
                         resumen_diario=resumen_diario,
                         top_cajeros=top_cajeros,
                         fechas=json.dumps(fechas) if fechas else '[]',
                         ventas_diarias=json.dumps(ventas_diarias) if ventas_diarias else '[]')

@app.route('/reportes/empleados')
@login_required
@verificar_acceso_reportes
def reporte_empleados():
    """Reporte de rendimiento de empleados"""
    fecha_inicio = request.args.get('fecha_inicio', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_fin = request.args.get('fecha_fin', datetime.now().strftime('%Y-%m-%d'))
    
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reportes'))
    
    try:
        cursor = conn.cursor()
        
        # 1. Estadísticas generales de empleados
        cursor.execute("""
            SELECT 
                COUNT(*) as total_empleados,
                SUM(CASE WHEN activo = TRUE THEN 1 ELSE 0 END) as activos,
                SUM(CASE WHEN activo = FALSE THEN 1 ELSE 0 END) as inactivos,
                COUNT(DISTINCT especialidad) as especialidades
            FROM empleados
        """)
        
        estadisticas = cursor.fetchone()
        
        # 2. Rendimiento por empleado
        cursor.execute("""
            SELECT 
                e.id_empleado,
                CONCAT(e.nombre, ' ', e.apellido) as empleado,
                e.especialidad,
                COUNT(r.id_reserva) as total_reservas,
                SUM(CASE WHEN r.estado = 'completada' THEN 1 ELSE 0 END) as reservas_completadas,
                SUM(CASE WHEN f.estado = 'pagada' THEN f.total ELSE 0 END) as ingresos_generados
            FROM empleados e
            LEFT JOIN reservas r ON e.id_empleado = r.id_empleado
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            LEFT JOIN facturas f ON r.id_reserva = f.id_reserva
                AND DATE(f.fecha_emision) BETWEEN %s AND %s
            WHERE e.activo = TRUE
            GROUP BY e.id_empleado, e.nombre, e.apellido, e.especialidad
            ORDER BY ingresos_generados DESC
        """, (fecha_inicio, fecha_fin, fecha_inicio, fecha_fin))
        
        rendimiento_empleados = cursor.fetchall()
        
        # 3. Calcular métricas adicionales
        for empleado in rendimiento_empleados:
            total_reservas = empleado['total_reservas'] or 0
            completadas = empleado['reservas_completadas'] or 0
            
            if total_reservas > 0:
                empleado['tasa_exito'] = (completadas / total_reservas) * 100
            else:
                empleado['tasa_exito'] = 0
            
            # Calcular promedio por reserva
            ingresos = empleado['ingresos_generados'] or 0
            if completadas > 0:
                empleado['promedio_reserva'] = ingresos / completadas
            else:
                empleado['promedio_reserva'] = 0
        
        # 4. Distribución por especialidad
        cursor.execute("""
            SELECT 
                e.especialidad,
                COUNT(DISTINCT e.id_empleado) as cantidad_empleados,
                COUNT(r.id_reserva) as total_reservas,
                SUM(CASE WHEN r.estado = 'completada' THEN 1 ELSE 0 END) as reservas_completadas,
                SUM(f.total) as ingresos_especialidad
            FROM empleados e
            LEFT JOIN reservas r ON e.id_empleado = r.id_empleado
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            LEFT JOIN facturas f ON r.id_reserva = f.id_reserva
                AND f.estado = 'pagada'
                AND DATE(f.fecha_emision) BETWEEN %s AND %s
            WHERE e.activo = TRUE
            GROUP BY e.especialidad
            ORDER BY ingresos_especialidad DESC
        """, (fecha_inicio, fecha_fin, fecha_inicio, fecha_fin))
        
        por_especialidad = cursor.fetchall()
        
        # 5. Top empleados del mes
        cursor.execute("""
            SELECT 
                CONCAT(e.nombre, ' ', e.apellido) as empleado,
                COUNT(r.id_reserva) as reservas_mes,
                SUM(f.total) as ingresos_mes
            FROM empleados e
            LEFT JOIN reservas r ON e.id_empleado = r.id_empleado
                AND MONTH(r.fecha_reserva) = MONTH(CURRENT_DATE)
                AND YEAR(r.fecha_reserva) = YEAR(CURRENT_DATE)
            LEFT JOIN facturas f ON r.id_reserva = f.id_reserva
                AND f.estado = 'pagada'
                AND MONTH(f.fecha_emision) = MONTH(CURRENT_DATE)
                AND YEAR(f.fecha_emision) = YEAR(CURRENT_DATE)
            WHERE e.activo = TRUE
            GROUP BY e.id_empleado, e.nombre, e.apellido
            ORDER BY ingresos_mes DESC
            LIMIT 5
        """)
        
        top_mes = cursor.fetchall()
        
        # Preparar datos para gráficos
        nombres_empleados = []
        ingresos_empleados = []
        especialidades_nombres = []
        especialidades_cantidades = []
        
        for empleado in rendimiento_empleados[:10]:  # Top 10 empleados
            if empleado['empleado']:
                nombres_empleados.append(empleado['empleado'])
                ingresos_empleados.append(float(empleado['ingresos_generados'] or 0))
        
        for esp in por_especialidad:
            if esp['especialidad']:
                especialidades_nombres.append(esp['especialidad'])
                especialidades_cantidades.append(esp['cantidad_empleados'])
        
    except Error as e:
        flash(f'Error generando reporte: {e}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reportes/empleados.html',
                         fecha_inicio=fecha_inicio,
                         fecha_fin=fecha_fin,
                         estadisticas=estadisticas,
                         rendimiento_empleados=rendimiento_empleados,
                         por_especialidad=por_especialidad,
                         top_mes=top_mes,
                         nombres_empleados=json.dumps(nombres_empleados) if nombres_empleados else '[]',
                         ingresos_empleados=json.dumps(ingresos_empleados) if ingresos_empleados else '[]',
                         especialidades_nombres=json.dumps(especialidades_nombres) if especialidades_nombres else '[]',
                         especialidades_cantidades=json.dumps(especialidades_cantidades) if especialidades_cantidades else '[]')

@app.route('/reportes/servicios')
@login_required
@verificar_acceso_reportes
def reporte_servicios():
    """Reporte de servicios"""
    fecha_inicio = request.args.get('fecha_inicio', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_fin = request.args.get('fecha_fin', datetime.now().strftime('%Y-%m-%d'))
    
    conn = get_db_connection()
    if not conn:
        flash('No hay conexión a la base de datos.', 'danger')
        return redirect(url_for('reportes'))
    
    try:
        cursor = conn.cursor()
        
        # 1. Estadísticas generales de servicios
        cursor.execute("""
            SELECT 
                COUNT(*) as total_servicios,
                SUM(CASE WHEN activo = TRUE THEN 1 ELSE 0 END) as activos,
                SUM(CASE WHEN activo = FALSE THEN 1 ELSE 0 END) as inactivos,
                COUNT(DISTINCT categoria) as categorias,
                AVG(precio) as precio_promedio,
                MIN(precio) as precio_minimo,
                MAX(precio) as precio_maximo,
                SUM(precio - costo) as ganancia_potencial
            FROM servicios
        """)
        
        estadisticas = cursor.fetchone()
        
        # 2. Servicios más vendidos
        cursor.execute("""
            SELECT 
                s.id_servicio,
                s.nombre,
                s.categoria,
                s.precio,
                s.costo,
                (s.precio - s.costo) as margen,
                COUNT(rs.id_detalle) as veces_vendido,
                SUM(rs.cantidad) as cantidad_total,
                SUM(rs.subtotal) as ingresos_totales,
                (SUM(rs.subtotal) - (SUM(rs.cantidad) * s.costo)) as ganancia_real
            FROM servicios s
            LEFT JOIN reserva_servicios rs ON s.id_servicio = rs.id_servicio
            LEFT JOIN reservas r ON rs.id_reserva = r.id_reserva
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            WHERE s.activo = TRUE
            GROUP BY s.id_servicio, s.nombre, s.categoria, s.precio, s.costo
            ORDER BY ingresos_totales DESC
        """, (fecha_inicio, fecha_fin))
        
        servicios_vendidos = cursor.fetchall()
        
        # 3. Servicios por categoría
        cursor.execute("""
            SELECT 
                s.categoria,
                COUNT(DISTINCT s.id_servicio) as cantidad_servicios,
                COUNT(rs.id_detalle) as veces_vendido,
                SUM(rs.cantidad) as cantidad_total,
                SUM(rs.subtotal) as ingresos_categoria,
                AVG(s.precio) as precio_promedio_categoria,
                AVG(s.precio - s.costo) as margen_promedio
            FROM servicios s
            LEFT JOIN reserva_servicios rs ON s.id_servicio = rs.id_servicio
            LEFT JOIN reservas r ON rs.id_reserva = r.id_reserva
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            WHERE s.activo = TRUE
            GROUP BY s.categoria
            ORDER BY ingresos_categoria DESC
        """, (fecha_inicio, fecha_fin))
        
        por_categoria = cursor.fetchall()
        
        # 4. Servicios por mascota
        cursor.execute("""
            SELECT 
                m.especie,
                COUNT(rs.id_detalle) as servicios_realizados,
                SUM(rs.subtotal) as ingresos_especie,
                AVG(s.precio) as precio_promedio_especie
            FROM reserva_servicios rs
            JOIN servicios s ON rs.id_servicio = s.id_servicio
            JOIN reservas r ON rs.id_reserva = r.id_reserva
            JOIN mascotas m ON r.id_mascota = m.id_mascota
            WHERE DATE(r.fecha_reserva) BETWEEN %s AND %s
                AND s.activo = TRUE
            GROUP BY m.especie
            ORDER BY servicios_realizados DESC
        """, (fecha_inicio, fecha_fin))
        
        servicios_por_especie = cursor.fetchall()
        
        # Preparar datos para gráficos
        categorias_nombres = []
        categorias_ingresos = []
        top_servicios_nombres = []
        top_servicios_cantidades = []
        
        for cat in por_categoria:
            if cat['categoria']:
                categorias_nombres.append(cat['categoria'])
                categorias_ingresos.append(float(cat['ingresos_categoria'] or 0))
        
        for servicio in servicios_vendidos[:10]:  # Top 10 servicios
            if servicio['nombre']:
                top_servicios_nombres.append(servicio['nombre'][:20])  # Limitar longitud
                top_servicios_cantidades.append(servicio['veces_vendido'] or 0)
        
    except Error as e:
        flash(f'Error generando reporte: {e}', 'danger')
        return redirect(url_for('reportes'))
    finally:
        cursor.close()
        conn.close()
    
    return render_template('reportes/servicios.html',
                         fecha_inicio=fecha_inicio,
                         fecha_fin=fecha_fin,
                         estadisticas=estadisticas,
                         servicios_vendidos=servicios_vendidos,
                         por_categoria=por_categoria,
                         servicios_por_especie=servicios_por_especie,
                         categorias_nombres=json.dumps(categorias_nombres) if categorias_nombres else '[]',
                         categorias_ingresos=json.dumps(categorias_ingresos) if categorias_ingresos else '[]',
                         top_servicios_nombres=json.dumps(top_servicios_nombres) if top_servicios_nombres else '[]',
                         top_servicios_cantidades=json.dumps(top_servicios_cantidades) if top_servicios_cantidades else '[]')

# ==================== APIs para estadísticas ====================

@app.route('/api/reportes/estadisticas-dia')
@login_required
def api_estadisticas_dia():
    """API para obtener estadísticas del día actual"""
    conn = get_db_connection()
    if not conn:
        return jsonify({'success': False, 'error': 'No hay conexión a la base de datos'}), 500
    
    try:
        cursor = conn.cursor()
        
        # Ventas del día
        cursor.execute("""
            SELECT COALESCE(SUM(total), 0) as ventas_hoy
            FROM facturas
            WHERE DATE(fecha_emision) = CURRENT_DATE
                AND estado = 'pagada'
        """)
        ventas = cursor.fetchone()
        
        # Reservas del día
        cursor.execute("""
            SELECT COUNT(*) as reservas_hoy
            FROM reservas
            WHERE DATE(fecha_reserva) = CURRENT_DATE
        """)
        reservas = cursor.fetchone()
        
        # Servicios del día
        cursor.execute("""
            SELECT COUNT(*) as servicios_hoy
            FROM reserva_servicios rs
            JOIN reservas r ON rs.id_reserva = r.id_reserva
            WHERE DATE(r.fecha_reserva) = CURRENT_DATE
        """)
        servicios = cursor.fetchone()
        
        # Empleados activos
        cursor.execute("""
            SELECT COUNT(*) as empleados_activos
            FROM empleados
            WHERE activo = TRUE
        """)
        empleados = cursor.fetchone()
        
        return jsonify({
            'success': True,
            'ventas_hoy': float(ventas['ventas_hoy'] or 0),
            'reservas_hoy': int(reservas['reservas_hoy'] or 0),
            'servicios_hoy': int(servicios['servicios_hoy'] or 0),
            'empleados_activos': int(empleados['empleados_activos'] or 0)
        })
        
    except Error as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

# ==================== Exportación de reportes ====================

@app.route('/api/reportes/exportar/<tipo>')
@login_required
@admin_required
def exportar_reporte(tipo):
    """Exportar reporte a diferentes formatos"""
    formatos = ['excel', 'pdf', 'word']
    
    if tipo not in formatos:
        return jsonify({'success': False, 'error': 'Formato no válido'}), 400
    
    # Obtener parámetros
    reporte = request.args.get('reporte', 'ventas')
    fecha_inicio = request.args.get('fecha_inicio', '')
    fecha_fin = request.args.get('fecha_fin', '')
    
    try:
        if tipo == 'excel':
            if reporte == 'ventas':
                return exportar_excel(reporte, fecha_inicio, fecha_fin)
            elif reporte == 'caja':
                return exportar_excel_caja(fecha_inicio, fecha_fin)
            elif reporte == 'empleados':
                return exportar_excel_empleados(fecha_inicio, fecha_fin)
            elif reporte == 'servicios':
                return exportar_excel_servicios(fecha_inicio, fecha_fin)
            else:
                return jsonify({'success': False, 'error': 'Reporte no válido'}), 400
                
        elif tipo == 'pdf':
            # Para PDF y Word, puedes usar versiones simplificadas como hicimos para ventas
            # O crear funciones específicas para cada reporte
            return exportar_pdf_simple(reporte, fecha_inicio, fecha_fin)
            
        elif tipo == 'word':
            return exportar_word_simple(reporte, fecha_inicio, fecha_fin)
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    
    return redirect(url_for('reportes'))

def exportar_excel(reporte, fecha_inicio, fecha_fin):
    """Exportar a Excel - IMPLEMENTACIÓN COMPLETA"""
    try:
        import pandas as pd
        from io import BytesIO
        from flask import send_file
        from datetime import datetime
        
        # Obtener datos según el tipo de reporte
        conn = get_db_connection()
        if not conn:
            flash('Error de conexión.', 'danger')
            return redirect(url_for('reportes'))
        
        cursor = conn.cursor()
        
        if reporte == 'ventas':
            # Datos principales de ventas
            cursor.execute("""
                SELECT 
                    f.numero as 'N° Factura',
                    f.tipo_comprobante as 'Tipo',
                    DATE(f.fecha_emision) as 'Fecha',
                    f.metodo_pago as 'Método Pago',
                    CONCAT(c.nombre, ' ', c.apellido) as 'Cliente',
                    f.subtotal as 'Subtotal',
                    f.igv as 'IGV',
                    f.total as 'Total',
                    f.estado as 'Estado',
                    CASE 
                        WHEN f.estado = 'pagada' THEN 'Pagado'
                        WHEN f.estado = 'pendiente' THEN 'Pendiente'
                        WHEN f.estado = 'credito' THEN 'Crédito'
                        ELSE f.estado
                    END as 'Estado Texto'
                FROM facturas f
                LEFT JOIN clientes c ON f.id_cliente = c.id_cliente
                WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                ORDER BY f.fecha_emision DESC
            """, (fecha_inicio, fecha_fin))
            
            datos_principales = cursor.fetchall()
            
            # Resumen por método de pago
            cursor.execute("""
                SELECT 
                    COALESCE(metodo_pago, 'No especificado') as 'Método de Pago',
                    COUNT(*) as 'Cantidad Facturas',
                    SUM(total) as 'Total Recaudado',
                    ROUND(AVG(total), 2) as 'Promedio por Factura'
                FROM facturas
                WHERE DATE(fecha_emision) BETWEEN %s AND %s
                    AND estado = 'pagada'
                GROUP BY metodo_pago
                ORDER BY SUM(total) DESC
            """, (fecha_inicio, fecha_fin))
            
            resumen_metodos = cursor.fetchall()
            
            # Resumen por día
            cursor.execute("""
                SELECT 
                    DATE(fecha_emision) as 'Fecha',
                    COUNT(*) as 'Facturas del Día',
                    SUM(total) as 'Total del Día',
                    ROUND(AVG(total), 2) as 'Promedio del Día',
                    SUM(CASE WHEN metodo_pago = 'efectivo' THEN total ELSE 0 END) as 'Efectivo',
                    SUM(CASE WHEN metodo_pago = 'tarjeta' THEN total ELSE 0 END) as 'Tarjeta',
                    SUM(CASE WHEN metodo_pago IN ('yape', 'plin', 'transferencia') THEN total ELSE 0 END) as 'Digital'
                FROM facturas
                WHERE DATE(fecha_emision) BETWEEN %s AND %s
                    AND estado = 'pagada'
                GROUP BY DATE(fecha_emision)
                ORDER BY DATE(fecha_emision)
            """, (fecha_inicio, fecha_fin))
            
            resumen_diario = cursor.fetchall()
            
            # Top servicios
            cursor.execute("""
                SELECT 
                    s.nombre as 'Servicio',
                    s.categoria as 'Categoría',
                    COUNT(fs.id_detalle) as 'Veces Vendido',
                    SUM(fs.cantidad) as 'Cantidad Total',
                    SUM(fs.subtotal) as 'Ingresos Totales',
                    ROUND(AVG(fs.precio_unitario), 2) as 'Precio Promedio'
                FROM factura_servicios fs
                JOIN servicios s ON fs.id_servicio = s.id_servicio
                JOIN facturas f ON fs.id_factura = f.id_factura
                WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                    AND f.estado = 'pagada'
                GROUP BY s.id_servicio, s.nombre, s.categoria
                ORDER BY SUM(fs.subtotal) DESC
                LIMIT 20
            """, (fecha_inicio, fecha_fin))
            
            top_servicios = cursor.fetchall()
            
        cursor.close()
        conn.close()
        
        # Crear archivo Excel en memoria
        output = BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Hoja 1: Detalle de ventas
            if datos_principales:
                df_principales = pd.DataFrame(datos_principales)
                df_principales.to_excel(writer, sheet_name='Ventas Detalladas', index=False)
            
            # Hoja 2: Resumen por método de pago
            if resumen_metodos:
                df_metodos = pd.DataFrame(resumen_metodos)
                df_metodos.to_excel(writer, sheet_name='Por Método de Pago', index=False)
            
            # Hoja 3: Resumen diario
            if resumen_diario:
                df_diario = pd.DataFrame(resumen_diario)
                df_diario.to_excel(writer, sheet_name='Resumen Diario', index=False)
            
            # Hoja 4: Top servicios
            if top_servicios:
                df_servicios = pd.DataFrame(top_servicios)
                df_servicios.to_excel(writer, sheet_name='Servicios Más Vendidos', index=False)
            
            # Hoja 5: Estadísticas generales
            estadisticas_data = [{
                'Período': f'{fecha_inicio} al {fecha_fin}',
                'Total Facturas': len(datos_principales) if datos_principales else 0,
                'Total Recaudado': sum(float(d['Total'] or 0) for d in datos_principales if d.get('Estado') == 'pagada'),
                'Total Pendiente': sum(float(d['Total'] or 0) for d in datos_principales if d.get('Estado') == 'pendiente'),
                'Facturas Pagadas': sum(1 for d in datos_principales if d.get('Estado') == 'pagada'),
                'Facturas Pendientes': sum(1 for d in datos_principales if d.get('Estado') == 'pendiente'),
                'Fecha Generación': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }]
            
            df_estadisticas = pd.DataFrame(estadisticas_data)
            df_estadisticas.to_excel(writer, sheet_name='Estadísticas', index=False)
        
        output.seek(0)
        
        # Devolver archivo
        nombre_archivo = f"reporte_ventas_{fecha_inicio}_{fecha_fin}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except ImportError:
        flash('Error: Para exportar a Excel necesitas instalar pandas y openpyxl.', 'danger')
        return redirect(url_for('reporte_ventas'))
    except Exception as e:
        flash(f'Error exportando a Excel: {str(e)}', 'danger')
        return redirect(url_for('reporte_ventas'))

def exportar_pdf(reporte, fecha_inicio, fecha_fin):
    """Exportar a PDF usando FPDF - VERSIÓN MEJORADA"""
    try:
        from fpdf import FPDF
        from io import BytesIO
        from flask import send_file
        import datetime as dt
        
        # Obtener datos
        conn = get_db_connection()
        if not conn:
            return jsonify({'success': False, 'error': 'Error de conexión a la base de datos'}), 500
        
        cursor = conn.cursor()
        
        # Datos principales
        cursor.execute("""
            SELECT 
                f.numero,
                f.tipo_comprobante,
                DATE(f.fecha_emision) as fecha,
                f.metodo_pago,
                CONCAT(c.nombre, ' ', c.apellido) as cliente,
                f.subtotal,
                f.igv,
                f.total,
                f.estado
            FROM facturas f
            LEFT JOIN clientes c ON f.id_cliente = c.id_cliente
            WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                AND f.estado = 'pagada'
            ORDER BY f.fecha_emision DESC
        """, (fecha_inicio, fecha_fin))
        
        ventas = cursor.fetchall()
        
        # Métodos de pago
        cursor.execute("""
            SELECT 
                COALESCE(metodo_pago, 'No especificado') as metodo_pago,
                COUNT(*) as cantidad,
                SUM(total) as total
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
            GROUP BY metodo_pago
            ORDER BY SUM(total) DESC
        """, (fecha_inicio, fecha_fin))
        
        metodos = cursor.fetchall()
        
        # Estadísticas
        cursor.execute("""
            SELECT 
                COUNT(*) as total_facturas,
                SUM(total) as total_ingresos,
                AVG(total) as promedio_venta,
                MIN(total) as venta_minima,
                MAX(total) as venta_maxima
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
        """, (fecha_inicio, fecha_fin))
        
        estadisticas = cursor.fetchone()
        
        cursor.close()
        conn.close()
        
        # Crear PDF con diseño profesional
        pdf = FPDF(orientation='L')  # Horizontal
        pdf.add_page()
        
        # Configuración de márgenes
        pdf.set_margins(10, 10, 10)
        pdf.set_auto_page_break(True, margin=15)
        
        # Título con estilo
        pdf.set_font('Arial', 'B', 18)
        pdf.set_text_color(102, 126, 234)  # Color PetGlow
        pdf.cell(0, 15, 'REPORTE DE VENTAS', 0, 1, 'C')
        pdf.set_font('Arial', 'B', 12)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, 'PetGlow Peluquería Canina', 0, 1, 'C')
        
        # Información del período
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 6, f'Período: {fecha_inicio} al {fecha_fin}', 0, 1, 'C')
        pdf.cell(0, 6, f'Generado: {dt.datetime.now().strftime("%d/%m/%Y %H:%M:%S")}', 0, 1, 'C')
        pdf.ln(5)
        
        # Línea decorativa
        pdf.set_draw_color(102, 126, 234)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 280, pdf.get_y())
        pdf.ln(8)
        
        # Estadísticas generales en una tabla
        pdf.set_font('Arial', 'B', 12)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 10, 'ESTADÍSTICAS GENERALES', 0, 1, 'L', True)
        pdf.ln(2)
        
        pdf.set_font('Arial', '', 10)
        col_widths = [60, 40, 60, 40, 40]
        headers = ['Total Facturas', 'Ingresos Totales', 'Ticket Promedio', 'Venta Mínima', 'Venta Máxima']
        valores = [
            f"{estadisticas['total_facturas'] or 0}",
            f"S/ {float(estadisticas['total_ingresos'] or 0):,.2f}",
            f"S/ {float(estadisticas['promedio_venta'] or 0):,.2f}",
            f"S/ {float(estadisticas['venta_minima'] or 0):,.2f}",
            f"S/ {float(estadisticas['venta_maxima'] or 0):,.2f}"
        ]
        
        for i, (header, valor) in enumerate(zip(headers, valores)):
            pdf.set_font('Arial', 'B', 9)
            pdf.set_fill_color(245, 245, 245)
            pdf.cell(col_widths[i], 8, header, 1, 0, 'L', True)
        pdf.ln()
        
        for i, valor in enumerate(valores):
            pdf.set_font('Arial', '', 9)
            pdf.set_fill_color(255, 255, 255)
            pdf.cell(col_widths[i], 8, valor, 1, 0, 'C', True)
        pdf.ln(10)
        
        # Métodos de pago
        pdf.set_font('Arial', 'B', 12)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 10, 'DISTRIBUCIÓN POR MÉTODO DE PAGO', 0, 1, 'L', True)
        pdf.ln(2)
        
        # Tabla de métodos de pago
        col_widths_metodos = [80, 50, 60]
        pdf.set_font('Arial', 'B', 10)
        pdf.set_fill_color(118, 75, 162)  # Color morado
        pdf.set_text_color(255, 255, 255)
        
        headers_metodos = ['Método de Pago', 'Cantidad', 'Total (S/)']
        for i, header in enumerate(headers_metodos):
            pdf.cell(col_widths_metodos[i], 8, header, 1, 0, 'C', True)
        pdf.ln()
        
        pdf.set_text_color(0, 0, 0)
        total_general = 0
        for i, metodo in enumerate(metodos):
            # Alternar colores de fondo
            if i % 2 == 0:
                pdf.set_fill_color(255, 255, 255)
            else:
                pdf.set_fill_color(248, 249, 250)
            
            pdf.set_font('Arial', '', 9)
            pdf.cell(col_widths_metodos[0], 8, metodo['metodo_pago'], 1, 0, 'L', True)
            pdf.cell(col_widths_metodos[1], 8, str(metodo['cantidad']), 1, 0, 'C', True)
            pdf.cell(col_widths_metodos[2], 8, f"{float(metodo['total'] or 0):,.2f}", 1, 0, 'R', True)
            pdf.ln()
            total_general += float(metodo['total'] or 0)
        
        # Total general
        pdf.set_font('Arial', 'B', 10)
        pdf.set_fill_color(67, 233, 123)  # Color verde
        pdf.set_text_color(255, 255, 255)
        pdf.cell(col_widths_metodos[0] + col_widths_metodos[1], 8, 'TOTAL GENERAL', 1, 0, 'R', True)
        pdf.cell(col_widths_metodos[2], 8, f"S/ {total_general:,.2f}", 1, 0, 'R', True)
        pdf.ln(12)
        
        # Detalle de ventas
        if ventas:
            pdf.add_page()
            pdf.set_font('Arial', 'B', 12)
            pdf.set_fill_color(240, 240, 240)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, 'DETALLE DE VENTAS', 0, 1, 'L', True)
            pdf.ln(2)
            
            # Encabezados de la tabla de ventas
            col_widths_ventas = [30, 30, 25, 60, 30, 30, 30, 25]
            headers_ventas = ['N° Factura', 'Tipo', 'Fecha', 'Cliente', 'Método', 'Subtotal', 'IGV', 'Total']
            
            pdf.set_fill_color(245, 111, 108)  # Color rojo
            pdf.set_text_color(255, 255, 255)
            pdf.set_font('Arial', 'B', 8)
            
            for i, header in enumerate(headers_ventas):
                pdf.cell(col_widths_ventas[i], 8, header, 1, 0, 'C', True)
            pdf.ln()
            
            # Datos de ventas
            pdf.set_text_color(0, 0, 0)
            for i, venta in enumerate(ventas[:80]):  # Limitar a 80 registros
                if i % 2 == 0:
                    pdf.set_fill_color(255, 255, 255)
                else:
                    pdf.set_fill_color(248, 249, 250)
                
                pdf.set_font('Arial', '', 7)
                pdf.cell(col_widths_ventas[0], 8, venta['numero'][:12], 1, 0, 'L', True)
                pdf.cell(col_widths_ventas[1], 8, venta['tipo_comprobante'][:1].upper(), 1, 0, 'C', True)
                pdf.cell(col_widths_ventas[2], 8, str(venta['fecha']), 1, 0, 'C', True)
                pdf.cell(col_widths_ventas[3], 8, (venta['cliente'] or 'N/A')[:25], 1, 0, 'L', True)
                pdf.cell(col_widths_ventas[4], 8, venta['metodo_pago'][:8], 1, 0, 'C', True)
                pdf.cell(col_widths_ventas[5], 8, f"{float(venta['subtotal'] or 0):,.2f}", 1, 0, 'R', True)
                pdf.cell(col_widths_ventas[6], 8, f"{float(venta['igv'] or 0):,.2f}", 1, 0, 'R', True)
                pdf.cell(col_widths_ventas[7], 8, f"{float(venta['total'] or 0):,.2f}", 1, 0, 'R', True)
                pdf.ln()
        
        # Pie de página
        pdf.set_y(-20)
        pdf.set_font('Arial', 'I', 8)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 10, f'Página {pdf.page_no()}', 0, 0, 'C')
        pdf.ln(5)
        pdf.cell(0, 10, 'PetGlow Peluquería Canina - Sistema de Gestión Integral', 0, 0, 'C')
        
        # Guardar en buffer
        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin1')
        buffer.write(pdf_output)
        buffer.seek(0)
        
        # Enviar archivo
        nombre_archivo = f"reporte_ventas_{fecha_inicio}_{fecha_fin}.pdf"
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except Exception as e:
        flash(f'Error exportando a PDF: {str(e)}', 'danger')
        return redirect(url_for('reporte_ventas'))

def exportar_word(reporte, fecha_inicio, fecha_fin):
    """Exportar a Word - VERSIÓN SIMPLIFICADA"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from flask import send_file
        
        # Obtener datos
        conn = get_db_connection()
        if not conn:
            flash('Error de conexión.', 'danger')
            return redirect(url_for('reporte_ventas'))
        
        cursor = conn.cursor()
        
        # 1. Estadísticas generales
        cursor.execute("""
            SELECT 
                COUNT(*) as total_facturas,
                SUM(total) as total_ingresos,
                AVG(total) as promedio_venta,
                MIN(total) as venta_minima,
                MAX(total) as venta_maxima
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
        """, (fecha_inicio, fecha_fin))
        
        estadisticas = cursor.fetchone()
        
        # 2. Métodos de pago
        cursor.execute("""
            SELECT 
                COALESCE(metodo_pago, 'No especificado') as metodo_pago,
                COUNT(*) as cantidad,
                SUM(total) as total
            FROM facturas
            WHERE DATE(fecha_emision) BETWEEN %s AND %s
                AND estado = 'pagada'
            GROUP BY metodo_pago
            ORDER BY SUM(total) DESC
        """, (fecha_inicio, fecha_fin))
        
        metodos = cursor.fetchall()
        
        # 3. Top 10 clientes
        cursor.execute("""
            SELECT 
                c.nombre,
                c.apellido,
                COUNT(f.id_factura) as cantidad_facturas,
                SUM(f.total) as total_gastado
            FROM facturas f
            JOIN clientes c ON f.id_cliente = c.id_cliente
            WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                AND f.estado = 'pagada'
            GROUP BY c.id_cliente, c.nombre, c.apellido
            ORDER BY total_gastado DESC
            LIMIT 10
        """, (fecha_inicio, fecha_fin))
        
        clientes = cursor.fetchall()
        
        # 4. Top 10 servicios
        cursor.execute("""
            SELECT 
                s.nombre,
                COUNT(fs.id_detalle) as veces_vendido,
                SUM(fs.subtotal) as ingresos_totales
            FROM factura_servicios fs
            JOIN servicios s ON fs.id_servicio = s.id_servicio
            JOIN facturas f ON fs.id_factura = f.id_factura
            WHERE DATE(f.fecha_emision) BETWEEN %s AND %s
                AND f.estado = 'pagada'
            GROUP BY s.id_servicio, s.nombre
            ORDER BY SUM(fs.subtotal) DESC
            LIMIT 10
        """, (fecha_inicio, fecha_fin))
        
        servicios = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Crear documento Word MUY SIMPLE
        doc = Document()
        
        # Título
        title = doc.add_heading('REPORTE DE VENTAS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitle = doc.add_heading('PetGlow Peluquería Canina', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del período
        doc.add_paragraph(f'Período: {fecha_inicio} al {fecha_fin}')
        doc.add_paragraph(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph()
        
        # 1. Estadísticas generales
        doc.add_heading('ESTADÍSTICAS GENERALES', level=1)
        
        stats_data = [
            ('Total Facturas', f"{estadisticas['total_facturas'] or 0}"),
            ('Ingresos Totales', f"S/ {float(estadisticas['total_ingresos'] or 0):,.2f}"),
            ('Ticket Promedio', f"S/ {float(estadisticas['promedio_venta'] or 0):,.2f}"),
            ('Venta Mínima', f"S/ {float(estadisticas['venta_minima'] or 0):,.2f}"),
            ('Venta Máxima', f"S/ {float(estadisticas['venta_maxima'] or 0):,.2f}")
        ]
        
        for label, value in stats_data:
            p = doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(value)
        
        doc.add_paragraph()
        
        # 2. Métodos de pago
        doc.add_heading('DISTRIBUCIÓN POR MÉTODO DE PAGO', level=1)
        
        if metodos:
            table = doc.add_table(rows=len(metodos) + 1, cols=3)
            table.style = 'Table Grid'
            
            # Encabezado
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Método'
            hdr_cells[1].text = 'Cantidad'
            hdr_cells[2].text = 'Total (S/)'
            
            # Datos
            for i, metodo in enumerate(metodos, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = metodo['metodo_pago']
                row_cells[1].text = str(metodo['cantidad'])
                row_cells[2].text = f"{float(metodo['total'] or 0):,.2f}"
        else:
            doc.add_paragraph('No hay datos de métodos de pago')
        
        doc.add_paragraph()
        
        # 3. Top clientes
        doc.add_heading('TOP 10 CLIENTES', level=1)
        
        if clientes:
            table_clientes = doc.add_table(rows=len(clientes) + 1, cols=3)
            table_clientes.style = 'Table Grid'
            
            # Encabezado
            hdr_cells = table_clientes.rows[0].cells
            hdr_cells[0].text = 'Cliente'
            hdr_cells[1].text = 'Facturas'
            hdr_cells[2].text = 'Total Gastado (S/)'
            
            # Datos
            for i, cliente in enumerate(clientes, 1):
                row_cells = table_clientes.rows[i].cells
                row_cells[0].text = f"{cliente['nombre']} {cliente['apellido']}"
                row_cells[1].text = str(cliente['cantidad_facturas'])
                row_cells[2].text = f"{float(cliente['total_gastado'] or 0):,.2f}"
        else:
            doc.add_paragraph('No hay datos de clientes')
        
        doc.add_paragraph()
        
        # 4. Top servicios
        doc.add_heading('TOP 10 SERVICIOS', level=1)
        
        if servicios:
            table_servicios = doc.add_table(rows=len(servicios) + 1, cols=3)
            table_servicios.style = 'Table Grid'
            
            # Encabezado
            hdr_cells = table_servicios.rows[0].cells
            hdr_cells[0].text = 'Servicio'
            hdr_cells[1].text = 'Veces Vendido'
            hdr_cells[2].text = 'Ingresos Totales (S/)'
            
            # Datos
            for i, servicio in enumerate(servicios, 1):
                row_cells = table_servicios.rows[i].cells
                row_cells[0].text = servicio['nombre']
                row_cells[1].text = str(servicio['veces_vendido'])
                row_cells[2].text = f"{float(servicio['ingresos_totales'] or 0):,.2f}"
        else:
            doc.add_paragraph('No hay datos de servicios')
        
        # Pie de página
        doc.add_page_break()
        doc.add_paragraph('---')
        doc.add_paragraph('PetGlow Peluquería Canina')
        doc.add_paragraph('Sistema de Gestión Integral')
        doc.add_paragraph(f'Documento generado el {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        
        # Guardar en buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Enviar archivo
        nombre_archivo = f"reporte_ventas_{fecha_inicio}_{fecha_fin}.docx"
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except ImportError:
        flash('Error: Para exportar a Word necesitas instalar python-docx.', 'danger')
        return redirect(url_for('reporte_ventas'))
    except Exception as e:
        flash(f'Error exportando a Word: {str(e)}', 'danger')
        return redirect(url_for('reporte_ventas'))

def exportar_excel_caja(fecha_inicio, fecha_fin):
    """Exportar reporte de caja a Excel"""
    try:
        import pandas as pd
        from io import BytesIO
        from flask import send_file
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Cierres de caja
        cursor.execute("""
            SELECT 
                c.fecha,
                CONCAT(e.nombre, ' ', e.apellido) as cajero,
                c.monto_apertura,
                c.venta_efectivo,
                c.venta_tarjeta,
                c.venta_digital,
                c.total_ventas,
                c.diferencia,
                c.estado
            FROM caja_diaria c
            JOIN empleados e ON c.id_empleado_cajero = e.id_empleado
            WHERE c.fecha BETWEEN %s AND %s
            ORDER BY c.fecha DESC
        """, (fecha_inicio, fecha_fin))
        
        cierres = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Crear DataFrame
        df = pd.DataFrame(cierres)
        
        # Crear archivo Excel
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Cierres de Caja', index=False)
        
        output.seek(0)
        
        nombre_archivo = f"reporte_caja_{fecha_inicio}_{fecha_fin}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except Exception as e:
        flash(f'Error exportando a Excel: {str(e)}', 'danger')
        return redirect(url_for('reporte_caja'))

def exportar_excel_empleados(fecha_inicio, fecha_fin):
    """Exportar reporte de empleados a Excel"""
    try:
        import pandas as pd
        from io import BytesIO
        from flask import send_file
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Rendimiento de empleados
        cursor.execute("""
            SELECT 
                CONCAT(e.nombre, ' ', e.apellido) as empleado,
                e.especialidad,
                COUNT(r.id_reserva) as total_reservas,
                SUM(CASE WHEN r.estado = 'completada' THEN 1 ELSE 0 END) as reservas_completadas,
                SUM(CASE WHEN f.estado = 'pagada' THEN f.total ELSE 0 END) as ingresos_generados
            FROM empleados e
            LEFT JOIN reservas r ON e.id_empleado = r.id_empleado
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            LEFT JOIN facturas f ON r.id_reserva = f.id_reserva
                AND DATE(f.fecha_emision) BETWEEN %s AND %s
            WHERE e.activo = TRUE
            GROUP BY e.id_empleado, e.nombre, e.apellido, e.especialidad
            ORDER BY ingresos_generados DESC
        """, (fecha_inicio, fecha_fin, fecha_inicio, fecha_fin))
        
        empleados = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Calcular métricas adicionales
        for emp in empleados:
            total = emp['total_reservas'] or 0
            completadas = emp['reservas_completadas'] or 0
            ingresos = emp['ingresos_generados'] or 0
            
            if total > 0:
                emp['tasa_exito'] = (completadas / total) * 100
            else:
                emp['tasa_exito'] = 0
            
            if completadas > 0:
                emp['promedio_reserva'] = ingresos / completadas
            else:
                emp['promedio_reserva'] = 0
        
        df = pd.DataFrame(empleados)
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Empleados', index=False)
        
        output.seek(0)
        
        nombre_archivo = f"reporte_empleados_{fecha_inicio}_{fecha_fin}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except Exception as e:
        flash(f'Error exportando a Excel: {str(e)}', 'danger')
        return redirect(url_for('reporte_empleados'))

def exportar_excel_servicios(fecha_inicio, fecha_fin):
    """Exportar reporte de servicios a Excel"""
    try:
        import pandas as pd
        from io import BytesIO
        from flask import send_file
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Servicios más vendidos
        cursor.execute("""
            SELECT 
                s.nombre,
                s.categoria,
                s.precio,
                s.costo,
                COUNT(rs.id_detalle) as veces_vendido,
                SUM(rs.cantidad) as cantidad_total,
                SUM(rs.subtotal) as ingresos_totales
            FROM servicios s
            LEFT JOIN reserva_servicios rs ON s.id_servicio = rs.id_servicio
            LEFT JOIN reservas r ON rs.id_reserva = r.id_reserva
                AND DATE(r.fecha_reserva) BETWEEN %s AND %s
            WHERE s.activo = TRUE
            GROUP BY s.id_servicio, s.nombre, s.categoria, s.precio, s.costo
            ORDER BY ingresos_totales DESC
        """, (fecha_inicio, fecha_fin))
        
        servicios = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        # Calcular ganancia y margen
        for serv in servicios:
            ingresos = serv['ingresos_totales'] or 0
            costo_total = (serv['cantidad_total'] or 0) * (serv['costo'] or 0)
            serv['ganancia'] = ingresos - costo_total
            
            if serv['costo'] > 0:
                serv['margen'] = ((serv['precio'] - serv['costo']) / serv['costo']) * 100
            else:
                serv['margen'] = 0
        
        df = pd.DataFrame(servicios)
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Servicios', index=False)
        
        output.seek(0)
        
        nombre_archivo = f"reporte_servicios_{fecha_inicio}_{fecha_fin}.xlsx"
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=nombre_archivo
        )
        
    except Exception as e:
        flash(f'Error exportando a Excel: {str(e)}', 'danger')
        return redirect(url_for('reporte_servicios'))


# ================= EJECUCIÓN =================

if __name__ == '__main__':
    print("=" * 50)
    print("🚀 PETGLOW WEB SYSTEM - SIN LOGIN")
    print("=" * 50)
    print("📊 URL: http://localhost:5000")
    print("👤 Acceso directo al dashboard")
    print("=" * 50)
    
    # Verificar conexión a BD
    conn = get_db_connection()
    if conn:
        print("✅ Conectado a MySQL")
        conn.close()
    else:
        print("⚠️  Modo DEMO - Sin conexión a MySQL")
        print("   Configura tu .env con las credenciales correctas")
    
    print("=" * 50)
    
    app.run(debug=True, host='0.0.0.0', port=5000)


